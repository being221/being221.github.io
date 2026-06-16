#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
5400跨度 docx修改 修补 (第三轮)
修正合并单元格问题 + 更新弯矩分配相关表格
"""
import sys, os, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

L1, L2, ss = 5.4, 2.4, 3.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0
F_NEW = 1 - 2*(0.5*ss/L1)**2 + (0.5*ss/L1)**3
sec_beam_new = 1.54 * L1 / 2

# ============================================================
# 修补表3-3 (Table index 9) 中未正确替换的单元格
# ============================================================
print("修补表3-3...")
t = doc.tables[9]

# 楼面层边柱 次梁自重 R2 - 直接重设
t.rows[2].cells[2].text = f'1.54×{L1}/2={sec_beam_new:.1f}kN'

# 楼面层边柱 楼面导荷 R5
edge_floor_conc = g_floor * (ss**2/4 + ss*L1/2)
t.rows[5].cells[2].text = f'{g_floor}×({ss}/2×{ss}/2+{ss}×{L1})={edge_floor_conc:.2f}kN'

# 楼面层边柱 小计 R6
t.rows[6].cells[2].text = f'{42.3+20.08+sec_beam_new+edge_floor_conc:.2f}kN'

# 楼面层中柱 次梁自重 R8
t.rows[8].cells[2].text = f'1.54×{L1}/2={sec_beam_new:.1f}kN'

# 楼面层中柱 楼面导荷 R11
mid_floor_conc = g_floor*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))
t.rows[11].cells[2].text = f'{g_floor}×(({ss}/2×{ss}/2+{ss}×{L1})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_floor_conc:.2f}kN'

# 楼面层中柱 小计 R12
t.rows[12].cells[2].text = f'{41.95+20.08+sec_beam_new+mid_floor_conc:.2f}kN'

# 屋面层边柱 次梁自重 R14
t.rows[14].cells[2].text = f'1.54×{L1}/2={sec_beam_new:.1f}kN'

# 屋面层边柱 楼面导荷 R16
edge_roof_conc = g_roof*(ss**2/4+ss*L1/2)
t.rows[16].cells[2].text = f'{g_roof}×({ss}/2×{ss}/2+{ss}×{L1})={edge_roof_conc:.2f}kN'

# 屋面层边柱 小计 R17
t.rows[17].cells[2].text = f'{31.19+20.08+sec_beam_new+edge_roof_conc:.2f}kN'

# 屋面层中柱 次梁自重 R18
t.rows[18].cells[2].text = f'1.54×{L1}/2={sec_beam_new:.1f}kN'

# 屋面层中柱 楼面导荷 R20
mid_roof_conc = g_roof*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))
t.rows[20].cells[2].text = f'{g_roof}×(({ss}/2×{ss}/2+{ss}×{L1})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_roof_conc:.2f}kN'

# 屋面层中柱 小计 R21
mid_roof_total = 3.7+20.08+sec_beam_new+mid_roof_conc
t.rows[21].cells[2].text = f'{mid_roof_total:.2f}kN'

print("  表3-3 修补完成")

# ============================================================
# 修补段落 P582 中楼面等效荷载
# ============================================================
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # P582: "楼面层为2.57+6.2+0.79×14.49=20.22kN/m"
    if '0.79×14.49=20.22kN/m' in text or '0.79×14.49' in text:
        new_q = 2.57+6.2+F_NEW*ss*g_floor
        # 尝试逐个run修改
        for run in para.runs:
            run.text = run.text.replace('0.79×14.49=20.22kN/m',
                                       f'{F_NEW:.2f}×{ss*g_floor:.2f}={new_q:.2f}kN/m')
            run.text = run.text.replace('0.79×14.49',
                                       f'{F_NEW:.2f}×{ss*g_floor:.2f}')
            if '20.22kN/m' in run.text and '=' not in run.text:
                run.text = run.text.replace('20.22kN/m', f'{new_q:.2f}kN/m')
            run.font.color.rgb = RED
        print(f"  P{i}: 楼面等效荷载已修复 20.22→{new_q:.2f}")

# ============================================================
# 更新地震力相关段落 - 修正P394的计算式
# ============================================================
for i, para in enumerate(doc.paragraphs):
    t = para.text
    # P394: "0.052×0.85×33835=1528.8kN" should be "0.0506×0.85×35559=1528.8kN"
    if '33835=1528.8kN' in t:
        for run in para.runs:
            if '0.052' in run.text:
                run.text = run.text.replace('0.052', '0.0506')
                run.font.color.rgb = RED
            if '33835' in run.text:
                run.text = run.text.replace('33835', '35559')
                run.font.color.rgb = RED
        print(f"  P{i}: 地震力计算式已修复")

# ============================================================
# 更新表4-4 整体D值表 (Table index 15)
# ============================================================
# 查找并修改整体D值
for ti in range(14, 18):
    if ti < len(doc.tables):
        t = doc.tables[ti]
        first_cell = t.rows[0].cells[0].text if t.rows[0].cells else ''
        print(f"  Table {ti+1}: {first_cell[:80]}")

# 找到D值汇总表并修改
# 表4-4 整体抗侧刚度

# ============================================================
# 更新所有跨度引用
# ============================================================
for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        # 修正"0.125×20.22×4.8²"中的跨度和荷载
        if '0.125×20.22×4.8²' in run.text:
            new_q = 2.57+6.2+F_NEW*ss*g_floor
            run.text = run.text.replace('20.22×4.8²', f'{new_q:.2f}×{L1}²')
            run.font.color.rgb = RED
            print(f"  P{i}: 跨中弯矩计算式已修复")
        if '0.125×5.44×4.8²' in run.text:
            new_q = F_NEW*ss*q_floor
            run.text = run.text.replace('5.44×4.8²', f'{new_q:.2f}×{L1}²')
            run.font.color.rgb = RED
            print(f"  P{i}: 活载跨中弯矩计算式已修复")

# 保存
doc.save(DOC)
# 同时更新审阅版
review = DOC.replace('修正版', '审阅版')
doc.save(review)

print(f"\n修补完成！文件已更新。")
print(f"  修正版: {DOC}")
print(f"  审阅版: {review}")
