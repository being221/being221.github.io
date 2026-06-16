#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
第4章(地震) + 第5章(风荷载) 5400跨度完整重算
L1=5.4m, L2=2.4m, 7度0.15g, II类场地, Tg=0.35s, 三级抗震
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

# ============================================================
# 基本参数
# ============================================================
L1, L2 = 5.4, 2.4
L_long = 6.9
ss = 3.45
E = 30e6  # kN/m²

# 截面
b_col, h_col = 0.5, 0.5
b_be, h_be = 0.25, 0.5   # 边跨梁
b_bm, h_bm = 0.25, 0.4   # 中跨梁

# 层高
h_1st, h_std = 4.0, 3.0
H_total = h_1st + 5 * h_std  # 19.0m
heights = [19.0, 16.0, 13.0, 10.0, 7.0, 4.0]  # 6F→1F

# 截面惯性矩
I0_e = b_be * h_be**3 / 12  # 0.002604
I0_m = b_bm * h_bm**3 / 12  # 0.001333
Ic = b_col * h_col**3 / 12  # 0.005208

# 梁线刚度 (考虑楼板放大)
ib = {
    'edge_e': E * 1.5 * I0_e / L1,   # 边榀边跨
    'edge_m': E * 1.5 * I0_m / L2,   # 边榀中跨
    'mid_e':  E * 2.0 * I0_e / L1,   # 中间榀边跨
    'mid_m':  E * 2.0 * I0_m / L2,   # 中间榀中跨
}

# 柱线刚度
ic_top = E * Ic / h_std  # 52083
ic_1st = E * Ic / h_1st  # 39063

print("="*60)
print("梁线刚度 (kN·m)")
for k, v in ib.items():
    print(f"  {k}: {v:.0f}")
print(f"  柱标准层: {ic_top:.0f}, 首层: {ic_1st:.0f}")

# ============================================================
# D值计算
# ============================================================
def calc_D(ic, h, ib_left, ib_right, is_first):
    K = (ib_left + ib_right) / ic
    if is_first:
        alpha = (0.5 + K) / (2 + K)
    else:
        alpha = K / (2 + K)
    D = alpha * 12 * ic / h**2
    return K, alpha, D

# 各框架类型
frames = {}
for ftype, ib_e, ib_m in [('edge', ib['edge_e'], ib['edge_m']),
                            ('mid',  ib['mid_e'],  ib['mid_m'])]:
    frames[ftype] = {}
    for loc, h, ic, isf in [('std', h_std, ic_top, False),
                             ('1st', h_1st, ic_1st, True)]:
        K_e, a_e, D_e = calc_D(ic, h, ib_e, 0, isf)
        K_m, a_m, D_m = calc_D(ic, h, ib_e, ib_m, isf)
        sum_D = 2*D_e + 2*D_m  # 每榀4柱(2边+2中)
        frames[ftype][loc] = {
            'K_e': K_e, 'a_e': a_e, 'D_e': D_e,
            'K_m': K_m, 'a_m': a_m, 'D_m': D_m,
            'sum_D': sum_D
        }

print("\n" + "="*60)
print("D值计算结果")
for ftype in ['mid', 'edge']:
    for loc in ['std', '1st']:
        d = frames[ftype][loc]
        print(f"  {ftype}榀 {loc}: K_e={d['K_e']:.3f} D_e={d['D_e']:.0f} "
              f"K_m={d['K_m']:.3f} D_m={d['D_m']:.0f} Σ={d['sum_D']:.0f}")

# 总D值(7榀=2边+5中)
D_total_std = 2*frames['edge']['std']['sum_D'] + 5*frames['mid']['std']['sum_D']
D_total_1st = 2*frames['edge']['1st']['sum_D'] + 5*frames['mid']['1st']['sum_D']
print(f"\n总D值: 标准层={D_total_std:.0f}, 首层={D_total_1st:.0f}")

# 表4-6用的D值(中间榀，每柱含所有7榀的贡献)
# 实际: D_e_total = 2*D_e_edge + 5*D_e_mid (所有边柱的D值总和)
D_e_std_total = 2*frames['edge']['std']['D_e'] + 5*frames['mid']['std']['D_e']
D_m_std_total = 2*frames['edge']['std']['D_m'] + 5*frames['mid']['std']['D_m']
D_e_1st_total = 2*frames['edge']['1st']['D_e'] + 5*frames['mid']['1st']['D_e']
D_m_1st_total = 2*frames['edge']['1st']['D_m'] + 5*frames['mid']['1st']['D_m']
sum_D_frame_std = 2*frames['mid']['std']['D_e'] + 2*frames['mid']['std']['D_m']  # 单榀
sum_D_frame_1st = 2*frames['mid']['1st']['D_e'] + 2*frames['mid']['1st']['D_m']

# ============================================================
# 重力荷载代表值估算
# ============================================================
# 面积比: (5.4×2+2.4)/(4.8×2+2.4) = 13.2/12 = 1.10
area_ratio = (L1*2+L2) / (4.8*2+L2)
print(f"\n面积比: {area_ratio:.3f}")

# 旧值
Gi_old = {6: 5028, 5: 5679, 4: 5679, 3: 5679, 2: 5679, 1: 5679}
Gi_new = {}
for fl in range(1, 7):
    Gi_new[fl] = round(Gi_old[fl] * area_ratio)
    print(f"  {fl}F: Gi {Gi_old[fl]}→{Gi_new[fl]} kN")

# ============================================================
# 结构自振周期
# ============================================================
T1_old = 0.56
stiff_ratio = D_total_std / 579164  # 旧总D值=579164
# T ∝ 1/√K, 质量也增加了
mass_ratio = sum(Gi_new.values()) / sum(Gi_old.values())
T1_new = T1_old * math.sqrt(mass_ratio / stiff_ratio)
print(f"\nT₁ (估算): {T1_old:.2f}→{T1_new:.2f}s (刚度比={stiff_ratio:.3f}, 质量比={mass_ratio:.3f})")

# ============================================================
# 地震影响系数
# ============================================================
Tg, alpha_max = 0.35, 0.12
gamma = 0.9  # 阻尼比0.05
eta2 = 1.0
# α₁ = (Tg/T₁)^γ × η₂ × α_max, 且 ≥ 0.2^γ × η₂ × α_max (用于T₁>5Tg)
alpha1_new = (Tg / T1_new)**gamma * eta2 * alpha_max
alpha1_old = (Tg / T1_old)**gamma * eta2 * alpha_max
print(f"α₁: {alpha1_old:.4f}→{alpha1_new:.4f}")

# ============================================================
# 基底剪力FEK
# ============================================================
Geq_old = 0.85 * sum(Gi_old.values())
Geq_new = 0.85 * sum(Gi_new.values())
FEK_old = alpha1_old * Geq_old
FEK_new = alpha1_new * Geq_new
print(f"Geq: {Geq_old:.0f}→{Geq_new:.0f} kN")
print(f"FEK: {FEK_old:.1f}→{FEK_new:.1f} kN")

# ============================================================
# 楼层地震力分配
# ============================================================
# δn = 0.08T₁ + 0.07 (T₁ ≤ 0.7s? 内插)
# T₁=0.56: δn=0.08×0.56+0.07=0.0448+0.07=0.115... 实际用的0.14
# 规范: T₁>1.4Tg=0.49时需考虑顶部附加, δn=0.08T₁+0.07 (T₁≤0.7s部分)
# δn计算按规范格式
delta_n_old = 0.14  # 原值
delta_n_new = 0.08 * T1_new + 0.07
delta_n_new = min(delta_n_new, 0.15)  # 上限
print(f"δn: {delta_n_old:.2f}→{delta_n_new:.3f}")

# ΣGiHi
fl_heights = {6: 19.0, 5: 16.0, 4: 13.0, 3: 10.0, 2: 7.0, 1: 4.0}
sum_GiHi_new = sum(Gi_new[fl] * fl_heights[fl] for fl in range(1, 7))
sum_GiHi_old = sum(Gi_old[fl] * fl_heights[fl] for fl in range(1, 7))
print(f"ΣGiHi: {sum_GiHi_old:.0f}→{sum_GiHi_new:.0f} kN·m")

# 各层Fi和Vi
Fi_new = {}
Vi_new = {}
Vi_prev = 0
for fl in [6, 5, 4, 3, 2, 1]:
    Fi_base = FEK_new * (1 - delta_n_new) * Gi_new[fl] * fl_heights[fl] / sum_GiHi_new
    if fl == 6:
        Fi_new[fl] = Fi_base + delta_n_new * FEK_new
    else:
        Fi_new[fl] = Fi_base
    Vi_new[fl] = Vi_prev + Fi_new[fl]
    Vi_prev = Vi_new[fl]

print("\n地震剪力:")
for fl in [6, 5, 4, 3, 2, 1]:
    print(f"  {fl}F: Fi={Fi_new[fl]:.1f}, Vi={Vi_new[fl]:.1f}")

# ============================================================
# 表4-4 层间位移角
# ============================================================
# D值按层取: 6F-2F用std, 1F用1st
def get_D_for_floor(fl):
    return D_total_std if fl >= 2 else D_total_1st

print("\n层间位移角:")
drifts = {}
for fl in [6, 5, 4, 3, 2, 1]:
    h = h_std if fl >= 2 else h_1st
    Df = get_D_for_floor(fl)
    delta = Vi_new[fl] / Df  # m
    drift_angle = delta / h
    drifts[fl] = {'delta_m': delta, 'drift': drift_angle, 'D': Df}
    print(f"  {fl}F: Vi={Vi_new[fl]:.1f}, ΣD={Df:.0f}, Δu={delta*1000:.2f}mm, "
          f"θ=1/{1/drift_angle:.0f}")

# ============================================================
# 反弯点高度 (表4-5)
# ============================================================
# 使用 mid 榀的K值计算y0
def calc_y0(K, fl, is_first):
    """标准反弯点高度比 (查表插值)"""
    m = 6  # 总层数
    n = fl  # 所在层(从下往上数)
    # 均布水平荷载下的标准反弯点比
    y0_table = {
        # K: [n=1, n=2, n=3, n=4, n=5, n=6]
        0.4: [0.80, 0.45, 0.35, 0.30, 0.25, 0.20],
        0.5: [0.78, 0.45, 0.35, 0.30, 0.25, 0.20],
        0.6: [0.75, 0.45, 0.35, 0.30, 0.25, 0.20],
        0.7: [0.73, 0.45, 0.35, 0.30, 0.25, 0.20],
        0.8: [0.70, 0.45, 0.35, 0.30, 0.30, 0.25],
        0.9: [0.68, 0.45, 0.35, 0.30, 0.30, 0.25],
        1.0: [0.65, 0.45, 0.40, 0.35, 0.30, 0.30],
        1.1: [0.63, 0.45, 0.40, 0.35, 0.35, 0.30],
        1.2: [0.60, 0.45, 0.40, 0.35, 0.35, 0.30],
        1.3: [0.58, 0.45, 0.40, 0.35, 0.35, 0.35],
    }
    # 线性插值
    K_vals = sorted(y0_table.keys())
    if K <= K_vals[0]:
        y0s = y0_table[K_vals[0]]
    elif K >= K_vals[-1]:
        y0s = y0_table[K_vals[-1]]
    else:
        for i in range(len(K_vals)-1):
            if K_vals[i] <= K <= K_vals[i+1]:
                y0s_low = y0_table[K_vals[i]]
                y0s_high = y0_table[K_vals[i+1]]
                t = (K - K_vals[i]) / (K_vals[i+1] - K_vals[i])
                y0s = [y0s_low[j] + t*(y0s_high[j]-y0s_low[j]) for j in range(6)]
                break
    return y0s[n-1]  # n is 1-indexed

mid_frame = frames['mid']
print("\n反弯点高度比y0 (mid榀):")
y0_vals = {}
for fl in [6, 5, 4, 3, 2, 1]:
    is_1st = (fl == 1)
    d = mid_frame['1st'] if is_1st else mid_frame['std']
    n_from_bottom = fl  # 6F=n=6, 1F=n=1
    y0_e = calc_y0(d['K_e'], n_from_bottom, is_1st)
    y0_m = calc_y0(d['K_m'], n_from_bottom, is_1st)
    y0_vals[fl] = {'edge': y0_e, 'mid': y0_m}
    print(f"  {fl}F: K_e={d['K_e']:.2f} y0_e={y0_e:.3f}, K_m={d['K_m']:.2f} y0_m={y0_m:.3f}")

# ============================================================
# 表4-6 柱剪力 (用总D值)
# ============================================================
print("\n柱剪力(地震):")
col_shear = {}
for fl in [6, 5, 4, 3, 2, 1]:
    is_1st = (fl == 1)
    De_total = D_e_1st_total if is_1st else D_e_std_total
    Dm_total = D_m_1st_total if is_1st else D_m_std_total
    Df = get_D_for_floor(fl)
    V_e = Vi_new[fl] * De_total / Df
    V_m = Vi_new[fl] * Dm_total / Df
    col_shear[fl] = {'Ve': V_e, 'Vm': V_m}
    print(f"  {fl}F: Vi={Vi_new[fl]:.1f}, De={De_total:.0f}, Dm={Dm_total:.0f}, "
          f"ΣD={Df:.0f}, V_e={V_e:.2f}, V_m={V_m:.2f}")

# ============================================================
# 表4-7 柱端弯矩
# ============================================================
print("\n柱端弯矩(地震):")
col_moments = {}
for fl in [6, 5, 4, 3, 2, 1]:
    h = h_1st if fl == 1 else h_std
    y0e = y0_vals[fl]['edge']
    y0m = y0_vals[fl]['mid']
    Ve = col_shear[fl]['Ve']
    Vm = col_shear[fl]['Vm']
    # M上 = V × (1-y0) × h, M下 = V × y0 × h
    M_e_top = Ve * (1 - y0e) * h
    M_e_bot = Ve * y0e * h
    M_m_top = Vm * (1 - y0m) * h
    M_m_bot = Vm * y0m * h
    col_moments[fl] = {'Me_top': M_e_top, 'Me_bot': M_e_bot,
                        'Mm_top': M_m_top, 'Mm_bot': M_m_bot}
    print(f"  {fl}F: M_e上={M_e_top:.1f} M_e下={M_e_bot:.1f}, "
          f"M_m上={M_m_top:.1f} M_m下={M_m_bot:.1f}")

# ============================================================
# 表4-8 梁端弯矩 (节点平衡)
# ============================================================
print("\n梁端弯矩(地震):")
beam_moments = {}
# 从顶层往下: 上柱下端 + 下柱上端 → 梁端
for fl in [6, 5, 4, 3, 2, 1]:
    # 边节点: ΣM_c = M_e_bot(上层) + M_e_top(本层)
    # 顶层: 只有本层下柱(无上层)
    if fl == 6:
        M_c_e = col_moments[fl]['Me_bot']  # 只有下柱(上端无柱)
        M_c_m = col_moments[fl]['Mm_bot']
    else:
        M_c_e = col_moments[fl+1]['Me_bot'] + col_moments[fl]['Me_top']
        M_c_m = col_moments[fl+1]['Mm_bot'] + col_moments[fl]['Mm_top']

    # 边节点: 边梁端 = ΣM_c, 中节点需分配
    M_be = M_c_e
    # 中节点: 按梁线刚度分配
    ib_e_mid = ib['mid_e']
    ib_m_mid = ib['mid_m']
    M_bm_left = M_c_m * ib_e_mid / (ib_e_mid + ib_m_mid)  # 中跨左端
    M_be_right = M_c_m * ib_m_mid / (ib_e_mid + ib_m_mid)  # 边跨右端

    beam_moments[fl] = {
        'M_be_left': M_be,      # 边跨左端
        'M_be_right': M_bm_left, # 边跨右端(≈中节点左梁分配)
        'M_bm_left': M_bm_left,  # 中跨左端
        'M_bm_right': M_be_right, # 中跨右端(对称)
    }
    print(f"  {fl}F: M_be左={M_be:.1f}, M_be右={M_bm_left:.1f}, "
          f"M_bm={M_be_right:.1f}")

# Wait, I need to reconsider. In Chinese frame analysis:
# - At edge joint (A): M_beam = M_col_upper + M_col_lower
# - At middle joint (B): total M = M_col_upper + M_col_lower
#   distributed to left beam and right beam by stiffness ratio
# For symmetric structure, left beam gets ib_edge/(ib_edge+ib_mid)

# But the beam moments in the table show the beam end moments transferred
# from column moments. Let me redo this more carefully.

# Actually, looking at the table structure:
# T21 (表4-8):
# C[2] = 边柱节点弯矩 (from col)
# C[3] = 中柱节点弯矩
# C[4] = 边柱分配系数 (ib_e/(ib_e+ib_m))... no
# C[4] = 边跨梁左端 (from 边柱节点)
# C[5] = 中跨梁左端 (from 中柱节点, 由边跨传递)
# C[6] = 边跨梁右端

# Let me re-examine T21:
# R2: 楼层 | 边柱 | 中柱 | 边柱 | 中柱 | 边梁 | 中梁 | 边梁
# R3: 6F | -30.43 | -44.96 | -13.04 | -25.66 | 30.43 | 22.75 | 22.21

# Columns: [0]=楼层, [1]=边柱M, [2]=中柱M, [3]=?, [4]=?, [5]=边梁左, [6]=中梁, [7]=边梁右

# Let me re-read the headers:
# R1: 楼层 | | | | | | |
# R2: 楼层 | 边柱 | 中柱 | 边柱 | 中柱 | 边梁 | 中梁 | 边梁

# Hmm, the column headers are duplicated. Let me think about what makes sense:
# For seismic frame analysis with inflection point method:
# 1. From column shear and y0 → column end moments (M上, M下)
# 2. Joint equilibrium → beam end moments
#    - Edge joint: M_beam = M_col_upper + M_col_lower
#    - Middle joint: M_left_beam + M_right_beam = M_col_upper + M_col_lower
#      M_left_beam / M_right_beam = ib_left / ib_right

# In T21, "边柱" column likely shows the column end moment at the joint
# "边梁" column shows the beam end moment

# Wait, looking at T20 first:
# R2: 楼层 | | 边柱 | 中柱 | 边柱 | 中柱 | 边柱 | 中柱
# R3: 6F | 3.00 | 14.49 | 23.54 | 0.90 | 1.09 | -30.43 | -44.96
# Columns: [0]=楼层, [1]=h, [2]=V_e, [3]=V_m, [4]=M_e下(h×y0), [5]=M_m下, [6]=M_e上, [7]=M_m上

# So in T20: [4]=M_e下, [5]=M_m下, [6]=M_e上, [7]=M_m上
# R3: M_e下=0.90, M_m下=1.09, M_e上=-30.43, M_m上=-44.96
# Check: Ve=14.49, h=3.0, y0=0.30
# M_e下 = 14.49 × 0.30 × 3.0 = 13.04 ← not 0.90
#
# Hmm, that doesn't work. Let me reconsider.
# Actually, y0 for 6F might not be 0.30. Looking at T18:
# R3: 6F, 3.00, 0.63, 0.3000, 0.3000, 1.27, 0.3640, 0.3640
# So for 6F: y0_e=0.3000
# M_e_bot = V_e × y0 × h = 14.49 × 0.30 × 3.0 = 13.04
# But T20 says column [4] = 0.90...

# This doesn't add up. Let me look at [4]=0.90 more carefully.
# 0.90 is the "y" value shown in some versions of the table, not the moment.

# Let me re-examine T20 column structure:
# R1: 楼层 |  |  |  |  |  |  |
# R2: 楼层 |  | 边柱 | 中柱 | 边柱 | 中柱 | 边柱 | 中柱
# R3: 6F | 3.00 | 14.49 | 23.54 | 0.90 | 1.09 | -30.43 | -44.96

# Actually, T18 gives y values and T19 gives V values. T20 should have:
# [0]=楼层, [1]=h, [2]=V_e, [3]=V_m, [4]=y_e×h, [5]=y_m×h, [6]=M_e上, [7]=M_m上
# y_e×h = 0.30×3.0 = 0.90 ✓
# y_m×h = 0.364×3.0 = 1.09 ✓
# M_e上 = V_e×(1-y0)×h = 14.49×0.7×3.0 = 30.43 ✓ (negative sign convention)

# So in T20, the moment columns show only M_top (上端). The M_bot (下端) is implied
# and M_top + M_bot = V × h:
# 30.43 + V_e×y0×h = 30.43 + 13.04 = 43.47 = 14.49×3.0 ✓

# Now for T21 (梁端弯矩):
# At each joint, the beam moment = sum of column moments
# 6F edge joint: M_beam = M_e_top(6F) = 30.43 (no upper column, 顶层)
# 6F middle joint: M_total = M_m_top(6F) = 44.96
#   M_be_right = 44.96 × ib_e/(ib_e+ib_m) = 44.96 × 28935/(28935+33333) = 44.96 × 0.465 = 20.90
#   M_bm_left = 44.96 × ib_m/(ib_e+ib_m) = 44.96 × 33333/(28935+33333) = 44.96 × 0.535 = 24.06

# T21 R3: 6F | -30.43 | -44.96 | -13.04 | -25.66 | 30.43 | 22.75 | 22.21
# Hmm, that doesn't match 20.90/24.06.

# Wait, the values in the table use OLD beam stiffness (32552 and 33333):
# Old ratio: 32552/(32552+33333) = 0.494, 33333/(32552+33333) = 0.506
# M_be_right = 44.96 × 0.494 = 22.21 (close to 22.21 but this is the 边梁 right value)
# Wait: 44.96 × 0.506 = 22.75 → M_bm = 22.75 (中梁)
# 44.96 × 0.494 = 22.21 → M_be_right = 22.21 (边梁右)

# Hmm, but the table also has [3] = -13.04 and [4] = -25.66. What are these?
# [3] = -13.04 = M_e_bot(6F) = Ve × y0 × h = 14.49 × 0.3 × 3.0 = 13.04
# [4] = -25.66 = M_m_bot(6F) = Vm × y0 × h = 23.54 × 0.364 × 3.0 = 25.70 ≈ 25.66

# So T21 columns are:
# [0]=楼层, [1]=M_e_top(边柱上端), [2]=M_m_top(中柱上端),
# [3]=M_e_bot(边柱下端), [4]=M_m_bot(中柱下端),
# [5]=M_边梁左, [6]=M_中梁, [7]=M_边梁右

# Actually, looking more carefully:
# R3: 6F | -30.43 | -44.96 | -13.04 | -25.66 | 30.43 | 22.75 | 22.21
# [1] = -30.43 = M_e_top (上端弯矩)
# [2] = -44.96 = M_m_top
# [3] = -13.04 = ???
# [4] = -25.66 = ???

# Hmm, let me think again. The table header shows:
# R2: 楼层 | 边柱 | 中柱 | 边柱 | 中柱 | 边梁 | 中梁 | 边梁

# The structure might be showing moments at the column-beam joint:
# [1] = 边柱节点: M from column = node moment
# [2] = 中柱节点: M from column
# [3] = 边柱分配: M distributed to beam at edge node = 0 (only one beam)
#   Wait no, [3] = -13.04 = M_e_bot...

# Actually I think [3] and [4] are the lower column end moments for the floor below:
# For 6F: there's no upper column, so [3]=M_e_bot(6F)=13.04, [4]=M_m_bot(6F)=25.66
# These are carried down to 5F for the joint equilibrium calculation at 5F.
# [5] = M_beam_edge_left = 30.43 (from edge joint, no upper col, equals M_e_top)
# [6] = M_beam_mid = 22.75 (中跨梁端弯矩)
# [7] = M_beam_edge_right = 22.21 (边跨右端弯矩)

# Then for 5F:
# Joint equilibrium at edge: M_beam = M_e_bot(6F) + M_e_top(5F) = 13.04 + 42.43 = 55.47
# Joint equilibrium at middle: M_total = M_m_bot(6F) + M_m_top(5F) = 25.66 + 63.20 = 88.86
# M_be_right(5F) = 88.86 × 0.494 = 43.90, M_bm_left(5F) = 88.86 × 0.506 = 44.96

# R4: 5F | -42.43 | -63.20 | -28.28 | -51.71 | 55.47 | 44.96 | 43.90
# [1]=-42.43=M_e_top(5F), [2]=-63.20=M_m_top(5F)
# [3]=-28.28=M_e_bot(5F)=Ve×y0×h=23.57×0.4×3.0=28.28 ✓
# [4]=-51.71=M_m_bot(5F)=38.3×0.45×3.0=51.71 ✓
# [5]=55.47=M_beam_edge_left from joint: 13.04+42.43=55.47 ✓
# [6]=44.96=M_beam_mid
# [7]=43.90=M_beam_edge_right

# Great, now I understand the table structure. Let me proceed.

# For M_beam at middle joint: M_left + M_right = ΣM_col
# M_left (边跨梁右端) = ΣM_col × ib_edge / (ib_edge + ib_mid)
# M_right (中跨梁左端) = ΣM_col × ib_mid / (ib_edge + ib_mid)

# OK, now let me recalculate the beam moments properly.

print("\n梁端弯矩(地震) - 节点平衡法:")
beam_moments = {}
for fl in [6, 5, 4, 3, 2, 1]:
    is_1st = (fl == 1)
    h = h_1st if is_1st else h_std

    Ve = col_shear[fl]['Ve']
    Vm = col_shear[fl]['Vm']
    y0e = y0_vals[fl]['edge']
    y0m = y0_vals[fl]['mid']

    M_e_top = Ve * (1 - y0e) * h  # 边柱上端
    M_e_bot = Ve * y0e * h         # 边柱下端
    M_m_top = Vm * (1 - y0m) * h  # 中柱上端
    M_m_bot = Vm * y0m * h         # 中柱下端

    # 节点弯矩(梁端弯矩之和)
    if fl == 6:
        M_node_e = M_e_top  # 顶层: 无上柱
        M_node_m = M_m_top
    else:
        M_node_e = M_e_top + col_moments[fl+1]['Me_bot']  # 本层上端+上层下端
        M_node_m = M_m_top + col_moments[fl+1]['Mm_bot']

    # 中节点分配 (mid榀 ib值)
    ib_e_m = ib['mid_e']
    ib_m_m = ib['mid_m']
    M_be_right = M_node_m * ib_e_m / (ib_e_m + ib_m_m)  # 边跨右端
    M_bm_left  = M_node_m * ib_m_m / (ib_e_m + ib_m_m)  # 中跨左端

    beam_moments[fl] = {
        'Me_top': M_e_top, 'Me_bot': M_e_bot,
        'Mm_top': M_m_top, 'Mm_bot': M_m_bot,
        'M_node_e': M_node_e, 'M_node_m': M_node_m,
        'M_be_left': M_node_e,     # 边跨左端=边节点总弯矩
        'M_be_right': M_be_right,  # 边跨右端
        'M_bm_left': M_bm_left,    # 中跨左端
        'M_bm_right': M_bm_left,   # 中跨右端(对称)
    }
    print(f"  {fl}F: M_e=({M_e_top:.1f},{M_e_bot:.1f}) M_m=({M_m_top:.1f},{M_m_bot:.1f})")
    print(f"       节点: M_e_node={M_node_e:.1f} M_m_node={M_node_m:.1f}")
    print(f"       梁端: 边左={M_node_e:.1f} 边右={M_be_right:.1f} 中跨={M_bm_left:.1f}")

# ============================================================
# 表4-9 梁端剪力与柱轴力
# ============================================================
print("\n梁端剪力与柱轴力(地震):")
beam_shear_axial = {}
axial_e_cum, axial_m_cum = 0, 0
for fl in [6, 5, 4, 3, 2, 1]:
    bm = beam_moments[fl]
    # Vb = (M_left + M_right) / L
    Vb_edge = (bm['M_be_left'] + bm['M_be_right']) / L1
    Vb_mid  = (bm['M_bm_left'] + bm['M_bm_right']) / L2
    # 边柱轴力增量 = Vb_edge(左) - Vb_edge(右) 或直接用剪力
    # 边柱: N += Vb_edge (来自边跨)
    # 中柱: N += Vb_mid - Vb_edge (中柱两侧梁剪力差)
    axial_e_cum += Vb_edge
    axial_m_cum += Vb_mid - Vb_edge
    beam_shear_axial[fl] = {
        'Vb_edge': Vb_edge, 'Vb_mid': Vb_mid,
        'N_edge': axial_e_cum, 'N_mid': axial_m_cum
    }
    # Note: 柱受压为正, 但表里可能用负号表示压力
    print(f"  {fl}F: Vb_e={Vb_edge:.2f}, Vb_m={Vb_mid:.2f}, "
          f"N_e={axial_e_cum:.2f}, N_m={axial_m_cum:.2f}")

# ============================================================
# 第5章 风荷载
# ============================================================
print("\n" + "="*60)
print("第5章: 风荷载重算")
print("="*60)

w0 = 0.30  # 基本风压
mu_s = 1.3  # 体型系数
# 风压高度变化系数 μz (B类地面粗糙度)
mu_z_table = {5: 0.95, 10: 1.00, 15: 1.14, 20: 1.25, 30: 1.42}

def mu_z(height):
    """B类 风压高度变化系数 内插"""
    if height <= 5: return 0.95
    if height >= 30: return 1.42
    ks = sorted(mu_z_table.keys())
    for i in range(len(ks)-1):
        if ks[i] <= height <= ks[i+1]:
            t = (height - ks[i]) / (ks[i+1] - ks[i])
            return mu_z_table[ks[i]] + t * (mu_z_table[ks[i+1]] - mu_z_table[ks[i]])
    return 1.0

# 楼层风荷载面积: 上下各一半层高 × 纵向跨度6.9m
wind_heights = {6: 18.6, 5: 15.6, 4: 12.6, 3: 9.6, 2: 6.6, 1: 3.6}
wind_areas = {6: 2.4*6.9, 5: 3.0*6.9, 4: 3.0*6.9, 3: 3.0*6.9, 2: 3.0*6.9, 1: 3.9*6.9}

# 风振系数 βz = 1 + ξ×ν×φz/μz (GB50009-2012)
# w0×T1² = 0.30 × 0.61² = 0.111
# ξ = 1.47 (钢筋混凝土结构, w0T1²≈0.11, 查表)
xi = 1.47
# ν = 0.45 (H/B = 19/13.2 = 1.44, B类, 查表)
nu = 0.45
# φz: 振型系数, 近似取 z/H
H_bld = 19.0

def beta_z(z):
    phi_z = z / H_bld
    mz = mu_z(z)
    return 1.0 + xi * nu * phi_z / mz

print("\n风荷载计算 (含风振系数):")
wind_F = {}
for fl in [6, 5, 4, 3, 2, 1]:
    z = wind_heights[fl]
    mz = mu_z(z)
    bz = beta_z(z)
    wk = bz * mu_s * mz * w0
    A = wind_areas[fl]
    F_wind = wk * A
    wind_F[fl] = {'z': z, 'muz': mz, 'bz': bz, 'wk': wk, 'A': A, 'F': F_wind}
    print(f"  {fl}F: z={z}m, μz={mz:.2f}, βz={bz:.3f}, ωk={wk:.3f}, "
          f"A={A:.1f}m², F={F_wind:.2f}kN")

# Now compute Vi (story shear) for wind
wind_V = {}
V_prev = 0
for fl in [6, 5, 4, 3, 2, 1]:
    V_prev += wind_F[fl]['F']
    wind_V[fl] = V_prev
    print(f"  {fl}F Vi={V_prev:.2f}")

# Wind uses per-frame D values (single frame model)
# Actually from T24: ΣD=87500 = single mid frame
# Let me use single mid frame D values for wind
D_wind_std = frames['mid']['std']['sum_D']  # single frame
D_wind_1st = frames['mid']['1st']['sum_D']

# D per column (for single frame)
De_wind_std = frames['mid']['std']['D_e']
Dm_wind_std = frames['mid']['std']['D_m']
De_wind_1st = frames['mid']['1st']['D_e']
Dm_wind_1st = frames['mid']['1st']['D_m']

print(f"\n风荷载D值(单榀): 标准层Σ={D_wind_std:.0f}, 首层Σ={D_wind_1st:.0f}")
print(f"  De_std={De_wind_std:.0f}, Dm_std={Dm_wind_std:.0f}")
print(f"  De_1st={De_wind_1st:.0f}, Dm_1st={Dm_wind_1st:.0f}")

# Column shear for wind
wind_col_shear = {}
for fl in [6, 5, 4, 3, 2, 1]:
    is_1st = (fl == 1)
    Df = D_wind_1st if is_1st else D_wind_std
    De = De_wind_1st if is_1st else De_wind_std
    Dm = Dm_wind_1st if is_1st else Dm_wind_std
    Ve = wind_V[fl] * De / Df
    Vm = wind_V[fl] * Dm / Df
    wind_col_shear[fl] = {'Ve': Ve, 'Vm': Vm}
    print(f"  {fl}F: Vi={wind_V[fl]:.2f}, Ve={Ve:.3f}, Vm={Vm:.3f}")

# Column end moments for wind
print("\n柱端弯矩(风):")
wind_col_moments = {}
for fl in [6, 5, 4, 3, 2, 1]:
    h = h_1st if fl == 1 else h_std
    y0e = y0_vals[fl]['edge']
    y0m = y0_vals[fl]['mid']
    Ve = wind_col_shear[fl]['Ve']
    Vm = wind_col_shear[fl]['Vm']
    Me_top = Ve * (1 - y0e) * h
    Me_bot = Ve * y0e * h
    Mm_top = Vm * (1 - y0m) * h
    Mm_bot = Vm * y0m * h
    wind_col_moments[fl] = {'Me_top': Me_top, 'Me_bot': Me_bot,
                            'Mm_top': Mm_top, 'Mm_bot': Mm_bot}
    print(f"  {fl}F: Me=({Me_top:.2f},{Me_bot:.2f}) Mm=({Mm_top:.2f},{Mm_bot:.2f})")

# Beam moments for wind
print("\n梁端弯矩(风):")
wind_beam_moments = {}
for fl in [6, 5, 4, 3, 2, 1]:
    wcm = wind_col_moments[fl]
    if fl == 6:
        M_node_e = wcm['Me_top']
        M_node_m = wcm['Mm_top']
    else:
        M_node_e = wcm['Me_top'] + wind_col_moments[fl+1]['Me_bot']
        M_node_m = wcm['Mm_top'] + wind_col_moments[fl+1]['Mm_bot']

    ib_e_m = ib['mid_e']
    ib_m_m = ib['mid_m']
    M_be_right = M_node_m * ib_e_m / (ib_e_m + ib_m_m)
    M_bm = M_node_m * ib_m_m / (ib_e_m + ib_m_m)

    wind_beam_moments[fl] = {
        'Me_top': wcm['Me_top'], 'Me_bot': wcm['Me_bot'],
        'Mm_top': wcm['Mm_top'], 'Mm_bot': wcm['Mm_bot'],
        'M_node_e': M_node_e, 'M_node_m': M_node_m,
        'M_be_left': M_node_e, 'M_be_right': M_be_right,
        'M_bm': M_bm
    }
    print(f"  {fl}F: 边梁左={M_node_e:.2f} 边梁右={M_be_right:.2f} 中梁={M_bm:.2f}")

# ============================================================
# 写入DOCX
# ============================================================
print("\n" + "="*60)
print("写入DOCX表格")
print("="*60)

# --- T16 表4-3 地震剪力 ---
t = doc.tables[16]
for ri in range(2, 8):
    fl = 8 - ri  # R2=6F, R3=5F, ..., R7=1F
    row = t.rows[ri]
    row.cells[1].text = f'{fl_heights[fl]:.2f}'
    row.cells[2].text = f'{Gi_new[fl]}'
    row.cells[3].text = f'{Gi_new[fl]*fl_heights[fl]:.0f}'
    row.cells[4].text = f'{sum_GiHi_new:.0f}'
    row.cells[5].text = f'{Fi_new[fl]:.2f}'
    row.cells[6].text = f'{Vi_new[fl]:.2f}'
    row.cells[7].text = f'{Fi_new[fl]/FEK_new:.4f}'
print("  T16 表4-3 ✓")

# --- T17 表4-4 层间位移角 ---
t = doc.tables[17]
for ri in range(2, 8):
    fl = 8 - ri
    d = drifts[fl]
    row = t.rows[ri]
    row.cells[1].text = f'{Vi_new[fl]:.2f}'
    row.cells[2].text = f'{d["D"]:.0f}'
    row.cells[3].text = f'{h_std if fl>=2 else h_1st:.2f}'
    row.cells[4].text = f'1/{1/d["drift"]:.0f}'
    # C5 = limit, keep as is (1/550)
print("  T17 表4-4 ✓")

# --- T18 表4-5 反弯点 ---
t = doc.tables[18]
for ri in range(3, 9):
    fl = 9 - ri  # R3=6F, R4=5F, ..., R8=1F
    row = t.rows[ri]
    is_1st = (fl == 1)
    h = h_1st if is_1st else h_std
    d = frames['mid']['1st'] if is_1st else frames['mid']['std']
    row.cells[1].text = f'{h:.2f}'
    row.cells[2].text = f'{d["K_e"]:.2f}'
    row.cells[3].text = f'{y0_vals[fl]["edge"]:.4f}'
    row.cells[4].text = f'{y0_vals[fl]["edge"]:.4f}'  # y (same as y0 for simple case)
    row.cells[5].text = f'{d["K_m"]:.2f}'
    row.cells[6].text = f'{y0_vals[fl]["mid"]:.4f}'
    row.cells[7].text = f'{y0_vals[fl]["mid"]:.4f}'
print("  T18 表4-5 ✓")

# --- T19 表4-6 地震柱剪力 ---
t = doc.tables[19]
for ri in range(2, 8):
    fl = 8 - ri
    row = t.rows[ri]
    is_1st = (fl == 1)
    De = D_e_1st_total if is_1st else D_e_std_total
    Dm = D_m_1st_total if is_1st else D_m_std_total
    Df = D_total_1st if is_1st else D_total_std
    row.cells[1].text = f'{Vi_new[fl]:.1f}'
    row.cells[2].text = f'{De:.0f}'
    row.cells[3].text = f'{Dm:.0f}'
    row.cells[4].text = f'{Df:.0f}'
    row.cells[5].text = f'{col_shear[fl]["Ve"]:.2f}'
    row.cells[6].text = f'{col_shear[fl]["Vm"]:.2f}'
print("  T19 表4-6 ✓")

# --- T20 表4-7 地震柱端弯矩 ---
t = doc.tables[20]
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    is_1st = (fl == 1)
    h = h_1st if is_1st else h_std
    bm = beam_moments[fl]
    y0e = y0_vals[fl]['edge']
    y0m = y0_vals[fl]['mid']
    row.cells[1].text = f'{h:.2f}'
    row.cells[2].text = f'{col_shear[fl]["Ve"]:.2f}'
    row.cells[3].text = f'{col_shear[fl]["Vm"]:.2f}'
    row.cells[4].text = f'{y0e*h:.2f}'  # y*h for edge
    row.cells[5].text = f'{y0m*h:.2f}'  # y*h for mid
    row.cells[6].text = f'{-bm["Me_top"]:.2f}'
    row.cells[7].text = f'{-bm["Mm_top"]:.2f}'
print("  T20 表4-7 ✓")

# --- T21 表4-8 地震梁端弯矩 ---
t = doc.tables[21]
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    bm = beam_moments[fl]
    row.cells[1].text = f'{-bm["Me_top"]:.2f}'
    row.cells[2].text = f'{-bm["Mm_top"]:.2f}'
    row.cells[3].text = f'{-bm["Me_bot"]:.2f}'
    row.cells[4].text = f'{-bm["Mm_bot"]:.2f}'
    row.cells[5].text = f'{bm["M_be_left"]:.2f}'
    row.cells[6].text = f'{bm["M_bm_left"]:.2f}'
    row.cells[7].text = f'{bm["M_be_right"]:.2f}'
print("  T21 表4-8 ✓")

# --- T22 表4-9 地震梁端剪力与柱轴力 ---
t = doc.tables[22]
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    bsa = beam_shear_axial[fl]
    row.cells[1].text = f'{L1:.2f}'
    row.cells[2].text = f'{L2:.2f}'
    # 边梁 M左, M右
    row.cells[3].text = f'{beam_moments[fl]["M_be_left"]:.2f}'
    row.cells[4].text = f'{beam_moments[fl]["M_be_right"]:.2f}'
    # 中梁 M左, M右
    row.cells[5].text = f'{beam_moments[fl]["M_bm_left"]:.2f}'
    row.cells[6].text = f'{beam_moments[fl]["M_bm_right"]:.2f}'
    # 边梁 V, 柱N
    row.cells[7].text = f'{bsa["Vb_edge"]:.2f}'
    if len(row.cells) > 8:
        row.cells[8].text = f'{bsa["Vb_mid"]:.2f}'
    if len(row.cells) > 9:
        row.cells[9].text = f'{-bsa["N_edge"]:.2f}'  # 压力为负
    if len(row.cells) > 10:
        row.cells[10].text = f'{-bsa["N_mid"]:.2f}'
print("  T22 表4-9 ✓")

# ============================================================
# 第5章 风荷载表格
# ============================================================

# --- T23 表5-1 风荷载集中力 ---
t = doc.tables[23]
for ri in range(2, 8):
    fl = 8 - ri
    row = t.rows[ri]
    wf = wind_F[fl]
    row.cells[1].text = f'{wf["z"]:.2f}'
    row.cells[2].text = f'{wf["muz"]:.2f}'
    row.cells[3].text = f'{w0:.2f}'
    row.cells[4].text = f'{mu_s:.2f}'
    row.cells[5].text = f'{wf["wk"]:.2f}'
    # A: 保持原来的面积表述格式
    if fl == 6:
        row.cells[6].text = f'2.4×6.9=16.56m²'
    elif fl == 1:
        row.cells[6].text = f'3.9×6.9=26.91m²'
    else:
        row.cells[6].text = f'3×6.9=20.7m²'
    row.cells[7].text = f'{wf["F"]:.2f}'
print("  T23 表5-1 ✓")

# --- T24 表5-2 风荷载层间位移角 ---
t = doc.tables[24]
for ri in range(2, 8):
    fl = 8 - ri
    row = t.rows[ri]
    is_1st = (fl == 1)
    Df = D_wind_1st if is_1st else D_wind_std
    h = h_1st if is_1st else h_std
    delta = wind_V[fl] / Df
    drift = delta / h
    row.cells[1].text = f'{wind_V[fl]:.2f}'
    row.cells[2].text = f'{Df:.0f}'
    row.cells[3].text = f'{h:.2f}'
    row.cells[4].text = f'1/{1/drift:.0f}'
print("  T24 表5-2 ✓")

# --- T25 表5-3 风荷载柱剪力 ---
t = doc.tables[25]
for ri in range(2, 8):
    fl = 8 - ri
    row = t.rows[ri]
    is_1st = (fl == 1)
    Df = D_wind_1st if is_1st else D_wind_std
    De = De_wind_1st if is_1st else De_wind_std
    Dm = Dm_wind_1st if is_1st else Dm_wind_std
    row.cells[1].text = f'{wind_V[fl]:.2f}'
    row.cells[2].text = f'{De:.0f}'
    row.cells[3].text = f'{Dm:.0f}'
    row.cells[4].text = f'{Df:.0f}'
    row.cells[5].text = f'{wind_col_shear[fl]["Ve"]:.2f}'
    row.cells[6].text = f'{wind_col_shear[fl]["Vm"]:.2f}'
print("  T25 表5-3 ✓")

# --- T26 表5-4 风荷载柱端弯矩 ---
t = doc.tables[26]
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    is_1st = (fl == 1)
    h = h_1st if is_1st else h_std
    wcm = wind_col_moments[fl]
    y0e = y0_vals[fl]['edge']
    y0m = y0_vals[fl]['mid']
    row.cells[1].text = f'{h:.2f}'
    row.cells[2].text = f'{wind_col_shear[fl]["Ve"]:.2f}'
    row.cells[3].text = f'{wind_col_shear[fl]["Vm"]:.2f}'
    row.cells[4].text = f'{y0e*h:.2f}'
    row.cells[5].text = f'{y0m*h:.2f}'
    row.cells[6].text = f'{-wcm["Me_top"]:.2f}'
    row.cells[7].text = f'{-wcm["Mm_top"]:.2f}'
print("  T26 表5-4 ✓")

# --- T27 表5-5 风荷载梁端弯矩 ---
t = doc.tables[27]
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    wbm = wind_beam_moments[fl]
    row.cells[1].text = f'{-wbm["Me_top"]:.2f}'
    row.cells[2].text = f'{-wbm["Mm_top"]:.2f}'
    row.cells[3].text = f'{wbm["M_be_left"]:.2f}'
    row.cells[4].text = f'{wbm["M_be_right"]:.2f}'
    row.cells[5].text = f'{wbm["M_bm"]:.2f}'
    row.cells[6].text = f'{wbm["M_bm"]:.2f}'
print("  T27 表5-5 ✓")

# --- T28 表5-6 风荷载梁端剪力与柱轴力 ---
t = doc.tables[28]
wind_N_e_cum, wind_N_m_cum = 0, 0
for ri in range(3, 9):
    fl = 9 - ri
    row = t.rows[ri]
    wbm = wind_beam_moments[fl]
    Vb_e = (wbm['M_be_left'] + wbm['M_be_right']) / L1
    Vb_m = (wbm['M_bm'] + wbm['M_bm']) / L2
    wind_N_e_cum += Vb_e
    wind_N_m_cum += Vb_m - Vb_e
    row.cells[1].text = f'{L1:.2f}'
    row.cells[2].text = f'{L2:.2f}'
    row.cells[3].text = f'{wbm["M_be_left"]:.2f}'
    row.cells[4].text = f'{wbm["M_be_right"]:.2f}'
    row.cells[5].text = f'{wbm["M_bm"]:.2f}'
    row.cells[6].text = f'{wbm["M_bm"]:.2f}'
    row.cells[7].text = f'{Vb_e:.2f}'
    if len(row.cells) > 8:
        row.cells[8].text = f'{Vb_m:.2f}'
    if len(row.cells) > 9:
        row.cells[9].text = f'{-wind_N_e_cum:.2f}'
    if len(row.cells) > 10:
        row.cells[10].text = f'{-wind_N_m_cum:.2f}'
print("  T28 表5-6 ✓")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n第4章+第5章更新完成!")
print(f"修正版: {DOC}")
print(f"审阅版: {review}")
