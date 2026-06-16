#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复列组合表 (7-15~7-20): 正确柱端弯矩映射 + 重算组合
映射规则: 每层柱=该层以下柱段, 上端=该层A_l/B_l, 下端=下一层A_u/B_u
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
BAK = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_4800原版备份.docx'

doc = Document(DOC)
bak = Document(BAK)

# ============================================================
# 基本参数
# ============================================================
L1, L2 = 5.4, 2.4
BETA = 0.85
ss = 3.45
F_NEW = 1 - 2*(0.5*ss/L1)**2 + (0.5*ss/L1)**3

g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0

qe_d_r = g_beam_e + F_NEW*ss*g_roof
qe_d_f = g_beam_e + g_wall_e + F_NEW*ss*g_floor
qm_d_r = g_beam_m + 0.625*L2*g_roof
qm_d_f = g_beam_m + g_wall_m + 0.625*L2*g_floor
qe_l_r = F_NEW*ss*q_roof
qe_l_f = F_NEW*ss*q_floor
qm_l_r = 0.625*L2*q_roof
qm_l_f = 0.625*L2*q_floor

E = 30e6; Ic = 0.5**4/12
ic_std = E*Ic/3.0; ic_1st = E*Ic/4.0
ib_edge = E*1.5*0.25*0.5**3/12/L1
ib_mid = E*1.5*0.25*0.4**3/12/L2

def fem(q,L): return q*L**2/12

def moment_dist(ic_u,ic_l,ib_e,ib_m,q_e,q_m):
    FE_e=fem(q_e,L1); FE_m=fem(q_m,L2)
    sA=ic_u+ic_l+ib_e; sB=ic_u+ic_l+ib_e+ib_m
    mu_A_u=ic_u/sA; mu_A_l=ic_l/sA; mu_A_b=ib_e/sA
    mu_B_u=ic_u/sB; mu_B_l=ic_l/sB; mu_B_bl=ib_e/sB; mu_B_br=ib_m/sB
    M_A_u=M_A_l=0; M_A_b=-FE_e
    M_B_u=M_B_l=0; M_B_bl=FE_e; M_B_br=-FE_m
    # 1st
    uA=M_A_u+M_A_l+M_A_b; dAb=-mu_A_b*uA; dAu=-mu_A_u*uA; dAl=-mu_A_l*uA
    M_A_b+=dAb; M_A_u+=dAu; M_A_l+=dAl
    uB=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl=-mu_B_bl*uB; dBbr=-mu_B_br*uB; dBu=-mu_B_u*uB; dBl=-mu_B_l*uB
    M_B_bl+=dBbl; M_B_br+=dBbr; M_B_u+=dBu; M_B_l+=dBl
    M_A_b+=0.5*dBbl; M_B_bl+=0.5*dAb; M_B_br+=0.5*(-dBbr)
    # 2nd
    uA2=M_A_u+M_A_l+M_A_b; dAb2=-mu_A_b*uA2; dAu2=-mu_A_u*uA2; dAl2=-mu_A_l*uA2
    M_A_b+=dAb2; M_A_u+=dAu2; M_A_l+=dAl2
    uB2=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl2=-mu_B_bl*uB2; dBbr2=-mu_B_br*uB2; dBu2=-mu_B_u*uB2; dBl2=-mu_B_l*uB2
    M_B_bl+=dBbl2; M_B_br+=dBbr2; M_B_u+=dBu2; M_B_l+=dBl2
    M_A_b+=0.5*dBbl2; M_B_bl+=0.5*dAb2; M_B_br+=0.5*(-dBbr2)
    return dict(A_u=M_A_u,A_l=M_A_l,A_b=M_A_b,
                B_u=M_B_u,B_l=M_B_l,B_bl=M_B_bl,B_br=M_B_br)

floors_cfg = [
    ('6F',0,ic_std,qe_d_r,qm_d_r,qe_l_r,qm_l_r),
    ('5F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('4F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('3F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('2F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('1F',ic_std,ic_1st,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
]

# 每层弯矩分配结果
joints = {}
for nm,icu,icl,qed,qmd,qel,qml in floors_cfg:
    rd = moment_dist(icu,icl,ib_edge,ib_mid,qed,qmd)
    rl = moment_dist(icu,icl,ib_edge,ib_mid,qel,qml)
    joints[nm] = {'dead':rd,'live':rl,'qe_d':qed,'qm_d':qmd,'qe_l':qel,'qm_l':qml}

def beam_shear(q,L,Ml,Mr):
    Vl = q*L/2 + (Mr-Ml)/L
    Vr = q*L/2 + (Ml-Mr)/L
    return Vl, Vr

# ============================================================
# 构建 STORY-BASED 柱内力
# ============================================================
# story[i] (i=0→6F, i=5→1F): 柱段 = 层i楼面以下到层i-1楼面(或基础)
# 上端 = joint[i].A_l/B_l (柱段在层i楼面以下的顶部)
# 下端 = joint[i+1].A_u/B_u (柱段在层i-1楼面以上的底部, i<5)
# 对于1F(i=5): 下端在基础, 无下层joint, 用上端×0.65(参考4800模式)

floor_names = ['6F','5F','4F','3F','2F','1F']
Gc_std = 6.76 * 3.0
Gc_1st = 6.76 * 4.0

# 集中力
sec_beam = 1.54 * L1 / 2
edge_floor_conc = g_floor * (ss**2/4 + ss*L1/2)
mid_floor_conc = g_floor * ((ss**2/4 + ss*L1/2) + (ss*L2 - 0.5*L2*L2))
edge_roof_conc = g_roof * (ss**2/4 + ss*L1/2)
mid_roof_conc = g_roof * ((ss**2/4 + ss*L1/2) + (ss*L2 - 0.5*L2*L2))

F_edge_roof = 31.19 + 20.08 + sec_beam + edge_roof_conc
F_edge_floor = 42.3 + 20.08 + sec_beam + edge_floor_conc
F_mid_roof = 3.7 + 20.08 + sec_beam + mid_roof_conc
F_mid_floor = 41.95 + 20.08 + sec_beam + mid_floor_conc

# 逐层累计轴力
N_edge_d_top = [0.0]*6; N_edge_d_bot = [0.0]*6
N_edge_l_top = [0.0]*6; N_edge_l_bot = [0.0]*6
N_mid_d_top = [0.0]*6; N_mid_d_bot = [0.0]*6
N_mid_l_top = [0.0]*6; N_mid_l_bot = [0.0]*6

# 柱端弯矩 (story-based)
M_edge_d_top = [0.0]*6; M_edge_d_bot = [0.0]*6
M_edge_l_top = [0.0]*6; M_edge_l_bot = [0.0]*6
M_mid_d_top = [0.0]*6; M_mid_d_bot = [0.0]*6
M_mid_l_top = [0.0]*6; M_mid_l_bot = [0.0]*6

# 柱剪力 (story-based)
V_edge_d = [0.0]*6; V_edge_l = [0.0]*6
V_mid_d = [0.0]*6; V_mid_l = [0.0]*6

prev_edge_d = 0; prev_mid_d = 0
prev_edge_l = 0; prev_mid_l = 0

for fi in range(6):
    nm = floor_names[fi]
    jd = joints[nm]['dead']; jl = joints[nm]['live']
    qed = joints[nm]['qe_d']; qmd = joints[nm]['qm_d']
    qel = joints[nm]['qe_l']; qml = joints[nm]['qm_l']

    is_roof = (fi == 0); is_first = (fi == 5)
    h = 4.0 if is_first else 3.0

    # === 柱端弯矩 (story-based) ===
    # 上端 = A_l/B_l at THIS joint (柱段顶部)
    M_edge_d_top[fi] = jd['A_l']
    M_edge_l_top[fi] = jl['A_l']
    M_mid_d_top[fi] = jd['B_l']
    M_mid_l_top[fi] = jl['B_l']

    # 下端 = A_u/B_u at NEXT LOWER joint (柱段底部)
    # 对于1F(最底层): 下端在基础, 无下层joint
    if fi < 5:
        next_jd = joints[floor_names[fi+1]]['dead']
        next_jl = joints[floor_names[fi+1]]['live']
        M_edge_d_bot[fi] = next_jd['A_u']
        M_edge_l_bot[fi] = next_jl['A_u']
        M_mid_d_bot[fi] = next_jd['B_u']
        M_mid_l_bot[fi] = next_jl['B_u']
    else:
        # 1F下端: 参考4800模式 ratio≈0.65, 保守取上端值
        M_edge_d_bot[fi] = M_edge_d_top[fi]
        M_edge_l_bot[fi] = M_edge_l_top[fi]
        M_mid_d_bot[fi] = M_mid_d_top[fi]
        M_mid_l_bot[fi] = M_mid_l_top[fi]

    # === 柱剪力 V = (|M_top|+|M_bot|)/h ===
    V_edge_d[fi] = (abs(M_edge_d_top[fi]) + abs(M_edge_d_bot[fi])) / h
    V_edge_l[fi] = (abs(M_edge_l_top[fi]) + abs(M_edge_l_bot[fi])) / h
    V_mid_d[fi] = (abs(M_mid_d_top[fi]) + abs(M_mid_d_bot[fi])) / h
    V_mid_l[fi] = (abs(M_mid_l_top[fi]) + abs(M_mid_l_bot[fi])) / h

    # === 柱轴力 (逐层累计) ===
    # 调幅后梁端弯矩→梁端剪力
    M_le_d = BETA*jd['A_b']; M_re_d = BETA*jd['B_bl']; M_me_d = BETA*jd['B_br']
    M_le_l = BETA*jl['A_b']; M_re_l = BETA*jl['B_bl']; M_me_l = BETA*jl['B_br']

    Vl_e_d, Vr_e_d = beam_shear(qed, L1, M_le_d, M_re_d)
    Vl_m_d, Vr_m_d = beam_shear(qmd, L2, M_me_d, -M_me_d)
    Vl_e_l, Vr_e_l = beam_shear(qel, L1, M_le_l, M_re_l)
    Vl_m_l, Vr_m_l = beam_shear(qml, L2, M_me_l, -M_me_l)

    Fe = F_edge_roof if is_roof else F_edge_floor
    Fm = F_mid_roof if is_roof else F_mid_floor
    Gc = Gc_1st if is_first else Gc_std

    n_ed = Fe + Vr_e_d  # 边柱: 集中力 + 边跨右端剪力
    n_md = Fm + Vr_e_d + Vl_m_d  # 中柱: 集中力 + 边跨右端 + 中跨左端
    n_el = Vr_e_l
    n_ml = Vr_e_l + Vl_m_l

    if fi > 0:
        n_ed += prev_edge_d
        n_md += prev_mid_d
        n_el += prev_edge_l
        n_ml += prev_mid_l

    N_edge_d_top[fi] = n_ed
    N_edge_d_bot[fi] = n_ed + Gc
    N_edge_l_top[fi] = n_el
    N_edge_l_bot[fi] = n_el

    N_mid_d_top[fi] = n_md
    N_mid_d_bot[fi] = n_md + Gc
    N_mid_l_top[fi] = n_ml
    N_mid_l_bot[fi] = n_ml

    prev_edge_d = N_edge_d_bot[fi]
    prev_mid_d = N_mid_d_bot[fi]
    prev_edge_l = N_edge_l_bot[fi]
    prev_mid_l = N_mid_l_bot[fi]

# 打印验证
for fi, nm in enumerate(floor_names):
    print(f"{nm}: 边柱 M_top={M_edge_d_top[fi]:.2f} M_bot={M_edge_d_bot[fi]:.2f} V={V_edge_d[fi]:.2f} N_top={N_edge_d_top[fi]:.1f} N_bot={N_edge_d_bot[fi]:.1f}")
    print(f"     中柱 M_top={M_mid_d_top[fi]:.2f} M_bot={M_mid_d_bot[fi]:.2f} V={V_mid_d[fi]:.2f} N_top={N_mid_d_top[fi]:.1f} N_bot={N_mid_d_bot[fi]:.1f}")

# ============================================================
# 内力组合函数
# ============================================================
def combo(D, L, W, E):
    c1  = 1.3*D + 1.5*L
    c2  = 1.2*D + 1.4*W + 0.7*1.4*L
    c3  = 1.2*D + 0.6*L + 1.3*E
    c4  = 1.0*D + 1.4*W
    c5  = 1.0*D + 0.5*L + 1.3*E
    c6  = 1.2*D + 1.4*(-W) + 0.7*1.4*L
    c7  = 1.2*D + 0.6*L + 1.3*(-E)
    c8  = 1.0*D + 1.4*(-W)
    c9  = 1.0*D + 0.5*L + 1.3*(-E)
    c10 = max(abs(c1),abs(c2),abs(c3),abs(c4),abs(c5),abs(c6),abs(c7),abs(c8),abs(c9))
    c11 = max(abs(c1),abs(c4),abs(c5))
    c12 = max(abs(c2),abs(c3),abs(c6),abs(c7))
    return [c1,c2,c3,c4,c5,c6,c7,c8,c9,c10,c11,c12]

# ============================================================
# 更新列组合表 (7-15~7-20, T60-T65 = indices 59-64)
# ============================================================
print("\n" + "="*60)
print("更新列组合表 (7-15~7-20)")
print("="*60)

for ti in range(6):
    tbl_idx = 59 + ti
    # 表7-15(idx59)=1F, 表7-20(idx64)=6F → fi=5-ti
    fi = 5 - ti  # ti=0→fi=5(1F), ti=5→fi=0(6F)
    fn = floor_names[fi]
    t_doc = doc.tables[tbl_idx]
    t_bak = bak.tables[tbl_idx]

    print(f"\n表7-{15+ti} ({fn}):")

    for ri in range(4, min(16, len(t_doc.rows))):
        row_doc = t_doc.rows[ri]
        row_bak = t_bak.rows[ri]

        # 确定行列类型: 每3行一组 [M, N, V]
        group = (ri - 4) // 3  # 0=边柱上,1=边柱下,2=中柱上,3=中柱下
        rtype = (ri - 4) % 3   # 0=M,1=N,2=V
        is_edge = (group < 2)
        is_top = (group % 2 == 0)

        # 从4800备份恢复 cells[3](标签), cells[6](W), cells[7](E)
        bak_label = row_bak.cells[3].text.strip() if len(row_bak.cells) > 3 else ''
        bak_W = row_bak.cells[6].text.strip() if len(row_bak.cells) > 6 else ''
        bak_E = row_bak.cells[7].text.strip() if len(row_bak.cells) > 7 else ''

        row_doc.cells[3].text = bak_label
        row_doc.cells[6].text = bak_W
        row_doc.cells[7].text = bak_E

        try:
            W_val = float(bak_W) if bak_W else 0.0
            E_val = float(bak_E) if bak_E else 0.0
        except:
            W_val = 0.0; E_val = 0.0

        # 确定 D 和 L 新值
        if rtype == 0:  # M
            if is_edge:
                D_new = M_edge_d_top[fi] if is_top else M_edge_d_bot[fi]
                L_new = M_edge_l_top[fi] if is_top else M_edge_l_bot[fi]
            else:
                D_new = M_mid_d_top[fi] if is_top else M_mid_d_bot[fi]
                L_new = M_mid_l_top[fi] if is_top else M_mid_l_bot[fi]
        elif rtype == 1:  # N
            if is_edge:
                D_new = N_edge_d_top[fi] if is_top else N_edge_d_bot[fi]
                L_new = N_edge_l_top[fi] if is_top else N_edge_l_bot[fi]
            else:
                D_new = N_mid_d_top[fi] if is_top else N_mid_d_bot[fi]
                L_new = N_mid_l_top[fi] if is_top else N_mid_l_bot[fi]
        else:  # V
            if is_edge:
                D_new = V_edge_d[fi]
                L_new = V_edge_l[fi]
            else:
                D_new = V_mid_d[fi]
                L_new = V_mid_l[fi]

        # 写入 cells[4]=D, cells[5]=L
        row_doc.cells[4].text = f'{D_new:.2f}'
        row_doc.cells[5].text = f'{L_new:.2f}'

        # 重算组合 → cells[8+] (12列)
        combos = combo(D_new, L_new, W_val, E_val)
        for ci, val in enumerate(combos):
            col_idx = 8 + ci
            if col_idx < len(row_doc.cells) and col_idx < 20:
                row_doc.cells[col_idx].text = f'{val:.2f}'

        if ri == 4:
            label_type = ['M','N','V'][rtype]
            print(f"  {label_type}上端(边柱): D={D_new:.2f}, L={L_new:.2f}, W={W_val:.2f}, E={E_val:.2f}")
            print(f"    组合: 1.3D+1.5L={combos[0]:.2f}, 1.2D+1.4W+0.98L={combos[1]:.2f}")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n列组合表修复完成")
print(f"修正版: {DOC}")
print(f"审阅版: {review}")
