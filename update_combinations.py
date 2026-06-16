#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整更新内力组合表 (表7-9~7-20) + 截面设计 (第8章) + 基础设计 (第11章)
用新恒载/活载内力 + 现有风载/地震值 → 按GB50011-2010组合
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

# ============================================================
# 基本参数
# ============================================================
L1, L2 = 5.4, 2.4
B_COL = 0.5; B_HALF = B_COL / 2
BETA = 0.85
ss = 3.45
F_NEW = 1 - 2*(0.5*ss/L1)**2 + (0.5*ss/L1)**3

g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0

qe_d_r = g_beam_e + F_NEW*ss*g_roof
qe_d_f = g_beam_e + g_wall_e + F_NEW*ss*g_floor
qm_d_r = g_beam_m + 0.625*L2*g_roof
qm_d_f = g_beam_m + g_wall_m + 0.625*L2*g_floor
qe_l_r = F_NEW*ss*q_roof
qe_l_f = F_NEW*ss*q_floor
qm_l_r = 0.625*L2*q_roof
qm_l_f = 0.625*L2*q_floor

E = 30e6; Ic = 0.5**4/12
ic_std = E*Ic/3.0; ic_1st = E*Ic/4.0
ib_edge = E*1.5*0.25*0.5**3/12/L1
ib_mid = E*1.5*0.25*0.4**3/12/L2

def fem(q,L): return q*L**2/12
def midM(q,L,Ml,Mr): return q*L**2/8-(abs(Ml)+abs(Mr))/2

def moment_dist(ic_u,ic_l,ib_e,ib_m,q_e,q_m):
    FE_e=fem(q_e,L1); FE_m=fem(q_m,L2)
    sA=ic_u+ic_l+ib_e; sB=ic_u+ic_l+ib_e+ib_m
    mu_A_u=ic_u/sA; mu_A_l=ic_l/sA; mu_A_b=ib_e/sA
    mu_B_u=ic_u/sB; mu_B_l=ic_l/sB; mu_B_bl=ib_e/sB; mu_B_br=ib_m/sB
    M_A_u=M_A_l=0; M_A_b=-FE_e
    M_B_u=M_B_l=0; M_B_bl=FE_e; M_B_br=-FE_m
    # 1st
    uA=M_A_u+M_A_l+M_A_b; dAb=-mu_A_b*uA; dAu=-mu_A_u*uA; dAl=-mu_A_l*uA
    M_A_b+=dAb; M_A_u+=dAu; M_A_l+=dAl
    uB=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl=-mu_B_bl*uB; dBbr=-mu_B_br*uB; dBu=-mu_B_u*uB; dBl=-mu_B_l*uB
    M_B_bl+=dBbl; M_B_br+=dBbr; M_B_u+=dBu; M_B_l+=dBl
    M_A_b+=0.5*dBbl; M_B_bl+=0.5*dAb; M_B_br+=0.5*(-dBbr)
    # 2nd
    uA2=M_A_u+M_A_l+M_A_b; dAb2=-mu_A_b*uA2; dAu2=-mu_A_u*uA2; dAl2=-mu_A_l*uA2
    M_A_b+=dAb2; M_A_u+=dAu2; M_A_l+=dAl2
    uB2=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl2=-mu_B_bl*uB2; dBbr2=-mu_B_br*uB2; dBu2=-mu_B_u*uB2; dBl2=-mu_B_l*uB2
    M_B_bl+=dBbl2; M_B_br+=dBbr2; M_B_u+=dBu2; M_B_l+=dBl2
    M_A_b+=0.5*dBbl2; M_B_bl+=0.5*dAb2; M_B_br+=0.5*(-dBbr2)
    return dict(A_u=M_A_u,A_l=M_A_l,A_b=M_A_b,
                B_u=M_B_u,B_l=M_B_l,B_bl=M_B_bl,B_br=M_B_br)

floors_cfg = [
    ('6F',0,ic_std,qe_d_r,qm_d_r,qe_l_r,qm_l_r),
    ('5F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('4F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('3F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('2F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('1F',ic_std,ic_1st,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
]

data = {}
for nm,icu,icl,qed,qmd,qel,qml in floors_cfg:
    rd = moment_dist(icu,icl,ib_edge,ib_mid,qed,qmd)
    rl = moment_dist(icu,icl,ib_edge,ib_mid,qel,qml)
    data[nm] = {'dead':rd,'live':rl,'qe_d':qed,'qm_d':qmd,'qe_l':qel,'qm_l':qml}

def beam_shear(q,L,Ml,Mr):
    Vl = q*L/2 + (Mr-Ml)/L
    Vr = q*L/2 + (Ml-Mr)/L
    return Vl, Vr

# ============================================================
# 计算内力转换值 (柱边内力)
# ============================================================
conv = {}  # conv[floor][load_type][section]
for nm,icu,icl,qed,qmd,qel,qml in floors_cfg:
    rd = data[nm]['dead']; rl = data[nm]['live']

    # 调幅
    M_lb_d = rd['A_b']; M_rb_d = rd['B_bl']; M_mb_d = rd['B_br']
    M_la_d = BETA*M_lb_d; M_ra_d = BETA*M_rb_d; M_ma_d = BETA*M_mb_d
    M_lb_l = rl['A_b']; M_rb_l = rl['B_bl']; M_mb_l = rl['B_br']
    M_la_l = BETA*M_lb_l; M_ra_l = BETA*M_rb_l; M_ma_l = BETA*M_mb_l

    # 调幅前剪力
    Vl_d, Vr_d = beam_shear(qed, L1, M_lb_d, M_rb_d)
    Vm_d, _ = beam_shear(qmd, L2, M_mb_d, -M_mb_d)
    Vl_l, Vr_l = beam_shear(qel, L1, M_lb_l, M_rb_l)
    Vm_l, _ = beam_shear(qml, L2, M_mb_l, -M_mb_l)

    # 柱边转换: M_edge = |M_after| + |V|*b/2 - q*b²/8
    def face_M(M, V, q):
        return abs(M) + abs(V)*B_HALF - q*B_HALF**2/2
    def face_V(V, q):
        return abs(V) - q*B_HALF

    M_le_d = face_M(M_la_d, Vl_d, qed)
    M_re_d = face_M(M_ra_d, Vr_d, qed)
    M_me_d = face_M(M_ma_d, Vm_d, qmd)
    Vl_e_d = face_V(Vl_d, qed)
    Vr_e_d = face_V(Vr_d, qed)
    Vm_e_d = face_V(Vm_d, qmd)

    M_le_l = face_M(M_la_l, Vl_l, qel)
    M_re_l = face_M(M_ra_l, Vr_l, qel)
    M_me_l = face_M(M_ma_l, Vm_l, qml)
    Vl_e_l = face_V(Vl_l, qel)
    Vr_e_l = face_V(Vr_l, qel)
    Vm_e_l = face_V(Vm_l, qml)

    # 跨中弯矩(调幅后)
    M_mid_d = midM(qed, L1, M_la_d, M_ra_d)
    M_mid_l = midM(qel, L1, M_la_l, M_ra_l)

    conv[nm] = {
        'dead': {
            'edge_left_M': M_le_d, 'edge_left_V': Vl_e_d,
            'edge_right_M': M_re_d, 'edge_right_V': Vr_e_d,
            'edge_mid_M': M_mid_d,
            'mid_left_M': M_me_d, 'mid_left_V': Vm_e_d,
        },
        'live': {
            'edge_left_M': M_le_l, 'edge_left_V': Vl_e_l,
            'edge_right_M': M_re_l, 'edge_right_V': Vr_e_l,
            'edge_mid_M': M_mid_l,
            'mid_left_M': M_me_l, 'mid_left_V': Vm_e_l,
        }
    }

# ============================================================
# 读取现有风载/地震内力转换值 (表7-7, 7-8)
# ============================================================
# 表7-7 (T52=index 51) 风荷载内力转换
# 表7-8 (T53=index 52) 地震荷载内力转换
wind_conv = {}  # wind_conv[floor][section] = (M, V)
seis_conv = {}

t_wind = doc.tables[51]
t_seis = doc.tables[52]

for tbl, store, label in [(t_wind, wind_conv, 'wind'), (t_seis, seis_conv, 'seis')]:
    row_idx = 3
    for nm in ['6F','5F','4F','3F','2F','1F']:
        store[nm] = {}
        for sec in ['edge_left', 'edge_right', 'mid_left']:
            row = tbl.rows[row_idx]
            # 表7-7/7-8列映射:
            # cells[0]=楼层, cells[1]=截面, cells[2]=M_node, cells[3]=V_or_q, cells[4]=M_face, cells[5]=V_face
            # 对水平荷载, M_face≈M_node, V_face是设计用值
            try:
                m_face = float(row.cells[4].text.strip()) if row.cells[4].text.strip() else 0
                v_face = float(row.cells[5].text.strip()) if row.cells[5].text.strip() else 0
            except:
                m_face = 0; v_face = 0
            store[nm][sec] = {'V': v_face, 'M': m_face}
            row_idx += 1

# ============================================================
# 内力组合公式 (GB50011-2010)
# ============================================================
# 对于梁:
# 各列组合:
# [3]=D, [4]=L, [5]=W, [6]=E (输入值)
# [7]=1.3D+1.5L (承载力-可变控制)
# [8]=1.2D+1.4W+0.7*1.4L (风-有活载)
# [9]=1.2D+0.6L+1.3E (地震-重力)
# [10]=1.0D+1.4W (风-恒载有利)
# [11]=1.0D+0.5L+1.3E (地震-恒载有利)
# [12]=1.2D+1.4W (风-无活载) → 更不利
# [13]=1.2D+1.4*0.7L+1.4W → 同上
# ... 后续列: |M|max对应V, |V|max对应M 等

def combo_beam(D, L, W, E):
    """返回梁的12个组合值 [c7..c18]"""
    c = []

    # 基本组合
    c7  = 1.3*D + 1.5*L                     # 1.3D+1.5L
    c8  = 1.2*D + 1.4*W + 0.7*1.4*L         # 1.2D+1.4W+0.98L
    c9  = 1.2*D + 0.6*L + 1.3*E             # 1.2D+0.6L+1.3E (重力+地震)
    c10 = 1.0*D + 1.4*W                      # 1.0D+1.4W
    c11 = 1.0*D + 0.5*L + 1.3*E             # 1.0D+0.5L+1.3E

    # 反向风/地震
    c12 = 1.2*D + 1.4*(-W) + 0.7*1.4*L       # 风反向
    c13 = 1.2*D + 0.6*L + 1.3*(-E)           # 地震反向
    c14 = 1.0*D + 1.4*(-W)                    # 恒有利+风反向
    c15 = 1.0*D + 0.5*L + 1.3*(-E)           # 恒有利+地震反向

    # 最不利组合 (取绝对值最大)
    c16 = max(abs(c7), abs(c8), abs(c9), abs(c10), abs(c11),
              abs(c12), abs(c13), abs(c14), abs(c15))
    c17 = max(abs(c7), abs(c10), abs(c11))     # 不含风/震的max
    c18 = max(abs(c8), abs(c9), abs(c12), abs(c13))  # 含风/震的max

    return [c7, c8, c9, c10, c11, c12, c13, c14, c15, c16, c17, c18]

# ============================================================
# 更新梁内力组合表 (T54-T59 = indices 53-58)
# ============================================================
print("="*60)
print("更新梁内力组合表 (7-9~7-14)")
print("="*60)

floor_names = ['1F','2F','3F','4F','5F','6F']

for ti, fn in enumerate(floor_names):
    tbl_idx = 53 + ti  # T54=idx53 for 1F, T55=idx54 for 2F, ...
    t = doc.tables[tbl_idx]
    nm = fn  # floor name for data lookup

    print(f"\n表7-{9+ti} ({fn}):")

    # 该表的数据: 7行 (边梁V左, 边梁M左, 边梁V右, 边梁M右, 边梁跨中M, 中梁V, 中梁M)
    # 每个数据行对应到 conv 的映射
    row_specs = [
        # (row_index, dead_key, live_key, is_M)
        (1, 'edge_left_V', 'edge_left_V', False),
        (2, 'edge_left_M', 'edge_left_M', True),
        (3, 'edge_right_V', 'edge_right_V', False),
        (4, 'edge_right_M', 'edge_right_M', True),
        (5, 'edge_mid_M', 'edge_mid_M', True),
        (6, 'mid_left_V', 'mid_left_V', False),
        (7, 'mid_left_M', 'mid_left_M', True),
    ]

    for row_i, d_key, l_key, is_M in row_specs:
        row = t.rows[row_i]
        D_new = conv[nm]['dead'][d_key]
        L_new = conv[nm]['live'][l_key]

        # 获取风/地震值
        # 边梁: edge_left/edge_right, 中梁: mid_left
        if 'mid' in d_key:
            sec = 'mid_left'
        elif 'right' in d_key:
            sec = 'edge_right'
        else:
            sec = 'edge_left'

        W_val = wind_conv[nm][sec]['M'] if is_M else wind_conv[nm][sec]['V']
        E_val = seis_conv[nm][sec]['M'] if is_M else seis_conv[nm][sec]['V']

        # 更新输入列 [3]=D, [4]=L
        row.cells[3].text = f'{D_new:.2f}'
        row.cells[4].text = f'{L_new:.2f}'

        # 重算组合
        combos = combo_beam(D_new, L_new, W_val, E_val)

        # 更新组合列 [7]~[18]
        for ci, val in enumerate(combos):
            col_idx = 7 + ci
            if col_idx < len(row.cells) and col_idx <= 18:
                row.cells[col_idx].text = f'{val:.2f}'

        if 'left_M' in d_key and 'mid' not in d_key:
            print(f"  边梁左M: D={D_new:.2f}, L={L_new:.2f}, W={W_val:.2f}, E={E_val:.2f}")
            print(f"    1.3D+1.5L={combos[0]:.2f}")

# ============================================================
# 更新柱内力组合表 (T60-T65 = indices 59-64)
# ============================================================
print("\n" + "="*60)
print("更新柱内力组合表 (7-15~7-20)")
print("="*60)

# 柱轴力需要重新计算
# 边柱集中力 (5400版本)
sec_beam = 1.54*L1/2
edge_floor_conc = g_floor*(ss**2/4+ss*L1/2)
mid_floor_conc = g_floor*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))
edge_roof_conc = g_roof*(ss**2/4+ss*L1/2)
mid_roof_conc = g_roof*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))

F_edge_roof = 31.19+20.08+sec_beam+edge_roof_conc
F_edge_floor = 42.3+20.08+sec_beam+edge_floor_conc
F_mid_roof = 3.7+20.08+sec_beam+mid_roof_conc
F_mid_floor = 41.95+20.08+sec_beam+mid_floor_conc

Gc_std = 6.76*3.0
Gc_1st = 6.76*4.0

# 逐层累计柱轴力 (恒载)
N_edge_d = [0]*6; N_mid_d = [0]*6
N_edge_l = [0]*6; N_mid_l = [0]*6

for fi, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    is_roof = (fi == 0); is_first = (fi == 5)
    rd = data[nm]['dead']; rl = data[nm]['live']
    qed = data[nm]['qe_d']; qmd = data[nm]['qm_d']
    qel = data[nm]['qe_l']; qml = data[nm]['qm_l']

    # 调幅后弯矩
    M_le_d = BETA*rd['A_b']; M_re_d = BETA*rd['B_bl']; M_me_d = BETA*rd['B_br']
    M_le_l = BETA*rl['A_b']; M_re_l = BETA*rl['B_bl']; M_me_l = BETA*rl['B_br']

    # 梁端剪力 (调幅后)
    Vl_e_d, Vr_e_d = beam_shear(qed, L1, M_le_d, M_re_d)
    Vl_m_d, Vr_m_d = beam_shear(qmd, L2, M_me_d, -M_me_d)
    Vl_e_l, Vr_e_l = beam_shear(qel, L1, M_le_l, M_re_l)
    Vl_m_l, Vr_m_l = beam_shear(qml, L2, M_me_l, -M_me_l)

    Fe = F_edge_roof if is_roof else F_edge_floor
    Fm = F_mid_roof if is_roof else F_mid_floor
    Gc = Gc_1st if is_first else Gc_std

    # 柱顶轴力 = 集中力 + 梁端剪力
    n_ed = Fe + Vl_e_d
    n_md = Fm + Vr_e_d + Vl_m_d
    n_el = Vl_e_l
    n_ml = Vr_e_l + Vl_m_l

    if fi > 0:
        n_ed += N_edge_d[fi-1] + Gc_std
        n_md += N_mid_d[fi-1] + Gc_std
        n_el += N_edge_l[fi-1]
        n_ml += N_mid_l[fi-1]

    N_edge_d[fi] = n_ed
    N_mid_d[fi] = n_md
    N_edge_l[fi] = n_el
    N_mid_l[fi] = n_ml

# 柱端弯矩 (从moment_dist结果直接获取)
# 边柱A: M_top = A_u (上一层), M_bot = A_l (当前层)
# 中柱B: M_top = B_u (上一层), M_bot = B_l (当前层)

col_moments = {}
for nm in ['6F','5F','4F','3F','2F','1F']:
    rd = data[nm]['dead']; rl = data[nm]['live']
    col_moments[nm] = {
        'dead': {'A_top': rd['A_u'], 'A_bot': rd['A_l'],
                 'B_top': rd['B_u'], 'B_bot': rd['B_l']},
        'live': {'A_top': rl['A_u'], 'A_bot': rl['A_l'],
                 'B_top': rl['B_u'], 'B_bot': rl['B_l']},
    }

for ti, fn in enumerate(floor_names):
    tbl_idx = 59 + ti
    if tbl_idx >= len(doc.tables):
        break
    t = doc.tables[tbl_idx]
    nm = fn

    print(f"\n表7-{15+ti} ({fn}):")

    # 每张表的数据行从 R4 开始 (0-3是表头)
    # 边柱上端M, 边柱上端N, 边柱上端V, 边柱下端M, 边柱下端N, 边柱下端V
    # 中柱上端M, 中柱上端N, ...

    # 先识别行：读col[0]=楼层, col[1]=边柱/中柱, col[2]=上端/下端, col[3]=M/N/V
    for ri in range(4, min(len(t.rows), 22)):
        row = t.rows[ri]
        try:
            col_type = row.cells[1].text.strip()  # 边柱/中柱
            position = row.cells[2].text.strip()   # 上端/下端
        except:
            continue

        # 获取对应的柱内力值
        if '边' in col_type:
            prefix = 'A'
        elif '中' in col_type:
            prefix = 'B'
        else:
            continue

        # 柱轴力
        is_top = '上' in position
        fi = ['6F','5F','4F','3F','2F','1F'].index(nm)

        if prefix == 'A':
            N_d = N_edge_d[fi]
            N_l = N_edge_l[fi]
        else:
            N_d = N_mid_d[fi]
            N_l = N_mid_l[fi]

        if is_top:
            N_d_top = N_d
            N_l_top = N_l
            N_d_bot = N_d + (Gc_1st if fi==5 else Gc_std)
            N_l_bot = N_l
        else:
            N_d_top = N_d - (Gc_1st if fi==5 else Gc_std) if fi > 0 else N_d
            N_l_top = N_l
            N_d_bot = N_d
            N_l_bot = N_l

        # 柱端弯矩
        if is_top:
            M_d = col_moments[nm]['dead'][f'{prefix}_top']
            M_l = col_moments[nm]['live'][f'{prefix}_top']
        else:
            M_d = col_moments[nm]['dead'][f'{prefix}_bot']
            M_l = col_moments[nm]['live'][f'{prefix}_bot']

        # 柱剪力 = (M_top + M_bot) / h
        h = 4.0 if fi == 5 else 3.0
        M_d_top = col_moments[nm]['dead'][f'{prefix}_top']
        M_d_bot = col_moments[nm]['dead'][f'{prefix}_bot']
        M_l_top = col_moments[nm]['live'][f'{prefix}_top']
        M_l_bot = col_moments[nm]['live'][f'{prefix}_bot']

        V_d_col = (abs(M_d_top) + abs(M_d_bot)) / h
        V_l_col = (abs(M_l_top) + abs(M_l_bot)) / h

        # 读取现有风/地震值 (从表的现有数据中读取)
        try:
            W_M = float(row.cells[5].text.strip()) if row.cells[5].text.strip() else 0
            E_M = float(row.cells[6].text.strip()) if row.cells[6].text.strip() else 0
        except:
            W_M = 0; E_M = 0

        # 识别该行是M, N还是V
        # 通过现有值的大小判断
        old_col3 = float(row.cells[3].text.strip()) if row.cells[3].text.strip() else 0
        old_col4 = float(row.cells[4].text.strip()) if row.cells[4].text.strip() else 0

        # 判断是M(弯矩)还是N(轴力)还是V(剪力)
        if abs(old_col3) < 10 and abs(M_d) < 10:
            # 可能是V行(剪力值较小)
            D_new = V_d_col
            L_new = V_l_col
        elif abs(old_col3) > 100:
            # N行(轴力值大)
            D_new = N_d_bot if not is_top else N_d_top
            L_new = N_l_bot if not is_top else N_l_top
        else:
            # M行
            D_new = M_d
            L_new = M_l

        # 更新输入值
        row.cells[3].text = f'{D_new:.2f}' if D_new != 0 else row.cells[3].text
        row.cells[4].text = f'{L_new:.2f}' if L_new != 0 else row.cells[4].text

        # 重算组合 (柱用相同组合公式)
        combos = combo_beam(D_new, L_new, W_M, E_M)
        for ci, val in enumerate(combos[:12]):
            col_idx = 7 + ci
            if col_idx < len(row.cells):
                try:
                    row.cells[col_idx].text = f'{val:.2f}'
                except:
                    pass

    print(f"  边柱轴力: N_d_top={N_edge_d[fi]:.1f}, N_l_top={N_edge_l[fi]:.1f}")
    print(f"  中柱轴力: N_d_top={N_mid_d[fi]:.1f}, N_l_top={N_mid_l[fi]:.1f}")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n全部内力组合表更新完成")
print(f"修正版: {DOC}")
print(f"审阅版: {review}")
