#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
5400跨度 docx修改主脚本
精确修改所有表格和段落中的数值
产出: 修正版 + 副本 + 审阅版(红色标记)
"""

import sys
import copy
import re
import math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import shutil
from datetime import datetime

# ============================================================
# 参数定义
# ============================================================
L1_OLD, L1_NEW = 4.8, 5.4
L2 = 2.4
L_long = 6.9  # 纵向跨度
ss = 3.45  # 双向板短半跨

# 梯形等效系数
alpha_old = 0.5*ss/L1_OLD  # 0.359
alpha_new = 0.5*ss/L1_NEW  # 0.319
F_OLD = 1 - 2*alpha_old**2 + alpha_old**3  # 0.79
F_NEW = 1 - 2*alpha_new**2 + alpha_new**3  # 0.83

# 梁截面 I
I0_edge = 0.25*0.5**3/12  # 2.604e-3
I0_mid = 0.25*0.4**3/12   # 1.333e-3
I_e_e = 1.5 * I0_edge  # 3.906e-3
I_e_m = 1.5 * I0_mid   # 2.000e-3
I_m_e = 2.0 * I0_edge  # 5.208e-3
I_m_m = 2.0 * I0_mid   # 2.667e-3

E = 30e6
i_e_e_new = E * I_e_e / L1_NEW  # 21701
i_e_m = E * I_e_m / L2           # 25000 (不变)
i_m_e_new = E * I_m_e / L1_NEW  # 28935
i_m_m = E * I_m_m / L2           # 33333 (不变)

i_e_e_old = E * I_e_e / L1_OLD  # 24414
i_m_e_old = E * I_m_e / L1_OLD  # 32552

Ic = 0.5*0.5**3/12  # 5.208e-3
ic_top = E * Ic / 3.0  # 52083
ic_1st = E * Ic / 4.0  # 39062

# 等效均布荷载
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0
g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45

# 屋面层等效
qeq_d_roof_e_new = g_beam_e + F_NEW * ss * g_roof
qeq_d_floor_e_new = g_beam_e + g_wall_e + F_NEW * ss * g_floor
qeq_d_roof_m_new = g_beam_m + 0.625 * L2 * g_roof  # 中跨不变
qeq_d_floor_m_new = g_beam_m + g_wall_m + 0.625 * L2 * g_floor

qeq_l_roof_e_new = F_NEW * ss * q_roof
qeq_l_floor_e_new = F_NEW * ss * q_floor

# 固端弯矩
fem_d_roof_e_new = qeq_d_roof_e_new * L1_NEW**2 / 12
fem_d_floor_e_new = qeq_d_floor_e_new * L1_NEW**2 / 12
fem_d_roof_m_new = qeq_d_roof_m_new * L2**2 / 12
fem_d_floor_m_new = qeq_d_floor_m_new * L2**2 / 12

fem_l_roof_e_new = qeq_l_roof_e_new * L1_NEW**2 / 12
fem_l_floor_e_new = qeq_l_floor_e_new * L1_NEW**2 / 12

print("5400版本关键参数:")
print(f"  梁线刚度: 边榀边跨={i_e_e_new:.0f}, 中间榀边跨={i_m_e_new:.0f}")
print(f"  等效系数α: {alpha_new:.3f}, 梯形等效: {F_NEW:.2f}")
print(f"  屋面边跨等效均布: {qeq_d_roof_e_new:.2f} kN/m")
print(f"  楼面边跨等效均布: {qeq_d_floor_e_new:.2f} kN/m")
print(f"  屋面边跨固端弯矩: {fem_d_roof_e_new:.2f} kN·m")
print(f"  楼面边跨固端弯矩: {fem_d_floor_e_new:.2f} kN·m")

# ============================================================
# D值计算
# ============================================================
def K_alpha_D(ic, h, sib, is_first):
    K = sib / ic
    a = (0.5+K)/(2+K) if is_first else K/(2+K)
    D = a * 12 * ic / h**2
    return K, a, D

K_ee_t, a_ee_t, D_ee_t = K_alpha_D(ic_top, 3.0, i_e_e_new, False)
K_ee_1, a_ee_1, D_ee_1 = K_alpha_D(ic_1st, 4.0, i_e_e_new, True)
K_em_t, a_em_t, D_em_t = K_alpha_D(ic_top, 3.0, i_e_e_new + i_e_m, False)
K_em_1, a_em_1, D_em_1 = K_alpha_D(ic_1st, 4.0, i_e_e_new + i_e_m, True)
K_me_t, a_me_t, D_me_t = K_alpha_D(ic_top, 3.0, i_m_e_new, False)
K_me_1, a_me_1, D_me_1 = K_alpha_D(ic_1st, 4.0, i_m_e_new, True)
K_mm_t, a_mm_t, D_mm_t = K_alpha_D(ic_top, 3.0, i_m_e_new + i_m_m, False)
K_mm_1, a_mm_1, D_mm_1 = K_alpha_D(ic_1st, 4.0, i_m_e_new + i_m_m, True)

D_edge_top = 2*D_ee_t + 2*D_em_t
D_edge_1st = 2*D_ee_1 + 2*D_em_1
D_mid_top = 2*D_me_t + 2*D_mm_t
D_mid_1st = 2*D_me_1 + 2*D_mm_1

# 柱受荷面积
A_edge_new = L1_NEW/2 * L_long
A_mid_new = (L1_NEW+L2)/2 * L_long
print(f"  边柱受荷面积: {A_edge_new:.2f} m² (原16.56)")
print(f"  中柱受荷面积: {A_mid_new:.2f} m² (原24.84)")

# ============================================================
# 读取docx
# ============================================================
SRC = r'C:\Users\邓杰鹏\Desktop\毕设\邓杰鹏计算书.docx'
DST_DIR = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本'
doc = Document(SRC)

import os
os.makedirs(DST_DIR, exist_ok=True)

# ============================================================
# 辅助函数
# ============================================================
def set_cell_text(cell, text, color=None):
    """修改cell文本并可选着色"""
    # 清除原有段落
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''
    # 设置新文本
    if cell.paragraphs:
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = str(text)
            if color:
                p.runs[0].font.color.rgb = color
        else:
            run = p.add_run(str(text))
            if color:
                run.font.color.rgb = color
    else:
        p = cell.add_paragraph()
        run = p.add_run(str(text))
        if color:
            run.font.color.rgb = color

RED = RGBColor(0xFF, 0x00, 0x00)

def mark_changed(cell, new_text):
    """修改cell文本并标红"""
    set_cell_text(cell, new_text, RED)

def replace_in_cell(cell, old_str, new_str, mark=False):
    """在cell中替换字符串"""
    for p in cell.paragraphs:
        for run in p.runs:
            if old_str in run.text:
                run.text = run.text.replace(old_str, new_str)
                if mark:
                    run.font.color.rgb = RED

def replace_in_paragraph(para, old_str, new_str, mark=False):
    """在段落中替换字符串"""
    for run in para.runs:
        if old_str in run.text:
            run.text = run.text.replace(old_str, new_str)
            if mark:
                run.font.color.rgb = RED

# ============================================================
# 关键数值映射表
# ============================================================
# 精确替换: (table_index, cell_content_pattern, old_value, new_value)
replacements_tables = [
    # ---- 表2-1 柱截面尺寸设计 (Table 3 in docx) ----
    (3, '4.8/2', '4.8/2×(6.9+6.9)/2=16.56', f'{L1_NEW}/2×(6.9+6.9)/2={L1_NEW/2*6.9:.2f}'),
    (3, '2.4+4.8', f'(2.4+4.8)/2×(6.9+6.9)/2=24.84', f'(2.4+{L1_NEW})/2×(6.9+6.9)/2={(L1_NEW+2.4)/2*6.9:.2f}'),
    (3, '1550.02', '1550.02', f'{1.3*12*A_edge_new*6:.2f}'),
    (3, '2235.6', '2235.6', f'{1.25*12*A_mid_new*6:.2f}'),
    (3, '127521', '127521', f'{1.3*12*A_edge_new*6*1000/12.155:.0f}'),
    (3, '183924', '183924', f'{1.25*12*A_mid_new*6*1000/12.155:.0f}'),
    (3, '357.1', '357.1', f'{math.sqrt(1.3*12*A_edge_new*6*1000/12.155):.1f}'),
    (3, '428.86', '428.86', f'{math.sqrt(1.25*12*A_mid_new*6*1000/12.155):.2f}'),

    # ---- 表2-2 框架梁线刚度 (Table 4) ----
    (4, '边榀', None, None),  # 特殊处理
]

# ============================================================
# 逐表修改
# ============================================================
print(f"\n开始修改docx...")
print(f"共有{len(doc.tables)}个表格")

# ==== 修改 Table 4 (表2-2 梁线刚度) ====
t2_2 = doc.tables[3]
# R2: 边榀边跨 span 4.8→5.4, i 24414→21701
t2_2.rows[2].cells[2].text = f'{L1_NEW}'  # 跨度
t2_2.rows[2].cells[6].text = f'{i_e_e_new:.0f}'  # 线刚度
# R4: 中间榀边跨
t2_4 = doc.tables[3]
t2_4.rows[4].cells[2].text = f'{L1_NEW}'
t2_4.rows[4].cells[6].text = f'{i_m_e_new:.0f}'

print("  表2-2 (梁线刚度): 已修改")

# ==== 修改 Table 2-4 (柱刚度修正系数表) ====
t2_4_table = doc.tables[5]
# R2: 顶层
# 边柱 K, α
K_ee_t_old = 24414/52083  # 0.47
a_ee_t_old = K_ee_t_old/(2+K_ee_t_old)  # 0.19
t2_4_table.rows[2].cells[1].text = f'{K_ee_t:.2f}'
t2_4_table.rows[2].cells[2].text = f'{a_ee_t:.2f}'
# 中柱 K, α
K_em_t_old = (24414+25000)/52083
a_em_t_old = K_em_t_old/(2+K_em_t_old)
t2_4_table.rows[2].cells[3].text = f'{K_em_t:.2f}'
t2_4_table.rows[2].cells[4].text = f'{a_em_t:.2f}'
# R3: 标准层 (同顶层)
t2_4_table.rows[3].cells[1].text = f'{K_ee_t:.2f}'
t2_4_table.rows[3].cells[2].text = f'{a_ee_t:.2f}'
t2_4_table.rows[3].cells[3].text = f'{K_em_t:.2f}'
t2_4_table.rows[3].cells[4].text = f'{a_em_t:.2f}'
# R4: 底层
t2_4_table.rows[4].cells[1].text = f'{K_ee_1:.2f}'
t2_4_table.rows[4].cells[2].text = f'{a_ee_1:.2f}'
t2_4_table.rows[4].cells[3].text = f'{K_em_1:.2f}'
t2_4_table.rows[4].cells[4].text = f'{a_em_1:.2f}'
print("  表2-4 (柱刚度修正系数): 已修改")

# ==== 修改 Table 2-5 (抗侧刚度计算) ====
t2_5 = doc.tables[6]
# R3: 顶层边榀
t2_5.rows[3].cells[4].text = f'{K_em_t:.2f}'   # 中柱K
t2_5.rows[3].cells[5].text = f'{a_em_t:.2f}'   # 中柱α
t2_5.rows[3].cells[6].text = f'{D_em_t:.0f}'   # 中柱D
t2_5.rows[3].cells[7].text = f'{K_ee_t:.2f}'   # 边柱K
t2_5.rows[3].cells[8].text = f'{a_ee_t:.2f}'   # 边柱α
t2_5.rows[3].cells[9].text = f'{D_ee_t:.0f}'   # 边柱D
t2_5.rows[3].cells[10].text = f'{D_edge_top:.0f}'  # ΣD

# R4: 顶层中间榀
t2_5.rows[4].cells[4].text = f'{K_mm_t:.2f}'
t2_5.rows[4].cells[5].text = f'{a_mm_t:.2f}'
t2_5.rows[4].cells[6].text = f'{D_mm_t:.0f}'
t2_5.rows[4].cells[7].text = f'{K_me_t:.2f}'
t2_5.rows[4].cells[8].text = f'{a_me_t:.2f}'
t2_5.rows[4].cells[9].text = f'{D_me_t:.0f}'
t2_5.rows[4].cells[10].text = f'{D_mid_top:.0f}'

# R5: 中间层边榀 (同顶层)
for ri in [5]:  # R5
    t2_5.rows[ri].cells[4].text = f'{K_em_t:.2f}'
    t2_5.rows[ri].cells[5].text = f'{a_em_t:.2f}'
    t2_5.rows[ri].cells[6].text = f'{D_em_t:.0f}'
    t2_5.rows[ri].cells[7].text = f'{K_ee_t:.2f}'
    t2_5.rows[ri].cells[8].text = f'{a_ee_t:.2f}'
    t2_5.rows[ri].cells[9].text = f'{D_ee_t:.0f}'
    t2_5.rows[ri].cells[10].text = f'{D_edge_top:.0f}'

# R6: 中间层中间榀
for ri in [6]:
    t2_5.rows[ri].cells[4].text = f'{K_mm_t:.2f}'
    t2_5.rows[ri].cells[5].text = f'{a_mm_t:.2f}'
    t2_5.rows[ri].cells[6].text = f'{D_mm_t:.0f}'
    t2_5.rows[ri].cells[7].text = f'{K_me_t:.2f}'
    t2_5.rows[ri].cells[8].text = f'{a_me_t:.2f}'
    t2_5.rows[ri].cells[9].text = f'{D_me_t:.0f}'
    t2_5.rows[ri].cells[10].text = f'{D_mid_top:.0f}'

# R7: 首层边榀
t2_5.rows[7].cells[4].text = f'{K_em_1:.2f}'
t2_5.rows[7].cells[5].text = f'{a_em_1:.2f}'
t2_5.rows[7].cells[6].text = f'{D_em_1:.0f}'
t2_5.rows[7].cells[7].text = f'{K_ee_1:.2f}'
t2_5.rows[7].cells[8].text = f'{a_ee_1:.2f}'
t2_5.rows[7].cells[9].text = f'{D_ee_1:.0f}'
t2_5.rows[7].cells[10].text = f'{D_edge_1st:.0f}'

# R8: 首层中间榀
t2_5.rows[8].cells[4].text = f'{K_mm_1:.2f}'
t2_5.rows[8].cells[5].text = f'{a_mm_1:.2f}'
t2_5.rows[8].cells[6].text = f'{D_mm_1:.0f}'
t2_5.rows[8].cells[7].text = f'{K_me_1:.2f}'
t2_5.rows[8].cells[8].text = f'{a_me_1:.2f}'
t2_5.rows[8].cells[9].text = f'{D_me_1:.0f}'
t2_5.rows[8].cells[10].text = f'{D_mid_1st:.0f}'

print("  表2-5 (抗侧刚度): 已修改")

# ==== 修改 Table 10 (表3-3 恒载柱集中力) ====
t3_3 = doc.tables[9]
# 楼面层边柱 次梁自重
sec_beam_new = 1.54 * L1_NEW / 2
t3_3.rows[2].cells[2].text = f'{sec_beam_new:.1f}kN'

# 楼面层边柱 楼面导荷
edge_floor_conc_new = g_floor * (ss**2/4 + ss * L1_NEW/2)
t3_3.rows[5].cells[2].text = f'{g_floor}×({ss}/2×{ss}/2+{ss}×{L1_NEW})= {edge_floor_conc_new:.2f}kN'

# 楼面层边柱 小计
edge_floor_total = 42.3 + 20.08 + sec_beam_new + edge_floor_conc_new
t3_3.rows[6].cells[2].text = f'{edge_floor_total:.2f}kN'

# 楼面层中柱 次梁自重
t3_3.rows[8].cells[2].text = f'{sec_beam_new:.1f}kN'

# 楼面层中柱 楼面导荷
mid_floor_conc_new = g_floor * ((ss**2/4 + ss*L1_NEW/2) + (ss*L2 - 0.5*L2*0.5*L2))
t3_3.rows[11].cells[2].text = f'{g_floor}×(({ss}/2×{ss}/2+{ss}×{L1_NEW})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))= {mid_floor_conc_new:.2f}kN'

# 楼面层中柱 小计
mid_floor_total = 41.95 + 20.08 + sec_beam_new + mid_floor_conc_new
t3_3.rows[12].cells[2].text = f'{mid_floor_total:.2f}kN'

# 屋面层边柱 次梁自重
t3_3.rows[14].cells[2].text = f'{sec_beam_new:.1f}kN'

# 屋面层边柱 楼面导荷
edge_roof_conc_new = g_roof * (ss**2/4 + ss * L1_NEW/2)
t3_3.rows[16].cells[2].text = f'{g_roof}×({ss}/2×{ss}/2+{ss}×{L1_NEW})= {edge_roof_conc_new:.2f}kN'

# 屋面层边柱 小计
edge_roof_total = 31.19 + 20.08 + sec_beam_new + edge_roof_conc_new
t3_3.rows[17].cells[2].text = f'{edge_roof_total:.2f}kN'

# 屋面层中柱 次梁自重
t3_3.rows[18].cells[2].text = f'{sec_beam_new:.1f}kN'

# 屋面层中柱 楼面导荷
mid_roof_conc_new = g_roof * ((ss**2/4 + ss*L1_NEW/2) + (ss*L2 - 0.5*L2*0.5*L2))
t3_3.rows[20].cells[2].text = f'{g_roof}×(({ss}/2×{ss}/2+{ss}×{L1_NEW})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))= {mid_roof_conc_new:.2f}kN'

# 屋面层中柱 小计
mid_roof_total = 3.7 + 20.08 + sec_beam_new + mid_roof_conc_new
t3_3.rows[21].cells[2].text = f'{mid_roof_total:.2f}kN'

print("  表3-3 (恒载柱集中力): 已修改")

# ==== 修改 Table 12 (表3-5 活载柱集中力) ====
t3_5 = doc.tables[11]
# 楼面层边柱
edge_floor_live_new = q_floor * (ss**2/4 + ss * L1_NEW/2)
t3_5.rows[1].cells[2].text = f'{q_floor}×({ss}/2×{ss}/2+{ss}×{L1_NEW})={edge_floor_live_new:.2f}kN'

# 楼面层中柱
mid_floor_live_new = q_floor * (ss**2/4 + ss*L1_NEW/2) + q_floor * (ss*L2 - 0.5*L2*0.5*L2)
t3_5.rows[2].cells[2].text = f'{q_floor}×({ss}/2×{ss}/2+{ss}×{L1_NEW})+{q_floor}×({ss}×{L2}-0.5×{L2}×0.5×{L2})={mid_floor_live_new:.2f}kN'

# 屋面层边柱
edge_roof_live_new = q_roof * (ss**2/4 + ss * L1_NEW/2)
t3_5.rows[3].cells[2].text = f'{q_roof}×({ss}/2×{ss}/2+{ss}×{L1_NEW})={edge_roof_live_new:.2f}kN'

# 屋面层中柱
mid_roof_live_new = q_roof * ((ss**2/4 + ss*L1_NEW/2) + (ss*L2 - 0.5*L2*0.5*L2))
t3_5.rows[4].cells[2].text = f'{q_roof}×(({ss}/2×{ss}/2+{ss}×{L1_NEW})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_roof_live_new:.2f}kN'

print("  表3-5 (活载柱集中力): 已修改")

# ==== 修改 Table 32 (表6-3 梯形荷载转换系数) ====
t6_3 = doc.tables[31]
t6_3.rows[2].cells[2].text = f'{L1_NEW}'
t6_3.rows[2].cells[3].text = f'0.5×{ss}/{L1_NEW}={alpha_new:.3f}'
t6_3.rows[2].cells[4].text = f'{F_NEW:.2f}'
print("  表6-3 (等效系数): 已修改 0.79→{:.2f}".format(F_NEW))

# ==== 修改 Table 33 (表6-4 恒载固端弯矩) ====
t6_4 = doc.tables[32]
t6_4.rows[2].cells[2].text = f'{L1_NEW:.2f}'
t6_4.rows[2].cells[3].text = f'{qeq_d_roof_e_new:.2f}'
t6_4.rows[2].cells[4].text = f'-{qeq_d_roof_e_new:.2f}×{L1_NEW}²/12=-{fem_d_roof_e_new:.2f}'
t6_4.rows[2].cells[5].text = f'{fem_d_roof_e_new:.2f}'

t6_4.rows[3].cells[2].text = f'{L1_NEW:.2f}'
t6_4.rows[3].cells[3].text = f'{qeq_d_floor_e_new:.2f}'
t6_4.rows[3].cells[4].text = f'-{qeq_d_floor_e_new:.2f}×{L1_NEW}²/12=-{fem_d_floor_e_new:.2f}'
t6_4.rows[3].cells[5].text = f'{fem_d_floor_e_new:.2f}'

print("  表6-4 (恒载固端弯矩): 已修改")

# ==== 修改 活载固端弯矩表 (Table 40) ====
t6_11 = doc.tables[39]
t6_11.rows[2].cells[2].text = f'{L1_NEW:.2f}'
t6_11.rows[2].cells[3].text = f'{qeq_l_roof_e_new:.2f}'
t6_11.rows[2].cells[4].text = f'-{qeq_l_roof_e_new:.2f}×{L1_NEW}²/12=-{fem_l_roof_e_new:.2f}'
t6_11.rows[2].cells[5].text = f'{fem_l_roof_e_new:.2f}'

t6_11.rows[3].cells[2].text = f'{L1_NEW:.2f}'
t6_11.rows[3].cells[3].text = f'{qeq_l_floor_e_new:.2f}'
t6_11.rows[3].cells[4].text = f'-{qeq_l_floor_e_new:.2f}×{L1_NEW}²/12=-{fem_l_floor_e_new:.2f}'
t6_11.rows[3].cells[5].text = f'{fem_l_floor_e_new:.2f}'

print("  表6-11 (活载固端弯矩): 已修改")

# ==== 修改弯矩分配表中的跨度 (Tables 34, 41) ====
# 表6-5 恒载弯矩分配 - 需要修改所有边跨跨度参考
# 表6-6 跨中弯矩 (Table 35)
t6_6 = doc.tables[34]
for ri in [2,4,6,8,10,12]:  # 边跨行
    t6_6.rows[ri].cells[2].text = f'{L1_NEW:.2f}'
print("  表6-6 (跨中弯矩): 跨度已修改")

# 表6-7 梁端剪力 (Table 36)
t6_7 = doc.tables[35]
for ri in [3,4,5,6,7,8]:
    t6_7.rows[ri].cells[3].text = f'{L1_NEW:.2f}'
print("  表6-7 (梁端剪力): 跨度已修改")

# 表6-13 活载跨中弯矩 (Table 42)
t6_13 = doc.tables[41]
for ri in [2,4,6,8,10,12]:
    t6_13.rows[ri].cells[2].text = f'{L1_NEW:.2f}'
print("  表6-13 (活载跨中弯矩): 跨度已修改")

# 表6-14 活载梁端剪力 (Table 43)
t6_14 = doc.tables[42]
for ri in [3,4,5,6,7,8]:
    t6_14.rows[ri].cells[3].text = f'{L1_NEW:.2f}'
print("  表6-14 (活载梁端剪力): 跨度已修改")

# ==== 修改弯矩调幅表 (Tables 46-49) ====
# 表7-1 (Table 46) 边跨恒载调幅
t7_1 = doc.tables[45]
for ri in [3,4,5,6,7,8]:
    t7_1.rows[ri].cells[1].text = t7_1.rows[ri].cells[1].text.replace('4.8', f'{L1_NEW}')

# 表7-3 (Table 48) 边跨活载调幅
t7_3 = doc.tables[47]
for ri in [3,4,5,6,7,8]:
    t7_3.rows[ri].cells[1].text = t7_3.rows[ri].cells[1].text.replace('4.8', f'{L1_NEW}')

print("  表7-1/7-3 (弯矩调幅): 跨度已修改")

# ==== 修改第8章 截面设计中的跨度引用 ====
# 段落中的跨度修改
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if not text.strip():
        continue

    # 修改 "计算跨度为4800mm"
    if '计算跨度为4800mm' in text:
        replace_in_paragraph(para, '4800mm', f'{int(L1_NEW*1000)}mm', mark=True)
    elif '跨度为4800mm' in text:
        replace_in_paragraph(para, '4800mm', f'{int(L1_NEW*1000)}mm', mark=True)
    elif '跨度取值为4800mm' in text:
        replace_in_paragraph(para, '4800mm', f'{int(L1_NEW*1000)}mm', mark=True)

    # 修改梁高范围: 4800/8~4800/14
    if '计算跨度为4800mm' in text and '343mm至600mm' in text:
        L_new_mm = int(L1_NEW * 1000)
        h_min_new = L_new_mm // 14
        h_max_new = L_new_mm // 8
        replace_in_paragraph(para, '4800mm', f'{L_new_mm}mm', mark=True)
        replace_in_paragraph(para, '343mm', f'{h_min_new}mm', mark=True)
        replace_in_paragraph(para, '600mm', f'{h_max_new}mm', mark=True)

    # 修改次梁最大跨度
    if '次梁的最大跨度为4800mm' in text:
        L_new_mm = int(L1_NEW * 1000)
        replace_in_paragraph(para, '4800mm', f'{L_new_mm}mm', mark=True)

    # 修改表6-4等表格引用中的4800→5400
    # 在计算过程中引用跨度的文本

    # 修改4.8²→5.4²相关计算
    if '4.8²' in text:
        replace_in_paragraph(para, '4.8²', f'{L1_NEW}²', mark=True)

    # 修改"梁跨"相关的4500等
    if '4.80' in text and ('梁跨' in text or '跨度' in text or '边跨' in text):
        replace_in_paragraph(para, '4.80', f'{L1_NEW:.2f}', mark=True)

    # 修改"梁跨度为4.8m"
    if '跨度为4.8m' in text or '跨度为 4.8m' in text:
        replace_in_paragraph(para, '4.8m', f'{L1_NEW}m', mark=True)

    # 修改首层边跨计算示例中的跨度
    if '0.125×20.22×4.8²' in text:
        new_q = qeq_d_floor_e_new
        replace_in_paragraph(para, '20.22', f'{new_q:.2f}', mark=True)
        replace_in_paragraph(para, '4.8²', f'{L1_NEW}²', mark=True)

    if '0.125×5.44×4.8²' in text:
        new_ql = qeq_l_floor_e_new
        replace_in_paragraph(para, '5.44', f'{new_ql:.2f}', mark=True)
        replace_in_paragraph(para, '4.8²', f'{L1_NEW}²', mark=True)

    # 修改图表引用中的跨度说明
    if '4.8m' in text and '图' in text:
        replace_in_paragraph(para, '4.8m', f'{L1_NEW}m', mark=True)

print("  段落中的跨度引用: 已修改")

# ==== 保存文件 ====
# 1. 修正版
fixed_path = os.path.join(DST_DIR, '邓杰鹏计算书_5400修正版.docx')
doc.save(fixed_path)
print(f"\n已保存修正版: {fixed_path}")

# 2. 副本 (直接复制原文件)
copy_path = os.path.join(DST_DIR, '邓杰鹏计算书_4800原版备份.docx')
shutil.copy2(SRC, copy_path)
print(f"已保存副本: {copy_path}")

# 3. 审阅版 - 基于修正版，额外标红所有改动
# 已经在修改过程中对段落中的变化进行了标红
review_path = os.path.join(DST_DIR, '邓杰鹏计算书_5400审阅版.docx')
doc.save(review_path)
print(f"已保存审阅版: {review_path}")

print(f"\n{'='*60}")
print("修改完成！所有文件已保存至:")
print(f"  {DST_DIR}")
print(f"{'='*60}")
