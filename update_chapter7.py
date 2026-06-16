#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新第7章弯矩调幅表 + 标注剩余工作
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF,0x00,0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

L1, L2 = 5.4, 2.4; ss = 3.45
F_NEW = 1-2*(0.5*ss/L1)**2+(0.5*ss/L1)**3
g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0

# 等效均布
qe_d_r = g_beam_e + F_NEW*ss*g_roof        # 16.75
qe_d_f = g_beam_e+g_wall_e + F_NEW*ss*g_floor  # 20.78
qm_d_r = g_beam_m + 0.625*L2*g_roof        # 9.33
qm_d_f = g_beam_m+g_wall_m + 0.625*L2*g_floor  # 14.64
qe_l_r = F_NEW*ss*q_roof                   # 1.43
qe_l_f = F_NEW*ss*q_floor                  # 5.72
qm_l_r = 0.625*L2*q_roof                   # 0.75
qm_l_f = 0.625*L2*q_floor                  # 3.00

E=30e6; Ic=0.5**4/12
ic_std=E*Ic/3.0; ic_1st=E*Ic/4.0
ib_edge=E*1.5*0.25*0.5**3/12/L1
ib_mid=E*1.5*0.25*0.4**3/12/L2

def fem(q,L): return q*L**2/12
def midM(q,L,Ml,Mr): return q*L**2/8-(abs(Ml)+abs(Mr))/2

def moment_dist(ic_u,ic_l,ib_e,ib_m,q_e,q_m):
    FE_e=fem(q_e,L1); FE_m=fem(q_m,L2)
    sA=ic_u+ic_l+ib_e; sB=ic_u+ic_l+ib_e+ib_m
    mu_A_u=ic_u/sA; mu_A_l=ic_l/sA; mu_A_b=ib_e/sA
    mu_B_u=ic_u/sB; mu_B_l=ic_l/sB; mu_B_bl=ib_e/sB; mu_B_br=ib_m/sB
    M_A_u=M_A_l=0; M_A_b=-FE_e
    M_B_u=M_B_l=0; M_B_bl=FE_e; M_B_br=-FE_m
    # 1st
    uA=M_A_u+M_A_l+M_A_b; dAb=-mu_A_b*uA; dAu=-mu_A_u*uA; dAl=-mu_A_l*uA
    M_A_b+=dAb; M_A_u+=dAu; M_A_l+=dAl
    uB=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl=-mu_B_bl*uB; dBbr=-mu_B_br*uB; dBu=-mu_B_u*uB; dBl=-mu_B_l*uB
    M_B_bl+=dBbl; M_B_br+=dBbr; M_B_u+=dBu; M_B_l+=dBl
    M_A_b+=0.5*dBbl; M_B_bl+=0.5*dAb; M_B_br+=0.5*(-dBbr)
    # 2nd
    uA2=M_A_u+M_A_l+M_A_b; dAb2=-mu_A_b*uA2; dAu2=-mu_A_u*uA2; dAl2=-mu_A_l*uA2
    M_A_b+=dAb2; M_A_u+=dAu2; M_A_l+=dAl2
    uB2=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl2=-mu_B_bl*uB2; dBbr2=-mu_B_br*uB2; dBu2=-mu_B_u*uB2; dBl2=-mu_B_l*uB2
    M_B_bl+=dBbl2; M_B_br+=dBbr2; M_B_u+=dBu2; M_B_l+=dBl2
    M_A_b+=0.5*dBbl2; M_B_bl+=0.5*dAb2; M_B_br+=0.5*(-dBbr2)
    return dict(A_u=M_A_u,A_l=M_A_l,A_b=M_A_b, B_u=M_B_u,B_l=M_B_l,B_bl=M_B_bl,B_br=M_B_br)

floors = [
    ('6F',0,ic_std,qe_d_r,qm_d_r,qe_l_r,qm_l_r),
    ('5F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('4F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('3F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('2F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('1F',ic_std,ic_1st,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
]

data = {}
for nm,icu,icl,qed,qmd,qel,qml in floors:
    rd = moment_dist(icu,icl,ib_edge,ib_mid,qed,qmd)
    rl = moment_dist(icu,icl,ib_edge,ib_mid,qel,qml)
    data[nm] = {'dead':rd,'live':rl,'qe_d':qed,'qm_d':qmd,'qe_l':qel,'qm_l':qml}

# ============================================================
# 弯矩调幅 (0.85系数)
# ============================================================
BETA = 0.85  # 调幅系数
print("="*60)
print("弯矩调幅计算")
print("="*60)

# 表7-1 边跨梁恒载调幅 (Table 46)
t = doc.tables[45]
print("更新表7-1 (边跨恒载调幅)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; qed = data[nm]['qe_d']
    M0 = qed*L1**2/8
    # 调幅前
    M_l = rd['A_b']; M_r = rd['B_bl']; M_mid = midM(qed, L1, M_l, M_r)
    # 调幅后 (梁端乘0.85)
    M_l_a = BETA * M_l; M_r_a = BETA * M_r
    # 跨中调幅后: M0 - (M_l_a+M_r_a)/2, 且不低于 M_mid, 且不低于 M0/2
    M_mid_a = M0 - (abs(M_l_a)+abs(M_r_a))/2
    M_mid_a = max(M_mid_a, M_mid * 1.1)  # at least 10% increase
    M_mid_a = max(M_mid_a, M0/2)

    row = 3+ri
    t.rows[row].cells[1].text = f'{qed:.2f}×{L1}²/8={M0:.2f}'
    t.rows[row].cells[2].text = f'{M_l:.2f}'
    t.rows[row].cells[3].text = f'{M_mid:.2f}'
    t.rows[row].cells[4].text = f'{M_r:.2f}'
    t.rows[row].cells[5].text = f'{M_l_a:.2f}'
    t.rows[row].cells[6].text = f'{M_mid_a:.2f}'
    t.rows[row].cells[7].text = f'{M_r_a:.2f}'
    print(f"  {nm}: M0={M0:.2f}, 调幅前({M_l:.2f},{M_mid:.2f},{M_r:.2f}), 调幅后({M_l_a:.2f},{M_mid_a:.2f},{M_r_a:.2f})")

# 表7-2 中跨梁恒载调幅 (Table 47)
t = doc.tables[46]
print("更新表7-2 (中跨恒载调幅)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; qmd = data[nm]['qm_d']
    M0 = qmd*L2**2/8
    M_l = rd['B_br']; M_r = -M_l  # 对称
    M_mid = midM(qmd, L2, M_l, M_r)
    M_l_a = BETA*M_l; M_r_a = BETA*M_r
    M_mid_a = M0 - (abs(M_l_a)+abs(M_r_a))/2
    M_mid_a = max(M_mid_a, M_mid*1.1, M0/2)

    row = 3+ri
    t.rows[row].cells[1].text = f'{qmd:.2f}×{L2}²/8={M0:.2f}'
    t.rows[row].cells[2].text = f'{M_l:.2f}'
    t.rows[row].cells[3].text = f'{M_mid:.2f}'
    t.rows[row].cells[4].text = f'{M_r:.2f}'
    t.rows[row].cells[5].text = f'{M_l_a:.2f}'
    t.rows[row].cells[6].text = f'{M_mid_a:.2f}'
    t.rows[row].cells[7].text = f'{M_r_a:.2f}'

# 表7-3 边跨梁活载调幅 (Table 48)
t = doc.tables[47]
print("更新表7-3 (边跨活载调幅)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']; qel = data[nm]['qe_l']
    M0 = qel*L1**2/8
    M_l = rl['A_b']; M_r = rl['B_bl']; M_mid = midM(qel, L1, M_l, M_r)
    M_l_a = BETA*M_l; M_r_a = BETA*M_r
    M_mid_a = M0 - (abs(M_l_a)+abs(M_r_a))/2
    M_mid_a = max(M_mid_a, M_mid*1.1, M0/2)

    row = 3+ri
    t.rows[row].cells[1].text = f'{qel:.2f}×{L1}²/8={M0:.2f}'
    t.rows[row].cells[2].text = f'{M_l:.2f}'
    t.rows[row].cells[3].text = f'{M_mid:.2f}'
    t.rows[row].cells[4].text = f'{M_r:.2f}'
    t.rows[row].cells[5].text = f'{M_l_a:.2f}'
    t.rows[row].cells[6].text = f'{M_mid_a:.2f}'
    t.rows[row].cells[7].text = f'{M_r_a:.2f}'

# 表7-4 中跨梁活载调幅 (Table 49)
t = doc.tables[48]
print("更新表7-4 (中跨活载调幅)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']; qml = data[nm]['qm_l']
    M0 = qml*L2**2/8
    M_l = rl['B_br']; M_r = -M_l
    M_mid = midM(qml, L2, M_l, M_r)
    M_l_a = BETA*M_l; M_r_a = BETA*M_r
    M_mid_a = M0 - (abs(M_l_a)+abs(M_r_a))/2
    M_mid_a = max(M_mid_a, M_mid*1.1, M0/2)

    row = 3+ri
    t.rows[row].cells[1].text = f'{qml:.2f}×{L2}²/8={M0:.2f}'
    t.rows[row].cells[2].text = f'{M_l:.2f}'
    t.rows[row].cells[3].text = f'{M_mid:.2f}'
    t.rows[row].cells[4].text = f'{M_r:.2f}'
    t.rows[row].cells[5].text = f'{M_l_a:.2f}'
    t.rows[row].cells[6].text = f'{M_mid_a:.2f}'
    t.rows[row].cells[7].text = f'{M_r_a:.2f}'

print("\n弯矩调幅表 7-1~7-4 全部更新完成")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"文件已保存")
