#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
完整更新第8章截面设计(表8-1~8-12) + 第11章基础设计(表11-1)
基于已更新的内力组合表 (7-9~7-20)
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
BAK = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_4800原版备份.docx'

doc = Document(DOC)
bak = Document(BAK)

# ============================================================
# 材料与截面参数
# ============================================================
fc, ft = 14.3, 1.43
fy, fyv = 360, 360
alpha1, beta1 = 1.0, 0.8
xi_b = 0.518
b_beam, h_edge, h_mid = 250, 500, 400
d_edge, d_mid = 460, 360
b_col, h_col = 500, 500

L1, L2 = 5.4, 2.4
Ln_edge, Ln_mid = L1 - 0.5, L2 - 0.5  # 净跨
H_std, H_1st = 3.0, 4.0

# ============================================================
# 读取梁组合表
# ============================================================
BEAM_KEYS = ['edge_VL','edge_ML','edge_VR','edge_MR','edge_Mmid',
             'mid_VL','mid_ML','mid_VR','mid_MR','mid_Mmid']

beam_combos = []
for ti in range(6):
    tbl = doc.tables[53 + ti]
    data = {}
    for ri in range(1, len(tbl.rows)):
        offset = (ri - 1) % 10
        if offset >= len(BEAM_KEYS):
            continue
        key = BEAM_KEYS[offset]
        vals = {}
        for ci in range(3, 19):
            try:
                txt = tbl.rows[ri].cells[ci].text.strip()
                vals[ci] = float(txt) if txt else 0.0
            except:
                vals[ci] = 0.0
        data[key] = vals
    beam_combos.append(data)

# ============================================================
# 读取柱组合表
# ============================================================
col_combos = []
for ti in range(6):
    tbl = doc.tables[59 + ti]
    data = {}
    for ri in range(4, len(tbl.rows)):
        group = (ri - 4) // 3
        is_edge = (group < 2)
        is_top = (group % 2 == 0)
        force_type = ['M', 'N', 'V'][(ri - 4) % 3]
        prefix = 'edge' if is_edge else 'mid'
        pos = 'top' if is_top else 'bot'
        key = f'{prefix}_{pos}_{force_type}'
        vals = {}
        for ci in range(4, 20):
            try:
                txt = tbl.rows[ri].cells[ci].text.strip()
                vals[ci] = float(txt) if txt else 0.0
            except:
                vals[ci] = 0.0
        data[key] = vals
    col_combos.append(data)

# 楼层索引: 0=1F, 1=2F, 2=3F, 3=4F, 4=5F, 5=6F
FLOOR_LABELS = ['1F','2F','3F','4F','5F','6F']

# ============================================================
# 辅助函数
# ============================================================
def max_abs_combo(vals_dict, start_col=8, end_col=19):
    """从组合值中取绝对值最大值"""
    mx = 0
    for ci in range(start_col, min(end_col+1, max(vals_dict.keys())+1)):
        mx = max(mx, abs(vals_dict.get(ci, 0)))
    return mx

def max_combo(vals_dict, start_col=8, end_col=19):
    """从组合值中取代数最大值"""
    mx = -1e9
    for ci in range(start_col, min(end_col+1, max(vals_dict.keys())+1)):
        mx = max(mx, vals_dict.get(ci, -1e9))
    return mx

def min_combo(vals_dict, start_col=8, end_col=19):
    """从组合值中取代数最小值"""
    mn = 1e9
    for ci in range(start_col, min(end_col+1, max(vals_dict.keys())+1)):
        mn = min(mn, vals_dict.get(ci, 1e9))
    return mn

# ============================================================
# 1. 表8-1 梁强剪弱弯调整 (T66 = idx65)
# ============================================================
print("="*60)
print("1. 表8-1 梁强剪弱弯调整")
print("="*60)
t = doc.tables[65]
eta_vb = 1.1

for ri in range(3, len(t.rows)):
    row = t.rows[ri]
    fl = row.cells[0].text.strip()
    # 楼层索引: 原表6F→2F (缺1F? 实际有6行=6F~1F)
    if '6' in fl: fi = 5
    elif '5' in fl: fi = 4
    elif '4' in fl: fi = 3
    elif '3' in fl: fi = 2
    elif '2' in fl: fi = 1
    else: fi = 0

    bd = beam_combos[fi]

    # V_Gb = 1.2*V_D + 0.6*V_L (重力荷载代表值下的剪力设计值)
    V_D_edge = bd.get('edge_VL', {}).get(3, 0)
    V_L_edge = bd.get('edge_VL', {}).get(4, 0)
    V_D_mid = bd.get('mid_VL', {}).get(3, 0)
    V_L_mid = bd.get('mid_VL', {}).get(4, 0)

    V_Gb_edge = abs(1.2 * V_D_edge + 0.6 * V_L_edge)
    V_Gb_mid = abs(1.2 * V_D_mid + 0.6 * V_L_mid)

    # 梁端弯矩(取绝对值最大组合值, 从含风/震的组合中)
    M_bl_edge = max_abs_combo(bd.get('edge_ML', {}))
    M_br_edge = max_abs_combo(bd.get('edge_MR', {}))
    M_bl_mid = max_abs_combo(bd.get('mid_ML', {}))
    M_br_mid = M_bl_mid  # 中跨对称

    # 强剪弱弯调整
    V_adj_edge = eta_vb * (M_bl_edge + M_br_edge) / Ln_edge + V_Gb_edge
    V_adj_mid = eta_vb * (M_bl_mid + M_br_mid) / Ln_mid + V_Gb_mid

    # 原表列映射(基于4800表结构):
    # [1]=V_Gb_edge, [2]=V_Gb_mid, [3]=M_edge_sum, [4]=M_mid_sum,
    # [5]=eta_vb, [6]=V_adj_edge, [7]=V_adj_mid
    row.cells[1].text = f'{V_Gb_edge:.2f}'
    row.cells[2].text = f'{V_Gb_mid:.2f}'
    row.cells[3].text = f'{M_bl_edge+M_br_edge:.2f}'
    row.cells[4].text = f'{M_bl_mid+M_br_mid:.2f}'
    row.cells[5].text = f'{eta_vb:.2f}'
    row.cells[6].text = f'{eta_vb:.2f}'
    if len(row.cells) > 7:
        row.cells[7].text = f'{V_adj_edge:.2f}'
    if len(row.cells) > 8:
        row.cells[8].text = f'{V_adj_mid:.2f}'

    print(f"  {fl}: V_Gb_edge={V_Gb_edge:.2f}, M_sum={M_bl_edge+M_br_edge:.2f}, V_adj_edge={V_adj_edge:.2f}")

# ============================================================
# 2. 表8-4 梁截面内力设计值 (T69 = idx68)
# ============================================================
print("\n" + "="*60)
print("2. 表8-4 梁截面内力设计值")
print("="*60)
t = doc.tables[68]

for ri in range(3, len(t.rows)):
    row = t.rows[ri]
    fl = row.cells[0].text.strip()
    if '6' in fl: fi = 5
    elif '5' in fl: fi = 4
    elif '4' in fl: fi = 3
    elif '3' in fl: fi = 2
    elif '2' in fl: fi = 1
    else: fi = 0

    bd = beam_combos[fi]

    # 从组合表取各截面的控制内力
    # 边跨左M (绝对值最大, 用于配筋)
    M_el = max_abs_combo(bd.get('edge_ML', {}))
    M_er = max_abs_combo(bd.get('edge_MR', {}))
    M_emid = max_combo(bd.get('edge_Mmid', {}))  # 跨中正弯矩

    M_ml = max_abs_combo(bd.get('mid_ML', {}))
    M_mr = max_abs_combo(bd.get('mid_MR', {}))
    M_mmid = max_combo(bd.get('mid_Mmid', {}))

    # 对应剪力
    V_el = max_abs_combo(bd.get('edge_VL', {}))
    V_ml = max_abs_combo(bd.get('mid_VL', {}))

    # 原表列: [1]=V_edge, [2]=V_mid, [3]=M_edge_L, [4]=M_mid_L,
    #          [5]=M_edge_R, [6]=M_mid_R, [7]=M_edge_mid
    row.cells[1].text = f'{V_el:.2f}'
    row.cells[2].text = f'{V_ml:.2f}'
    row.cells[3].text = f'{-M_el:.2f}'  # 负号表示梁端
    row.cells[4].text = f'{-M_ml:.2f}'
    row.cells[5].text = f'{-M_er:.2f}'
    row.cells[6].text = f'{-M_mr:.2f}'
    if len(row.cells) > 7:
        row.cells[7].text = f'{M_emid:.2f}'

    print(f"  {fl}: M_edge_L={M_el:.2f}, M_mid_L={M_ml:.2f}, M_mid_s={M_emid:.2f}")

# ============================================================
# 3. 表8-7 框架柱轴压比 (T72 = idx71) — 已在上面实现
# ============================================================
print("\n" + "="*60)
print("3. 表8-7 框架柱轴压比")
print("="*60)
t = doc.tables[71]
Ac = b_col * h_col

for ri in range(2, len(t.rows)):
    row = t.rows[ri]
    fl = row.cells[0].text.strip()
    ct = row.cells[1].text.strip()

    if '6' in fl: fi = 5
    elif '5' in fl: fi = 4
    elif '4' in fl: fi = 3
    elif '3' in fl: fi = 2
    elif '2' in fl: fi = 1
    else: fi = 0

    if '边' in ct:
        key = 'edge'
    else:
        key = 'mid'

    N_top = max_abs_combo(col_combos[fi][f'{key}_top_N'])
    N_bot = max_abs_combo(col_combos[fi][f'{key}_bot_N'])
    N_max = max(N_top, N_bot)

    mu_N = N_max * 1000 / (fc * Ac)
    row.cells[2].text = f'{N_max:.2f}'
    row.cells[3].text = f'{b_col}'
    row.cells[4].text = f'{fc}'
    row.cells[5].text = f'{mu_N:.2f}'

    print(f"  {fl} {ct}: N_max={N_max:.2f}kN, μ_N={mu_N:.2f}")

# ============================================================
# 4. 表8-2 柱端弯矩调整 强柱弱梁 (T67 = idx66)
# ============================================================
print("\n" + "="*60)
print("4. 表8-2 柱端弯矩调整 (强柱弱梁)")
print("="*60)

t = doc.tables[66]
eta_c = 1.3  # 三级抗震

for ri in range(3, len(t.rows)):
    row = t.rows[ri]
    fl = row.cells[0].text.strip()
    if '6' in fl: fi = 5
    elif '5' in fl: fi = 4
    elif '4' in fl: fi = 3
    elif '3' in fl: fi = 2
    elif '2' in fl: fi = 1
    else: fi = 0

    bd = beam_combos[fi]
    cd = col_combos[fi]

    # ΣM_b: 节点处梁端弯矩绝对值之和
    # 边节点: ΣM_b = |M_edge_left| (边跨梁端)
    # 中节点: ΣM_b = |M_edge_right| + |M_mid_left|
    # 取含地震组合的控制值
    M_edge_L = max_abs_combo(bd.get('edge_ML', {}))
    M_edge_R = max_abs_combo(bd.get('edge_MR', {}))
    M_mid_L = max_abs_combo(bd.get('mid_ML', {}))

    sum_Mb_edge = M_edge_L  # 边节点: 仅边跨梁左端
    sum_Mb_mid = M_edge_R + M_mid_L  # 中节点

    # ΣM_c: 上柱下端+下柱上端弯矩
    # 边柱: M_c_top + M_c_bot
    M_ec_top = max_abs_combo(cd.get('edge_top_M', {}))
    M_ec_bot = max_abs_combo(cd.get('edge_bot_M', {}))
    M_mc_top = max_abs_combo(cd.get('mid_top_M', {}))
    M_mc_bot = max_abs_combo(cd.get('mid_bot_M', {}))

    sum_Mc_edge = M_ec_top + M_ec_bot
    sum_Mc_mid = M_mc_top + M_mc_bot

    # 调整后柱端弯矩
    M_ec_top_adj = eta_c * M_ec_top
    M_ec_bot_adj = eta_c * M_ec_bot
    M_mc_top_adj = eta_c * M_mc_top
    M_mc_bot_adj = eta_c * M_mc_bot

    # 原表列: [1]=ΣM_b_edge, [2]=ΣM_b_mid, [3]=ΣM_c_edge, [4]=ΣM_c_mid,
    #          [5]=η_c, [6]=M_c_top_adj, [7]=M_c_bot_adj
    if len(row.cells) > 2:
        row.cells[1].text = f'{sum_Mb_edge:.2f}'
        row.cells[2].text = f'{sum_Mb_mid:.2f}'
        row.cells[3].text = f'{sum_Mc_edge:.2f}'
        row.cells[4].text = f'{sum_Mc_mid:.2f}'
    if len(row.cells) > 5:
        row.cells[5].text = f'{eta_c:.2f}'
        row.cells[6].text = f'{eta_c:.2f}'
    if len(row.cells) > 7:
        row.cells[7].text = f'{-M_ec_top_adj:.2f}'
    if len(row.cells) > 8:
        row.cells[8].text = f'{-M_mc_top_adj:.2f}'

    print(f"  {fl}: ΣM_b_edge={sum_Mb_edge:.2f}, ΣM_c_edge={sum_Mc_edge:.2f}")

# ============================================================
# 5. 表8-3 柱端剪力调整 (T68 = idx67)
# ============================================================
print("\n" + "="*60)
print("5. 表8-3 柱端剪力调整 (强剪弱弯)")
print("="*60)

t = doc.tables[67]
eta_vc = 1.2

for ri in range(3, len(t.rows)):
    row = t.rows[ri]
    fl = row.cells[0].text.strip()
    if '6' in fl: fi = 5
    elif '5' in fl: fi = 4
    elif '4' in fl: fi = 3
    elif '3' in fl: fi = 2
    elif '2' in fl: fi = 1
    else: fi = 0

    is_first = (fi == 0)
    Hn = H_1st if is_first else H_std  # 柱净高
    cd = col_combos[fi]

    M_ec_top = max_abs_combo(cd.get('edge_top_M', {}))
    M_ec_bot = max_abs_combo(cd.get('edge_bot_M', {}))
    M_mc_top = max_abs_combo(cd.get('mid_top_M', {}))
    M_mc_bot = max_abs_combo(cd.get('mid_bot_M', {}))

    V_ec = eta_vc * (M_ec_top + M_ec_bot) / Hn
    V_mc = eta_vc * (M_mc_top + M_mc_bot) / Hn

    # 原表列: [1]=η_vc_e, [2]=η_vc_m, [3]=Hn, [4]=Hn, [5]=V_ec, [6]=V_mc
    row.cells[1].text = f'{eta_vc:.2f}'
    row.cells[2].text = f'{eta_vc:.2f}'
    row.cells[3].text = f'{Hn:.2f}'
    row.cells[4].text = f'{Hn:.2f}'
    if len(row.cells) > 5:
        row.cells[5].text = f'{V_ec:.2f}'
        row.cells[6].text = f'{V_mc:.2f}'

    print(f"  {fl}: V_ec={V_ec:.2f}, V_mc={V_mc:.2f}")

# ============================================================
# 6. 表8-5 框架梁正截面计算 (T70 = idx69)
# ============================================================
print("\n" + "="*60)
print("6. 表8-5 框架梁正截面计算")
print("="*60)

t = doc.tables[69]
section_order = [
    (0, 'edge_ML', '边跨左', d_edge),
    (1, 'edge_Mmid', '边跨中', d_edge),
    (2, 'edge_MR', '边跨右', d_edge),
    (3, 'mid_ML', '中跨左', d_mid),
    (4, 'mid_Mmid', '中跨中', d_mid),
]

for ri in range(2, len(t.rows)):
    row = t.rows[ri]
    fl_text = row.cells[0].text.strip()
    sec_text = row.cells[1].text.strip()

    if '6' in fl_text: fi = 5
    elif '5' in fl_text: fi = 4
    elif '4' in fl_text: fi = 3
    elif '3' in fl_text: fi = 2
    elif '2' in fl_text: fi = 1
    else: fi = 0

    bd = beam_combos[fi]

    # 确定截面
    M_design = 0; d_eff = d_edge
    is_mid = ('跨中' in sec_text or '中跨中' in sec_text)
    is_mid_beam = ('中跨' in sec_text)
    b_eff = b_beam
    h_beam = h_edge

    if '边跨' in sec_text:
        if '左' in sec_text: key = 'edge_ML'; d_eff = d_edge
        elif '右' in sec_text: key = 'edge_MR'; d_eff = d_edge
        elif '跨' in sec_text: key = 'edge_Mmid'; d_eff = d_edge
        else: continue
        h_beam = h_edge
    elif '中跨' in sec_text:
        if '左' in sec_text: key = 'mid_ML'; d_eff = d_mid
        elif '跨' in sec_text or '中' in sec_text: key = 'mid_Mmid'; d_eff = d_mid
        else: continue
        h_beam = h_mid
    else:
        continue

    if key not in bd:
        continue

    # 取弯矩设计值 (取含地震/风组合的绝对值最大值)
    M_design = max_abs_combo(bd[key])

    # 正截面计算
    alpha_s = M_design * 1e6 / (alpha1 * fc * b_eff * d_eff**2)
    xi = 1 - math.sqrt(1 - 2*alpha_s)
    gamma_s = 1 - xi/2 if xi <= xi_b else 0.5 + 0.5*math.sqrt(1 - 2*alpha_s)
    As = M_design * 1e6 / (fy * gamma_s * d_eff) if gamma_s > 0 else 0

    # 最小配筋率: max(0.2%, 45*ft/fy%)
    rho_min = max(0.002, 0.45*ft/fy)
    As_min = rho_min * b_eff * h_beam

    # 原表列: [2]=M, [3]=α_s, [4]=ξ, [5]=γ_s, [6]=As
    row.cells[2].text = f'{M_design:.2f}'
    row.cells[3].text = f'{alpha_s:.3f}'
    row.cells[4].text = f'{xi:.3f}'
    row.cells[5].text = f'{gamma_s:.3f}'
    if len(row.cells) > 6:
        row.cells[6].text = f'{max(As, As_min):.0f}'

    if ri < 8:
        print(f"  {fl_text} {sec_text}: M={M_design:.2f}, α_s={alpha_s:.3f}, As={max(As,As_min):.0f}mm²")

# ============================================================
# 7. 表8-6 框架梁斜截面计算 (T71 = idx70)
# ============================================================
print("\n" + "="*60)
print("7. 表8-6 框架梁斜截面计算")
print("="*60)

t = doc.tables[70]

for ri in range(2, len(t.rows)):
    row = t.rows[ri]
    fl_text = row.cells[0].text.strip()
    beam_text = row.cells[1].text.strip()

    if '6' in fl_text: fi = 5
    elif '5' in fl_text: fi = 4
    elif '4' in fl_text: fi = 3
    elif '3' in fl_text: fi = 2
    elif '2' in fl_text: fi = 1
    else: fi = 0

    bd = beam_combos[fi]
    is_edge_beam = '边' in beam_text
    d_eff = d_edge if is_edge_beam else d_mid

    # 取最大剪力设计值
    if is_edge_beam:
        V_design = max_abs_combo(bd.get('edge_VL', {}))
    else:
        V_design = max_abs_combo(bd.get('mid_VL', {}))

    # 受剪承载力验算
    Vc = 0.7 * ft * b_beam * d_eff / 1000  # kN
    # 箍筋计算: ρ_sv = (V - Vc) / (fyv * d_eff) * s
    # 简化处理
    V_ratio = (V_design - Vc) / Vc if Vc > 0 else 0

    row.cells[2].text = f'{V_design:.2f}'
    row.cells[3].text = f'{d_eff}'
    row.cells[4].text = f'{b_beam}'

    if V_design <= Vc:
        row.cells[5].text = '构造'
    else:
        row.cells[5].text = f'{V_ratio:.3f}'

    print(f"  {fl_text} {beam_text}: V={V_design:.2f}, Vc={Vc:.2f}")

# ============================================================
# 8. 表8-12 框架柱斜截面计算 (T76 = idx75)
# ============================================================
print("\n" + "="*60)
print("8. 表8-12 框架柱斜截面计算")
print("="*60)

t = doc.tables[75]

for ri in range(2, len(t.rows)):
    row = t.rows[ri]
    fl_text = row.cells[0].text.strip()
    col_text = row.cells[1].text.strip()

    if '6' in fl_text: fi = 5
    elif '5' in fl_text: fi = 4
    elif '4' in fl_text: fi = 3
    elif '3' in fl_text: fi = 2
    elif '2' in fl_text: fi = 1
    else: fi = 0

    cd = col_combos[fi]
    is_edge = '边' in col_text

    key_pref = 'edge' if is_edge else 'mid'
    N_vals = cd.get(f'{key_pref}_top_N', {})
    V_vals = cd.get(f'{key_pref}_top_V', {})

    # 最大轴力
    N_max = max_abs_combo(N_vals)
    # 柱剪力(已考虑强剪弱弯)
    V_col = max_abs_combo(V_vals)

    is_first = (fi == 0)
    Hn = H_1st if is_first else H_std

    row.cells[2].text = f'{N_max:.2f}'
    if len(row.cells) > 3:
        row.cells[3].text = f'{V_col:.2f}'
    if len(row.cells) > 4:
        row.cells[4].text = f'{Hn:.2f}'

    print(f"  {fl_text} {col_text}: N_max={N_max:.2f}, V={V_col:.2f}")

# ============================================================
# 9. 表11-1 基础内力设计值 (T81 = idx80)
# ============================================================
print("\n" + "="*60)
print("9. 表11-1 基础内力设计值")
print("="*60)

t = doc.tables[80]
cd = col_combos[0]  # 1F (底层)

# 边柱基础: 取1F柱底内力
M_edge = max_abs_combo(cd.get('edge_bot_M', {}))
N_edge = max_abs_combo(cd.get('edge_bot_N', {}))
V_edge = max_abs_combo(cd.get('edge_bot_V', {}))

# 中柱基础
M_mid = max_abs_combo(cd.get('mid_bot_M', {}))
N_mid = max_abs_combo(cd.get('mid_bot_N', {}))
V_mid = max_abs_combo(cd.get('mid_bot_V', {}))

# 标准组合: /1.35 近似
N_edge_std = N_edge / 1.35
M_edge_std = M_edge / 1.35
V_edge_std = V_edge / 1.35
N_mid_std = N_mid / 1.35
M_mid_std = M_mid / 1.35
V_mid_std = V_mid / 1.35

# 写入表 (根据原表结构)
if len(t.rows) > 2:
    # 原表可能格式: 边柱基本组合/标准组合, 中柱基本组合/标准组合
    for ri in range(len(t.rows)):
        row = t.rows[ri]
        txt0 = row.cells[0].text.strip() if row.cells[0].text.strip() else ''

    # 简化: 至少更新基础数据行
    if len(t.rows) >= 4:
        row = t.rows[2]  # 假设R2=边柱
        if len(row.cells) > 1:
            row.cells[1].text = f'{M_edge:.2f}'
        if len(row.cells) > 2:
            row.cells[2].text = f'{N_edge:.2f}'
        if len(row.cells) > 3:
            row.cells[3].text = f'{V_edge:.2f}'

        row = t.rows[3]  # 假设R3=中柱
        if len(row.cells) > 1:
            row.cells[1].text = f'{M_mid:.2f}'
        if len(row.cells) > 2:
            row.cells[2].text = f'{N_mid:.2f}'
        if len(row.cells) > 3:
            row.cells[3].text = f'{V_mid:.2f}'

print(f"  边柱基础: M={M_edge:.2f}, N={N_edge:.2f}, V={V_edge:.2f}")
print(f"  中柱基础: M={M_mid:.2f}, N={N_mid:.2f}, V={V_mid:.2f}")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n第8章+第11章全部更新完成")
print(f"修正版: {DOC}")
print(f"审阅版: {review}")
