#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
修复剩余表格: 表11-1(列映射), 表8-8(控制内力), 表8-9(轴力最大配筋), 表8-11(弯矩最大配筋)
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
BAK = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_4800原版备份.docx'

doc = Document(DOC)
bak = Document(BAK)

# ============================================================
# 材料与截面参数
# ============================================================
fc, ft = 14.3, 1.43
fy, fyv = 360, 360
alpha1, beta1 = 1.0, 0.8
xi_b = 0.518
b_col, h_col = 500, 500
h0_col = 465  # h - as = 500 - 35
as_col = 35
Ac = b_col * h_col

L1, L2 = 5.4, 2.4
H_std, H_1st = 3.0, 4.0

# ============================================================
# 读取柱组合表 (7-15~7-20 = tables[59]~[64])
# ============================================================
# 每表结构: R4-R6边柱上(M/N/V), R7-R9边柱下, R10-R12中柱上, R13-R15中柱下
# cells[3]=标签, [4]=D, [5]=L, [6]=W, [7]=E, [8-16]=c1-c9, [17]=c10, [18]=c11, [19]=c12

col_data = []  # col_data[fi] = {...}
FLOOR_LABELS = ['1F','2F','3F','4F','5F','6F']

for ti in range(6):
    tbl_idx = 59 + ti
    fi = 5 - ti  # ti=0→1F(fi=5), ti=5→6F(fi=0)
    tbl = doc.tables[tbl_idx]
    fn = FLOOR_LABELS[fi]

    def read_section(ri_m, ri_n, ri_v):
        """读取一个截面的M/N/V和所有组合值"""
        # 逐个读取组合值
        m_combos = []
        n_combos = []
        v_combos = []
        for ci in range(8, 17):
            try:
                m_combos.append(float(tbl.rows[ri_m].cells[ci].text.strip()))
            except: m_combos.append(0.0)
            try:
                n_combos.append(float(tbl.rows[ri_n].cells[ci].text.strip()))
            except: n_combos.append(0.0)
            try:
                v_combos.append(float(tbl.rows[ri_v].cells[ci].text.strip()))
            except: v_combos.append(0.0)

        # D, L, W, E
        try: D_M = float(tbl.rows[ri_m].cells[4].text.strip())
        except: D_M = 0.0
        try: L_M = float(tbl.rows[ri_m].cells[5].text.strip())
        except: L_M = 0.0
        try: D_N = float(tbl.rows[ri_n].cells[4].text.strip())
        except: D_N = 0.0
        try: L_N = float(tbl.rows[ri_n].cells[5].text.strip())
        except: L_N = 0.0
        try: D_V = float(tbl.rows[ri_v].cells[4].text.strip())
        except: D_V = 0.0
        try: L_V = float(tbl.rows[ri_v].cells[5].text.strip())
        except: L_V = 0.0

        return {
            'M': m_combos, 'N': n_combos, 'V': v_combos,
            'D_M': D_M, 'L_M': L_M, 'D_N': D_N, 'L_N': L_N,
            'D_V': D_V, 'L_V': L_V,
        }

    edge_top = read_section(4, 5, 6)
    edge_bot = read_section(7, 8, 9)
    mid_top = read_section(10, 11, 12)
    mid_bot = read_section(13, 14, 15)

    col_data.append({
        'floor': fn,
        'fi': fi,
        'edge_top': edge_top,
        'edge_bot': edge_bot,
        'mid_top': mid_top,
        'mid_bot': mid_bot,
    })

# ============================================================
# 1. 修复 表11-1 (T80 = idx80)
# ============================================================
print("="*60)
print("1. 修复 表11-1 基础内力设计值")
print("="*60)

t = doc.tables[80]
t_bak = bak.tables[80]
cd = col_data[0]  # fi=0 → 1F

# 标准组合: D+L
def std_combo(sec):
    return sec['D_M']+sec['L_M'], sec['D_N']+sec['L_N'], sec['D_V']+sec['L_V']

# 设计组合: c10 = max(abs(c1..c9))
def design_combo(sec):
    M = max(abs(v) for v in sec['M'])
    N = max(abs(v) for v in sec['N'])
    V = max(abs(v) for v in sec['V'])
    return M, N, V

# 边柱
M_std_e, N_std_e, V_std_e = std_combo(cd['edge_bot'])
M_dsn_e, N_dsn_e, V_dsn_e = design_combo(cd['edge_bot'])

# 中柱
M_std_m, N_std_m, V_std_m = std_combo(cd['mid_bot'])
M_dsn_m, N_dsn_m, V_dsn_m = design_combo(cd['mid_bot'])

# 4800备份的结构:
# R1: ['内力', '柱', '', '', '']
# R2: 标准值 边柱 → [0]=标准值, [1]=边柱, [2]=N, [3]=M, [4]=V
# R3: 标准值 中柱
# R4: 设计值 边柱
# R5: 设计值 中柱

# 先恢复标签
for ri in range(1, min(6, len(t.rows))):
    for ci in range(min(5, len(t.rows[ri].cells))):
        if len(t_bak.rows[ri].cells) > ci:
            bak_text = t_bak.rows[ri].cells[ci].text.strip()
            # 恢复标签列(ci=0,1)
            if ci <= 1:
                t.rows[ri].cells[ci].text = bak_text

# 写入标准值 R2(边柱), R3(中柱)
# cells[1]=边柱/中柱, cells[2]=N, cells[3]=M, cells[4]=V
t.rows[2].cells[2].text = f'{N_std_e:.2f}'
t.rows[2].cells[3].text = f'{M_std_e:.2f}'
t.rows[2].cells[4].text = f'{V_std_e:.2f}'
t.rows[3].cells[2].text = f'{N_std_m:.2f}'
t.rows[3].cells[3].text = f'{M_std_m:.2f}'
t.rows[3].cells[4].text = f'{V_std_m:.2f}'

# 写入设计值 R4(边柱), R5(中柱)
t.rows[4].cells[2].text = f'{N_dsn_e:.2f}'
t.rows[4].cells[3].text = f'{M_dsn_e:.2f}'
t.rows[4].cells[4].text = f'{V_dsn_e:.2f}'
t.rows[5].cells[2].text = f'{N_dsn_m:.2f}'
t.rows[5].cells[3].text = f'{M_dsn_m:.2f}'
t.rows[5].cells[4].text = f'{V_dsn_m:.2f}'

print(f"  边柱 标准值: N={N_std_e:.2f}, M={M_std_e:.2f}, V={V_std_e:.2f}")
print(f"  边柱 设计值: N={N_dsn_e:.2f}, M={M_dsn_e:.2f}, V={V_dsn_e:.2f}")
print(f"  中柱 标准值: N={N_std_m:.2f}, M={M_std_m:.2f}, V={V_std_m:.2f}")
print(f"  中柱 设计值: N={N_dsn_m:.2f}, M={M_dsn_m:.2f}, V={V_dsn_m:.2f}")

# ============================================================
# 2. 更新 表8-8 边柱（三层）控制内力 (T72 = idx72)
# ============================================================
print("\n" + "="*60)
print("2. 更新 表8-8 边柱(三层)控制内力")
print("="*60)

t = doc.tables[72]
cd3 = col_data[2]  # fi=2 → 3F

def select_controlling(sec_top, sec_bot, combo_indices, criterion):
    """
    从指定组合中选取控制值.
    criterion: 'N_max', 'N_min', 'M_max'
    返回 (N, M) — 从控制截面(顶或底)选取
    """
    best_N = None
    best_M = None
    best_val = None
    best_section = None

    for sec, sec_name in [(sec_top, 'top'), (sec_bot, 'bot')]:
        for idx in combo_indices:
            N = sec['N'][idx]
            M = sec['M'][idx]
            if criterion == 'N_max':
                val = N
            elif criterion == 'N_min':
                val = N
            else:  # M_max
                val = abs(M)

            if best_val is None:
                best_val = val
                best_N = N
                best_M = M
                best_section = sec_name
            elif criterion == 'N_max' and val > best_val:
                best_val = val; best_N = N; best_M = M; best_section = sec_name
            elif criterion == 'N_min' and val < best_val:
                best_val = val; best_N = N; best_M = M; best_section = sec_name
            elif criterion == 'M_max' and val > best_val:
                best_val = val; best_N = N; best_M = M; best_section = sec_name

    return best_N, best_M, best_section

# 抗震组合: c2,c3,c6,c7 (indices 1,2,5,6)
seismic_idx = [1, 2, 5, 6]
# 非抗震组合: c1,c4,c5 (indices 0,3,4)
nonseismic_idx = [0, 3, 4]

# 抗震组合: N_max, N_min, |M|_max
s_Nmax_N, s_Nmax_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], seismic_idx, 'N_max')
s_Nmin_N, s_Nmin_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], seismic_idx, 'N_min')
s_Mmax_N, s_Mmax_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], seismic_idx, 'M_max')

# 非抗震组合: N_max, N_min, |M|_max
ns_Nmax_N, ns_Nmax_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], nonseismic_idx, 'N_max')
ns_Nmin_N, ns_Nmin_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], nonseismic_idx, 'N_min')
ns_Mmax_N, ns_Mmax_M, _ = select_controlling(cd3['edge_top'], cd3['edge_bot'], nonseismic_idx, 'M_max')

# 写入表8-8
# R2: N值 (抗震: N_max, N_min, |M|_max)
# R3: M值 (抗震: 对应M)
# R4: N值 (非抗震: N_max, N_min, |M|_max)
# R5: M值 (非抗震: 对应M)
t.rows[2].cells[1].text = f'{s_Nmax_N:.2f}'
t.rows[2].cells[2].text = f'{s_Nmin_N:.2f}'
t.rows[2].cells[3].text = f'{s_Mmax_N:.2f}'
t.rows[3].cells[1].text = f'{s_Nmax_M:.2f}'
t.rows[3].cells[2].text = f'{s_Nmin_M:.2f}'
t.rows[3].cells[3].text = f'{s_Mmax_M:.2f}'

t.rows[4].cells[1].text = f'{ns_Nmax_N:.2f}'
t.rows[4].cells[2].text = f'{ns_Nmin_N:.2f}'
t.rows[4].cells[3].text = f'{ns_Mmax_N:.2f}'
t.rows[5].cells[1].text = f'{ns_Nmax_M:.2f}'
t.rows[5].cells[2].text = f'{ns_Nmin_M:.2f}'
t.rows[5].cells[3].text = f'{ns_Mmax_M:.2f}'

print(f"  抗震 N_max: N={s_Nmax_N:.2f}, M={s_Nmax_M:.2f}")
print(f"  抗震 N_min: N={s_Nmin_N:.2f}, M={s_Nmin_M:.2f}")
print(f"  抗震 |M|_max: N={s_Mmax_N:.2f}, M={s_Mmax_M:.2f}")
print(f"  非抗震 N_max: N={ns_Nmax_N:.2f}, M={ns_Nmax_M:.2f}")
print(f"  非抗震 N_min: N={ns_Nmin_N:.2f}, M={ns_Nmin_M:.2f}")
print(f"  非抗震 |M|_max: N={ns_Mmax_N:.2f}, M={ns_Mmax_M:.2f}")

# ============================================================
# 3. 更新 表8-9 (轴力最大组合配筋) + 表8-11 (弯矩最大组合配筋)
# ============================================================
print("\n" + "="*60)
print("3. 更新 表8-9 和 表8-11 柱配筋计算")
print("="*60)

# 组合名称: c1到c9
COMBO_NAMES = ['1.3D+1.5L','1.2D+1.4W+0.98L','1.2D+0.6L+1.3E',
               '1.0D+1.4W','1.0D+0.5L+1.3E',
               '1.2D-1.4W+0.98L','1.2D+0.6L-1.3E',
               '1.0D-1.4W','1.0D+0.5L-1.3E']

def calc_column_reinforcement(N_val, M_top, M_bot, is_top, floor_label):
    """
    GB50010-2010 偏心受压柱配筋计算.
    N_val: 轴力设计值 (kN)
    M_top, M_bot: 柱端弯矩 (kN·m), 来自同一组合
    is_top: True=设计上端截面
    floor_label: '1F'~'6F'
    """
    if is_top:
        M2 = abs(M_top)
        M1 = abs(M_bot)
    else:
        M2 = abs(M_bot)
        M1 = abs(M_top)

    # 处理零弯矩情况
    if M2 < 0.01:
        M2 = 0.01
    if M1 < 0.01:
        M1 = 0.01 * M2  # 制造一个很小的M1/M2比值

    # M1/M2 符号: 同向弯曲为正
    M1_M2_ratio = M1 / M2 if M2 > 0 else 1.0
    # 确保 M1 ≤ M2 (M1是较小端, M2是较大端)
    if M1_M2_ratio > 1.0:
        M1_M2_ratio = 1.0 / M1_M2_ratio  # 交换

    # C_m = 0.7 + 0.3*(M1/M2) ≥ 0.7
    # M1/M2 sign: 同曲率取正(单曲率), 异曲率取负(双曲率)
    # 柱端弯矩同号→单曲率→M1/M2>0; 异号→双曲率→M1/M2<0
    if is_top:
        sign = 1 if M_top * M_bot >= 0 else -1
    else:
        sign = 1 if M_bot * M_top >= 0 else -1

    C_m = 0.7 + 0.3 * sign * M1_M2_ratio
    C_m = max(C_m, 0.7)

    # 柱计算长度
    is_first = ('1F' in floor_label or '1' == floor_label[0])
    if is_first:
        lc = 4.0 * 1000  # mm
    else:
        lc = 1.25 * 3.0 * 1000  # = 3750 mm

    # ζ_c = 0.5*fc*A/N ≤ 1.0 (考虑轴力对曲率的影响)
    if N_val > 1:
        zeta_c = min(0.5 * fc * Ac / (N_val * 1000), 1.0)
    else:
        zeta_c = 1.0

    # ea = max(20, h/30)
    ea = max(20, h_col / 30)

    # e0 = M2/N (初始偏心距, mm)
    if N_val > 0.01:
        e0 = M2 * 1000 / N_val  # mm
    else:
        e0 = 0

    # η_ns 弯矩放大系数
    # η_ns = 1 + (lc/h)² * ζ_c / (1300 * (M2/N + ea) / h0)
    denom = 1300 * (e0 + ea) / h0_col
    if denom > 0.01:
        eta_ns = 1 + (lc / h_col) ** 2 * zeta_c / denom
    else:
        eta_ns = 1.0

    # C_m * η_ns ≥ 1.0
    factor = max(C_m * eta_ns, 1.0)

    # 设计弯矩 M = C_m * η_ns * M2
    M_design = factor * M2

    # 偏心距计算
    e0_design = M_design * 1000 / N_val if N_val > 0.01 else 0  # mm
    ei = e0_design + ea  # mm
    e = ei + h_col / 2 - as_col  # mm

    # 受压区高度
    if N_val > 0:
        x = N_val * 1000 / (alpha1 * fc * b_col)  # mm
    else:
        x = 0

    # 判断大小偏压
    xi = x / h0_col if h0_col > 0 else 0
    if xi <= xi_b:
        judge = '大偏压'
    else:
        judge = '小偏压'

    # 配筋计算
    if judge == '大偏压':
        # x < 2as' 时取 x = 2as'
        if x < 2 * as_col:
            x = 2 * as_col
            As = N_val * 1000 * (ei - h_col/2 + as_col) / (fy * (h0_col - as_col)) if h0_col > as_col else 0
        else:
            As = (N_val * 1000 * e - alpha1 * fc * b_col * x * (h0_col - x/2)) / (fy * (h0_col - as_col))
            if As < 0:
                As = 0
    else:
        # 小偏压简化计算
        xi_calc = (N_val * 1000 - xi_b * alpha1 * fc * b_col * h0_col) / \
                  ((N_val * 1000 * e - 0.43 * alpha1 * fc * b_col * h0_col**2) / ((beta1 - xi_b) * (h0_col - as_col)) + alpha1 * fc * b_col * h0_col) + xi_b
        xi_calc = max(xi_calc, xi_b)
        As = (N_val * 1000 * e - xi_calc * (1 - 0.5 * xi_calc) * alpha1 * fc * b_col * h0_col**2) / (fy * (h0_col - as_col))
        if As < 0:
            As = 0

    # 最小配筋率: max(0.2%, 0.45*ft/fy%)
    rho_min = max(0.002, 0.45 * ft / fy)
    As_min = rho_min * b_col * h0_col / 2  # 单边最小

    As = max(As, As_min)

    # 选筋: 柱纵筋最小直径12mm (GB50010 §9.3.1)
    bar_options = {
        452: '4C12(452mm²)', 509: '2C18(509mm²)', 603: '3C16(603mm²)',
        628: '2C20(628mm²)', 763: '3C18(763mm²)', 804: '4C16(804mm²)',
        942: '3C20(942mm²)', 1017: '4C18(1018mm²)', 1140: '3C22(1140mm²)',
        1256: '4C20(1256mm²)', 1520: '3C25(1473mm²)', 1964: '4C25(1964mm²)',
    }
    bar_label = '4C16(804mm²)'
    As_selected = 804
    for area, label in sorted(bar_options.items()):
        if area >= As:
            bar_label = label
            As_selected = area
            break
    else:
        As_selected = As
        bar_label = f'{As:.0f}mm²'

    # ρ = As / (b*h0) * 100% (单边配筋率)
    rho = As_selected / (b_col * h0_col) * 100 if h0_col > 0 else 0

    return {
        'M2': M2, 'N': N_val,
        'C_m': C_m, 'M1_M2': sign * M1_M2_ratio,
        'eta_ns': eta_ns, 'M_design': M_design,
        'ei': ei, 'e': e, 'x': x,
        'judge': judge, 'As': As,
        'bar_label': bar_label, 'rho': rho,
        'lc': lc, 'zeta_c': zeta_c, 'ea': ea,
    }

def get_combo_for_section(sec, criterion):
    """从单个截面的组合中找到控制组合索引和值"""
    best_val = None
    best_idx = -1

    for idx in range(9):
        if criterion == 'N_max':
            val = abs(sec['N'][idx])
        elif criterion == 'N_min':
            val = sec['N'][idx]  # 代数值最小
        else:  # M_max
            val = abs(sec['M'][idx])

        if best_val is None:
            best_val = val; best_idx = idx
        elif criterion == 'N_min' and val < best_val:
            best_val = val; best_idx = idx
        elif criterion != 'N_min' and val > best_val:
            best_val = val; best_idx = idx

    return best_idx, sec['N'][best_idx], sec['M'][best_idx]

def design_section(sec, sec_opposite, is_top, criterion, fn):
    """设计单个截面: 找到控制组合, 取对端弯矩, 计算配筋"""
    idx, N_val, M_this = get_combo_for_section(sec, criterion)
    M_opposite = sec_opposite['M'][idx]  # 同一组合对端弯矩
    if is_top:
        return calc_column_reinforcement(N_val, M_this, M_opposite, True, fn)
    else:
        return calc_column_reinforcement(N_val, M_opposite, M_this, False, fn)

def write_section_rows(tbl, ri, fn, col_type, pos, result):
    """写入截面配筋结果的两行(M行+N行)"""
    if ri + 1 >= len(tbl.rows):
        return
    row_m = tbl.rows[ri]
    row_n = tbl.rows[ri+1]
    for row, label in [(row_m, 'M'), (row_n, 'N')]:
        row.cells[0].text = fn
        row.cells[1].text = f'{col_type}{pos}'
        row.cells[2].text = label
        row.cells[3].text = f'{result["M2"] if label=="M" else result["N"]:.2f}'
        row.cells[4].text = f'{result["C_m"]:.2f}'
        row.cells[5].text = f'{result["M1_M2"]:.3f}'
        row.cells[6].text = f'{result["eta_ns"]:.3f}'
        row.cells[7].text = f'{result["M_design"]:.2f}'
        row.cells[8].text = f'{result["ei"]:.1f}'
        row.cells[9].text = f'{result["e"]:.1f}'
        row.cells[10].text = f'{result["x"]:.1f}'
        row.cells[11].text = result['judge']
        row.cells[12].text = f'{result["As"]:.0f}'
        row.cells[15].text = result['bar_label']
        row.cells[16].text = f'{result["rho"]:.2f}'

# 更新表8-9 (轴力最大)
print("\n--- 表8-9 框架柱(轴力最大组合)配筋 ---")
t89 = doc.tables[73]
ri = 2
for fi in range(6):
    cd_f = col_data[fi]
    fn = FLOOR_LABELS[fi]
    for col_type, sec_top, sec_bot in [
        ('边柱', cd_f['edge_top'], cd_f['edge_bot']),
        ('中柱', cd_f['mid_top'], cd_f['mid_bot']),
    ]:
        # 顶截面: 在顶部找N_max组合
        r_top = design_section(sec_top, sec_bot, True, 'N_max', fn)
        write_section_rows(t89, ri, fn, col_type, '顶', r_top)
        ri += 2

        # 底截面: 在底部找N_max组合
        r_bot = design_section(sec_bot, sec_top, False, 'N_max', fn)
        write_section_rows(t89, ri, fn, col_type, '底', r_bot)
        ri += 2

        if fi == 0 and col_type == '边柱':
            print(f"  {fn} {col_type}顶: N={r_top['N']:.2f}, M2={r_top['M2']:.2f}, η_ns={r_top['eta_ns']:.3f}, {r_top['judge']}, As={r_top['As']:.0f}")
            print(f"  {fn} {col_type}底: N={r_bot['N']:.2f}, M2={r_bot['M2']:.2f}, η_ns={r_bot['eta_ns']:.3f}, {r_bot['judge']}, As={r_bot['As']:.0f}")

# 更新表8-11 (弯矩最大)
print("\n--- 表8-11 框架柱(弯矩绝对值最大组合)配筋 ---")
t811 = doc.tables[74]
ri = 2
for fi in range(6):
    cd_f = col_data[fi]
    fn = FLOOR_LABELS[fi]
    for col_type, sec_top, sec_bot in [
        ('边柱', cd_f['edge_top'], cd_f['edge_bot']),
        ('中柱', cd_f['mid_top'], cd_f['mid_bot']),
    ]:
        # 顶截面: 在顶部找|M|_max组合
        r_top = design_section(sec_top, sec_bot, True, 'M_max', fn)
        write_section_rows(t811, ri, fn, col_type, '顶', r_top)
        ri += 2

        # 底截面: 在底部找|M|_max组合
        r_bot = design_section(sec_bot, sec_top, False, 'M_max', fn)
        write_section_rows(t811, ri, fn, col_type, '底', r_bot)
        ri += 2

        if fi == 0 and col_type == '边柱':
            print(f"  {fn} {col_type}顶: N={r_top['N']:.2f}, M2={r_top['M2']:.2f}, η_ns={r_top['eta_ns']:.3f}, {r_top['judge']}, As={r_top['As']:.0f}")
            print(f"  {fn} {col_type}底: N={r_bot['N']:.2f}, M2={r_bot['M2']:.2f}, η_ns={r_bot['eta_ns']:.3f}, {r_bot['judge']}, As={r_bot['As']:.0f}")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n全部修复完成!")
print(f"修正版: {DOC}")
print(f"审阅版: {review}")
