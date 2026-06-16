#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于完整弯矩二次分配结果, 更新docx所有Ch6-Ch8表格
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF,0x00,0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

# ============================================================
# 参数
# ============================================================
L1, L2 = 5.4, 2.4; ss = 3.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0
F_NEW = 1-2*(0.5*ss/L1)**2+(0.5*ss/L1)**3
E = 30e6; Ic = 0.5**4/12
ic_std = E*Ic/3.0; ic_1st = E*Ic/4.0
I0_e = 0.25*0.5**3/12
ib_edge = E*1.5*I0_e/L1
ib_mid = E*1.5*0.25*0.4**3/12/L2

# 等效均布
qe_d_r = 2.57 + F_NEW*ss*g_roof        # 16.75
qe_d_f = 2.57+6.2 + F_NEW*ss*g_floor   # 20.78
qm_d_r = 1.89 + 0.625*L2*g_roof        # 9.33
qm_d_f = 1.89+6.45 + 0.625*L2*g_floor  # 14.64
qe_l_r = F_NEW*ss*q_roof               # 1.43
qe_l_f = F_NEW*ss*q_floor              # 5.72
qm_l_r = 0.625*L2*q_roof               # 0.75
qm_l_f = 0.625*L2*q_floor              # 3.00

def fem(q, L): return q*L**2/12
def midM(q, L, Ml, Mr): return q*L**2/8 - (abs(Ml)+abs(Mr))/2
def beamV(q, L, Ml, Mr):
    Vl = q*L/2 + (Mr-Ml)/L
    Vr = q*L/2 + (Ml-Mr)/L
    return Vl, Vr

def moment_dist(ic_u, ic_l, ib_e, ib_m, q_e, q_m):
    FE_e = fem(q_e, L1); FE_m = fem(q_m, L2)
    sA = ic_u+ic_l+ib_e; sB = ic_u+ic_l+ib_e+ib_m
    mu_A_u=ic_u/sA; mu_A_l=ic_l/sA; mu_A_b=ib_e/sA
    mu_B_u=ic_u/sB; mu_B_l=ic_l/sB; mu_B_bl=ib_e/sB; mu_B_br=ib_m/sB

    M_A_u=M_A_l=0; M_A_b=-FE_e
    M_B_u=M_B_l=0; M_B_bl=FE_e; M_B_br=-FE_m

    # 1st distribution
    uA=M_A_u+M_A_l+M_A_b; dAb=-mu_A_b*uA; dAu=-mu_A_u*uA; dAl=-mu_A_l*uA
    M_A_b+=dAb; M_A_u+=dAu; M_A_l+=dAl
    uB=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl=-mu_B_bl*uB; dBbr=-mu_B_br*uB; dBu=-mu_B_u*uB; dBl=-mu_B_l*uB
    M_B_bl+=dBbl; M_B_br+=dBbr; M_B_u+=dBu; M_B_l+=dBl

    # 1st carry-over
    M_A_b+=0.5*dBbl; M_B_bl+=0.5*dAb; M_B_br+=0.5*(-dBbr)

    # 2nd distribution
    uA2=M_A_u+M_A_l+M_A_b; dAb2=-mu_A_b*uA2; dAu2=-mu_A_u*uA2; dAl2=-mu_A_l*uA2
    M_A_b+=dAb2; M_A_u+=dAu2; M_A_l+=dAl2
    uB2=M_B_u+M_B_l+M_B_bl+M_B_br
    dBbl2=-mu_B_bl*uB2; dBbr2=-mu_B_br*uB2; dBu2=-mu_B_u*uB2; dBl2=-mu_B_l*uB2
    M_B_bl+=dBbl2; M_B_br+=dBbr2; M_B_u+=dBu2; M_B_l+=dBl2

    # 2nd carry-over
    M_A_b+=0.5*dBbl2; M_B_bl+=0.5*dAb2; M_B_br+=0.5*(-dBbr2)

    return dict(A_u=M_A_u,A_l=M_A_l,A_b=M_A_b,
                B_u=M_B_u,B_l=M_B_l,B_bl=M_B_bl,B_br=M_B_br)

# 所有6层计算
floors = [
    ('6F',0,ic_std,qe_d_r,qm_d_r,qe_l_r,qm_l_r),
    ('5F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('4F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('3F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('2F',ic_std,ic_std,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
    ('1F',ic_std,ic_1st,qe_d_f,qm_d_f,qe_l_f,qm_l_f),
]

data = {}
for nm, icu, icl, qed, qmd, qel, qml in floors:
    rd = moment_dist(icu, icl, ib_edge, ib_mid, qed, qmd)
    rl = moment_dist(icu, icl, ib_edge, ib_mid, qel, qml)
    data[nm] = {'dead': rd, 'live': rl, 'qe_d': qed, 'qm_d': qmd, 'qe_l': qel, 'qm_l': qml}

# ============================================================
# 更新表6-6 恒载跨中弯矩 (Table 35)
# ============================================================
t = doc.tables[34]
print("更新表6-6 (恒载跨中弯矩)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; qed = data[nm]['qe_d']; qmd = data[nm]['qm_d']
    Me = midM(qed, L1, rd['A_b'], rd['B_bl'])
    Mm = midM(qmd, L2, rd['B_br'], -rd['B_br'])
    row_e = 2+ri*2; row_m = 3+ri*2
    t.rows[row_e].cells[3].text = f'{qed:.2f}'
    t.rows[row_e].cells[4].text = f'{rd["A_b"]:.2f}'
    t.rows[row_e].cells[5].text = f'{rd["B_bl"]:.2f}'
    t.rows[row_e].cells[6].text = f'{Me:.2f}'
    if row_m < len(t.rows):
        t.rows[row_m].cells[3].text = f'{qmd:.2f}'
        t.rows[row_m].cells[4].text = f'{rd["B_br"]:.2f}'
        t.rows[row_m].cells[5].text = f'{-rd["B_br"]:.2f}'
        t.rows[row_m].cells[6].text = f'{Mm:.2f}'

# ============================================================
# 更新表6-13 活载跨中弯矩 (Table 42)
# ============================================================
t = doc.tables[41]
print("更新表6-13 (活载跨中弯矩)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']; qel = data[nm]['qe_l']; qml = data[nm]['qm_l']
    Me = midM(qel, L1, rl['A_b'], rl['B_bl'])
    Mm = midM(qml, L2, rl['B_br'], -rl['B_br'])
    row_e = 2+ri*2; row_m = 3+ri*2
    t.rows[row_e].cells[3].text = f'{qel:.2f}'
    t.rows[row_e].cells[4].text = f'{rl["A_b"]:.2f}'
    t.rows[row_e].cells[5].text = f'{rl["B_bl"]:.2f}'
    t.rows[row_e].cells[6].text = f'{Me:.2f}'
    if row_m < len(t.rows):
        t.rows[row_m].cells[3].text = f'{qml:.2f}'
        t.rows[row_m].cells[4].text = f'{rl["B_br"]:.2f}'
        t.rows[row_m].cells[5].text = f'{-rl["B_br"]:.2f}'
        t.rows[row_m].cells[6].text = f'{Mm:.2f}'

# ============================================================
# 更新表6-7 恒载梁端剪力 (Table 36)
# ============================================================
t = doc.tables[35]
print("更新表6-7 (恒载梁端剪力)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']
    qed = data[nm]['qe_d']; qmd = data[nm]['qm_d']
    # 边跨
    Vl_e, Vr_e = beamV(qed, L1, rd['A_b'], rd['B_bl'])
    # 中跨
    Vl_m, Vr_m = beamV(qmd, L2, rd['B_br'], -rd['B_br'])
    # ΔM/L
    dM_e = (rd['B_bl']-rd['A_b'])/L1
    dM_m = (-rd['B_br']-rd['B_br'])/L2 if abs(rd['B_br'])>0.01 else 0

    row = 3+ri
    t.rows[row].cells[1].text = f'{dM_e:.2f}'
    t.rows[row].cells[2].text = f'{dM_m:.2f}'
    # V_q
    Vq_e = qed*L1/2; Vq_m = qmd*L2/2
    t.rows[row].cells[5].text = f'{Vq_e:.2f}'
    t.rows[row].cells[6].text = f'{Vq_m:.2f}'
    # V_left, V_right
    t.rows[row].cells[7].text = f'{Vl_e:.2f}'
    t.rows[row].cells[8].text = f'{Vl_m:.2f}'
    t.rows[row].cells[9].text = f'{-Vr_e:.2f}'  # V_right downward negative
    t.rows[row].cells[10].text = f'{-Vr_m:.2f}'

# ============================================================
# 更新表6-14 活载梁端剪力 (Table 43)
# ============================================================
t = doc.tables[42]
print("更新表6-14 (活载梁端剪力)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']
    qel = data[nm]['qe_l']; qml = data[nm]['qm_l']
    Vl_e, Vr_e = beamV(qel, L1, rl['A_b'], rl['B_bl'])
    Vl_m, Vr_m = beamV(qml, L2, rl['B_br'], -rl['B_br'])
    dM_e = (rl['B_bl']-rl['A_b'])/L1
    dM_m = (-rl['B_br']-rl['B_br'])/L2 if abs(rl['B_br'])>0.01 else 0

    row = 3+ri
    t.rows[row].cells[1].text = f'{dM_e:.2f}'
    t.rows[row].cells[2].text = f'{dM_m:.2f}'
    Vq_e = qel*L1/2; Vq_m = qml*L2/2
    t.rows[row].cells[5].text = f'{Vq_e:.2f}'
    t.rows[row].cells[6].text = f'{Vq_m:.2f}'
    t.rows[row].cells[7].text = f'{Vl_e:.2f}'
    t.rows[row].cells[8].text = f'{Vl_m:.2f}'
    t.rows[row].cells[9].text = f'{-Vr_e:.2f}'
    t.rows[row].cells[10].text = f'{-Vr_m:.2f}'

# ============================================================
# 更新柱剪力和轴力表
# ============================================================
# 柱集中力 (5400版本)
sec_beam = 1.54*L1/2
F_e_r = 31.19+20.08+sec_beam+g_roof*(ss**2/4+ss*L1/2)
F_e_f = 42.3+20.08+sec_beam+g_floor*(ss**2/4+ss*L1/2)
F_m_r = 3.7+20.08+sec_beam+g_roof*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))
F_m_f = 41.95+20.08+sec_beam+g_floor*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))

E_l_r = q_roof*(ss**2/4+ss*L1/2)
E_l_f = q_floor*(ss**2/4+ss*L1/2)
M_l_r = q_roof*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))
M_l_f = q_floor*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))

Gc_std = 6.76*3.0; Gc_1st = 6.76*4.0

# 表6-8 恒载柱剪力 (Table 37)
t = doc.tables[36]
print("更新表6-8 (恒载柱剪力)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; h = 4.0 if nm=='1F' else 3.0
    # 边柱: 上端+下端弯矩(需要从相邻层取)
    # A柱的上端弯矩来自本层A_u, 下端来自本层A_l
    M_A_top = abs(rd['A_u']); M_A_bot = abs(rd['A_l'])
    M_B_top = abs(rd['B_u']); M_B_bot = abs(rd['B_l'])
    row = 3+ri
    t.rows[row].cells[1].text = f'{M_A_top+M_A_bot:.2f}'
    t.rows[row].cells[2].text = f'{M_B_top+M_B_bot:.2f}'
    if len(t.rows[row].cells) > 5:
        t.rows[row].cells[4].text = f'{-((M_A_top+M_A_bot)/h):.2f}'
        t.rows[row].cells[5].text = f'{((M_B_top+M_B_bot)/h):.2f}'

# 表6-15 活载柱剪力 (Table 44)
t = doc.tables[43]
print("更新表6-15 (活载柱剪力)...")
for ri, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']; h = 4.0 if nm=='1F' else 3.0
    M_A_top = abs(rl['A_u']); M_A_bot = abs(rl['A_l'])
    M_B_top = abs(rl['B_u']); M_B_bot = abs(rl['B_l'])
    row = 3+ri
    t.rows[row].cells[1].text = f'{M_A_top+M_A_bot:.2f}'
    t.rows[row].cells[2].text = f'{M_B_top+M_B_bot:.2f}'
    if len(t.rows[row].cells) > 5:
        t.rows[row].cells[4].text = f'{-((M_A_top+M_A_bot)/h):.2f}'
        t.rows[row].cells[5].text = f'{((M_B_top+M_B_bot)/h):.2f}'

# 柱轴力 (需要逐层累计)
# 表6-9 边柱轴力恒载 (Table 38)
t = doc.tables[37]
print("更新表6-9 (边柱轴力恒载)...")
N_edge_top = [0]*6; N_edge_bot = [0]*6
for fi, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; qed = data[nm]['qe_d']
    Vl_e, Vr_e = beamV(qed, L1, rd['A_b'], rd['B_bl'])
    F = F_e_r if nm=='6F' else F_e_f
    Gc = Gc_1st if nm=='1F' else Gc_std
    # N_top = F + V_beam_left (向下)
    N_top = F + Vl_e
    if fi > 0:
        N_top += N_edge_bot[fi-1]
    N_bot = N_top + Gc
    N_edge_top[fi] = N_top; N_edge_bot[fi] = N_bot
    row = 2+fi
    if fi==0:
        t.rows[row].cells[0].text = '6F'
        t.rows[row].cells[1].text = f'{F:.2f}'
        t.rows[row].cells[2].text = f'{Gc:.2f}'
        t.rows[row].cells[3].text = '0.00'
        t.rows[row].cells[4].text = f'{Vl_e:.2f}'
        t.rows[row].cells[5].text = f'{N_top:.2f}'
        t.rows[row].cells[6].text = f'{N_bot:.2f}'
    else:
        t.rows[row].cells[1].text = f'{F:.2f}'
        t.rows[row].cells[2].text = f'{Gc:.2f}'
        t.rows[row].cells[3].text = '0.00'
        t.rows[row].cells[4].text = f'{Vl_e:.2f}'
        t.rows[row].cells[5].text = f'{N_top:.2f}'
        t.rows[row].cells[6].text = f'{N_bot:.2f}'

# 表6-10 中柱轴力恒载 (Table 39)
t = doc.tables[38]
print("更新表6-10 (中柱轴力恒载)...")
N_mid_top = [0]*6; N_mid_bot = [0]*6
for fi, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rd = data[nm]['dead']; qed = data[nm]['qe_d']; qmd = data[nm]['qm_d']
    Vl_e, Vr_e = beamV(qed, L1, rd['A_b'], rd['B_bl'])
    Vl_m, Vr_m = beamV(qmd, L2, rd['B_br'], -rd['B_br'])
    F = F_m_r if nm=='6F' else F_m_f
    Gc = Gc_1st if nm=='1F' else Gc_std
    N_top = F + Vr_e + Vl_m  # 左梁右端+右梁左端 (均向下)
    if fi > 0:
        N_top += N_mid_bot[fi-1]
    N_bot = N_top + Gc
    N_mid_top[fi] = N_top; N_mid_bot[fi] = N_bot
    row = 2+fi
    if fi==0:
        t.rows[row].cells[0].text = '6F'
        t.rows[row].cells[1].text = f'{F:.2f}'
        t.rows[row].cells[2].text = f'{Gc:.2f}'
        t.rows[row].cells[3].text = f'{-Vr_e:.2f}'
        t.rows[row].cells[4].text = f'{Vl_m:.2f}'
        t.rows[row].cells[5].text = f'{N_top:.2f}'
        t.rows[row].cells[6].text = f'{N_bot:.2f}'
    else:
        t.rows[row].cells[1].text = f'{F:.2f}'
        t.rows[row].cells[2].text = f'{Gc:.2f}'
        t.rows[row].cells[3].text = f'{-Vr_e:.2f}'
        t.rows[row].cells[4].text = f'{Vl_m:.2f}'
        t.rows[row].cells[5].text = f'{N_top:.2f}'
        t.rows[row].cells[6].text = f'{N_bot:.2f}'

# 表6-16 柱轴力活载 (Table 45)
t = doc.tables[44]
print("更新表6-16 (柱轴力活载)...")
N_e_l_top = [0]*6; N_m_l_top = [0]*6
for fi, nm in enumerate(['6F','5F','4F','3F','2F','1F']):
    rl = data[nm]['live']; qel = data[nm]['qe_l']; qml = data[nm]['qm_l']
    Vl_e, Vr_e = beamV(qel, L1, rl['A_b'], rl['B_bl'])
    Vl_m, Vr_m = beamV(qml, L2, rl['B_br'], -rl['B_br'])
    Fe_l = E_l_r if nm=='6F' else E_l_f
    Fm_l = M_l_r if nm=='6F' else M_l_f
    N_e = Fe_l + Vl_e
    N_m = Fm_l + Vr_e + Vl_m
    if fi > 0:
        N_e += N_e_l_top[fi-1]; N_m += N_m_l_top[fi-1]
    N_e_l_top[fi] = N_e; N_m_l_top[fi] = N_m
    row = 2+fi
    # 活载轴力表格式: 楼层 | V_left | V_right | V_mid_left | F_edge | F_mid | N_edge | N_mid
    # 根据原表结构调整
    if fi==0 and len(t.rows[row].cells) >= 8:
        t.rows[row].cells[1].text = f'{-Vr_e:.2f}'
        t.rows[row].cells[2].text = f'{Vl_e:.2f}'
        t.rows[row].cells[3].text = f'{Vl_m:.2f}'
        t.rows[row].cells[4].text = f'{Fe_l:.2f}'
        t.rows[row].cells[5].text = f'{Fm_l:.2f}'
        t.rows[row].cells[6].text = f'{N_e:.2f}'
        t.rows[row].cells[7].text = f'{N_m:.2f}'
    elif len(t.rows[row].cells) >= 8:
        t.rows[row].cells[1].text = f'{-Vr_e:.2f}'
        t.rows[row].cells[2].text = f'{Vl_e:.2f}'
        t.rows[row].cells[3].text = f'{Vl_m:.2f}'
        t.rows[row].cells[4].text = f'{Fe_l:.2f}'
        t.rows[row].cells[5].text = f'{Fm_l:.2f}'
        t.rows[row].cells[6].text = f'{N_e:.2f}'
        t.rows[row].cells[7].text = f'{N_m:.2f}'

print(f"\n跨中弯矩、梁端剪力、柱剪力、柱轴力 - 全部更新完成")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"文件已保存")
