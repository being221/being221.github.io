#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新第7章 表7-5~7-8 内力转换
恒载+活载 → 梁端弯矩转柱边弯矩
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

# ============================================================
# 基本参数
# ============================================================
L1, L2 = 5.4, 2.4
ss = 3.45
B_COL = 0.5  # 柱宽
B_HALF = B_COL / 2  # 0.25m
BETA = 0.85

g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0

F_NEW = 1 - 2*(0.5*ss/L1)**2 + (0.5*ss/L1)**3

# 等效均布荷载
qe_d_r = g_beam_e + F_NEW*ss*g_roof
qe_d_f = g_beam_e + g_wall_e + F_NEW*ss*g_floor
qm_d_r = g_beam_m + 0.625*L2*g_roof
qm_d_f = g_beam_m + g_wall_m + 0.625*L2*g_floor
qe_l_r = F_NEW*ss*q_roof
qe_l_f = F_NEW*ss*q_floor
qm_l_r = 0.625*L2*q_roof
qm_l_f = 0.625*L2*q_floor

# 柱线刚度
E = 30e6
Ic = 0.5**4/12
ic_std = E*Ic/3.0
ic_1st = E*Ic/4.0
ib_edge = E*1.5*0.25*0.5**3/12/L1
ib_mid = E*1.5*0.25*0.4**3/12/L2

print(f"F_NEW={F_NEW:.4f}")
print(f"边跨恒载q: 屋面={qe_d_r:.2f}, 楼面={qe_d_f:.2f}")
print(f"中跨恒载q: 屋面={qm_d_r:.2f}, 楼面={qm_d_f:.2f}")
print(f"边跨活载q: 屋面={qe_l_r:.2f}, 楼面={qe_l_f:.2f}")
print(f"中跨活载q: 屋面={qm_l_r:.2f}, 楼面={qm_l_f:.2f}")

# ============================================================
# 弯矩二次分配
# ============================================================
def fem(q, L):
    return q*L**2/12

def midM(q, L, Ml, Mr):
    return q*L**2/8 - (abs(Ml) + abs(Mr))/2

def moment_dist(ic_u, ic_l, ib_e, ib_m, q_e, q_m):
    FE_e = fem(q_e, L1); FE_m = fem(q_m, L2)
    sA = ic_u + ic_l + ib_e; sB = ic_u + ic_l + ib_e + ib_m
    mu_A_u = ic_u/sA; mu_A_l = ic_l/sA; mu_A_b = ib_e/sA
    mu_B_u = ic_u/sB; mu_B_l = ic_l/sB; mu_B_bl = ib_e/sB; mu_B_br = ib_m/sB

    M_A_u = M_A_l = 0; M_A_b = -FE_e
    M_B_u = M_B_l = 0; M_B_bl = FE_e; M_B_br = -FE_m

    # 1st distribution
    uA = M_A_u + M_A_l + M_A_b
    dAb = -mu_A_b*uA; dAu = -mu_A_u*uA; dAl = -mu_A_l*uA
    M_A_b += dAb; M_A_u += dAu; M_A_l += dAl
    uB = M_B_u + M_B_l + M_B_bl + M_B_br
    dBbl = -mu_B_bl*uB; dBbr = -mu_B_br*uB; dBu = -mu_B_u*uB; dBl = -mu_B_l*uB
    M_B_bl += dBbl; M_B_br += dBbr; M_B_u += dBu; M_B_l += dBl
    # 1st carry-over
    M_A_b += 0.5*dBbl; M_B_bl += 0.5*dAb; M_B_br += 0.5*(-dBbr)

    # 2nd distribution
    uA2 = M_A_u + M_A_l + M_A_b
    dAb2 = -mu_A_b*uA2; dAu2 = -mu_A_u*uA2; dAl2 = -mu_A_l*uA2
    M_A_b += dAb2; M_A_u += dAu2; M_A_l += dAl2
    uB2 = M_B_u + M_B_l + M_B_bl + M_B_br
    dBbl2 = -mu_B_bl*uB2; dBbr2 = -mu_B_br*uB2; dBu2 = -mu_B_u*uB2; dBl2 = -mu_B_l*uB2
    M_B_bl += dBbl2; M_B_br += dBbr2; M_B_u += dBu2; M_B_l += dBl2
    # 2nd carry-over
    M_A_b += 0.5*dBbl2; M_B_bl += 0.5*dAb2; M_B_br += 0.5*(-dBbr2)

    return dict(A_u=M_A_u, A_l=M_A_l, A_b=M_A_b,
                B_u=M_B_u, B_l=M_B_l, B_bl=M_B_bl, B_br=M_B_br)

floors = [
    ('6F', 0, ic_std, qe_d_r, qm_d_r, qe_l_r, qm_l_r),
    ('5F', ic_std, ic_std, qe_d_f, qm_d_f, qe_l_f, qm_l_f),
    ('4F', ic_std, ic_std, qe_d_f, qm_d_f, qe_l_f, qm_l_f),
    ('3F', ic_std, ic_std, qe_d_f, qm_d_f, qe_l_f, qm_l_f),
    ('2F', ic_std, ic_std, qe_d_f, qm_d_f, qe_l_f, qm_l_f),
    ('1F', ic_std, ic_1st, qe_d_f, qm_d_f, qe_l_f, qm_l_f),
]

data = {}
for nm, icu, icl, qed, qmd, qel, qml in floors:
    rd = moment_dist(icu, icl, ib_edge, ib_mid, qed, qmd)
    rl = moment_dist(icu, icl, ib_edge, ib_mid, qel, qml)
    data[nm] = {'dead': rd, 'live': rl, 'qe_d': qed, 'qm_d': qmd, 'qe_l': qel, 'qm_l': qml}

# ============================================================
# 计算内力转换值
# ============================================================
def beam_shear(q, L, Ml, Mr):
    """V_left = qL/2 + (Mr-Ml)/L, V_right = qL/2 + (Ml-Mr)/L"""
    Vl = q*L/2 + (Mr - Ml)/L
    Vr = q*L/2 + (Ml - Mr)/L
    return Vl, Vr

def convert_to_face(M_node, V, q, b_half):
    """
    竖向荷载内力转换:
    M_face = M_node - V*b/2 + q*(b/2)²/2
    V_face = V - q*b/2
    使用绝对值(统一符号约定: 使柱边弯矩为正表示顶部受拉)
    """
    M_face = abs(M_node) + abs(V)*b_half - q*b_half**2/2
    V_face = abs(V) - q*b_half
    return M_face, V_face

print("\n" + "="*80)
print("表7-5 恒载内力转换")
print("="*80)

for nm in ['6F', '5F', '4F', '3F', '2F', '1F']:
    rd = data[nm]['dead']
    qed = data[nm]['qe_d']
    qmd = data[nm]['qm_d']

    # 调幅前
    M_lb = rd['A_b']   # 边跨左端
    M_rb = rd['B_bl']  # 边跨右端
    M_mb = rd['B_br']  # 中跨左端

    # 调幅后
    M_la = BETA * M_lb
    M_ra = BETA * M_rb
    M_ma = BETA * M_mb

    # 调幅前剪力 (用于col2)
    Vl_b, Vr_b = beam_shear(qed, L1, M_lb, M_rb)
    Vm_b, _ = beam_shear(qmd, L2, M_mb, -M_mb)  # 中跨对称

    # 调幅后剪力
    Vl_a, Vr_a = beam_shear(qed, L1, M_la, M_ra)
    Vm_a, _ = beam_shear(qmd, L2, M_ma, -M_ma)

    # 柱边转换 (用调幅后M + 调幅前V)
    M_le, Vl_e = convert_to_face(M_la, Vl_b, qed, B_HALF)
    M_re, Vr_e = convert_to_face(M_ra, Vr_b, qed, B_HALF)
    M_me, Vm_e = convert_to_face(M_ma, Vm_b, qmd, B_HALF)

    print(f"\n{nm}:")
    print(f"  边跨左: V={Vl_b:.2f}, q={qed:.2f}, M_调幅后={M_la:.2f}, M_柱边={M_le:.2f}, V_柱边={Vl_e:.2f}")
    print(f"  边跨右: V={Vr_b:.2f}, q={qed:.2f}, M_调幅后={M_ra:.2f}, M_柱边={M_re:.2f}, V_柱边={Vr_e:.2f}")
    print(f"  中跨左: V={Vm_b:.2f}, q={qmd:.2f}, M_调幅后={M_ma:.2f}, M_柱边={M_me:.2f}, V_柱边={Vm_e:.2f}")

# ============================================================
# 更新表7-5 (T50, index 49)
# ============================================================
t = doc.tables[49]
print("\n更新表7-5 (恒载内力转换)...")

floor_names = ['6F', '5F', '4F', '3F', '2F', '1F']
sections = ['边跨左', '边跨右', '中跨左']
row_idx = 3

for nm in floor_names:
    rd = data[nm]['dead']
    qed = data[nm]['qe_d']
    qmd = data[nm]['qm_d']

    M_lb = rd['A_b']; M_rb = rd['B_bl']; M_mb = rd['B_br']
    M_la = BETA * M_lb; M_ra = BETA * M_rb; M_ma = BETA * M_mb
    Vl_b, Vr_b = beam_shear(qed, L1, M_lb, M_rb)
    Vm_b, _ = beam_shear(qmd, L2, M_mb, -M_mb)

    M_le, Vl_e = convert_to_face(M_la, Vl_b, qed, B_HALF)
    M_re, Vr_e = convert_to_face(M_ra, Vr_b, qed, B_HALF)
    M_me, Vm_e = convert_to_face(M_ma, Vm_b, qmd, B_HALF)

    row_data = [
        (f'{Vl_b:.2f}', f'{qed:.2f}', f'{M_la:.2f}', f'{M_le:.2f}', f'{Vl_e:.2f}'),
        (f'{Vr_b:.2f}', f'{qed:.2f}', f'{M_ra:.2f}', f'{M_re:.2f}', f'{Vr_e:.2f}'),
        (f'{Vm_b:.2f}', f'{qmd:.2f}', f'{M_ma:.2f}', f'{M_me:.2f}', f'{Vm_e:.2f}'),
    ]

    for si, sec in enumerate(sections):
        row = t.rows[row_idx]
        row.cells[2].text = row_data[si][0]
        row.cells[3].text = row_data[si][1]
        row.cells[4].text = row_data[si][2]
        row.cells[5].text = row_data[si][3]
        row.cells[6].text = row_data[si][4]
        row_idx += 1

print("表7-5 更新完成")

# ============================================================
# 更新表7-6 (T51, index 50) 活载内力转换
# ============================================================
t = doc.tables[50]
print("\n更新表7-6 (活载内力转换)...")

row_idx = 3
for nm in floor_names:
    rl = data[nm]['live']
    qel = data[nm]['qe_l']
    qml = data[nm]['qm_l']

    M_lb = rl['A_b']; M_rb = rl['B_bl']; M_mb = rl['B_br']
    M_la = BETA * M_lb; M_ra = BETA * M_rb; M_ma = BETA * M_mb
    Vl_b, Vr_b = beam_shear(qel, L1, M_lb, M_rb)
    Vm_b, _ = beam_shear(qml, L2, M_mb, -M_mb)

    M_le, Vl_e = convert_to_face(M_la, Vl_b, qel, B_HALF)
    M_re, Vr_e = convert_to_face(M_ra, Vr_b, qel, B_HALF)
    M_me, Vm_e = convert_to_face(M_ma, Vm_b, qml, B_HALF)

    row_data = [
        (f'{Vl_b:.2f}', f'{qel:.2f}', f'{M_la:.2f}', f'{M_le:.2f}', f'{Vl_e:.2f}'),
        (f'{Vr_b:.2f}', f'{qel:.2f}', f'{M_ra:.2f}', f'{M_re:.2f}', f'{Vr_e:.2f}'),
        (f'{Vm_b:.2f}', f'{qml:.2f}', f'{M_ma:.2f}', f'{M_me:.2f}', f'{Vm_e:.2f}'),
    ]

    for si, sec in enumerate(sections):
        row = t.rows[row_idx]
        row.cells[2].text = row_data[si][0]
        row.cells[3].text = row_data[si][1]
        row.cells[4].text = row_data[si][2]
        row.cells[5].text = row_data[si][3]
        row.cells[6].text = row_data[si][4]
        row_idx += 1

print("表7-6 更新完成")

# ============================================================
# 保存
# ============================================================
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n文件已保存: {DOC}")
print(f"审阅版: {review}")
