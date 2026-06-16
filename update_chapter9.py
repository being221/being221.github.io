#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
更新第9章楼板设计: A板 4.8→5.4m, 高跨比 1.39→1.57
弯矩系数按规范表格内插, 含泊松比调整(ν=0.2)
"""
import sys, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)
nu = 0.2  # 泊松比

# 板格参数
lx_A, ly_A = 3.45, 5.4
lam_new = ly_A / lx_A  # 1.565
lam_old = 4.8 / lx_A    # 1.39

# 荷载
g, q = 4.2, 2.0
p_design = 1.3*g + 1.5*q     # 8.46
p_mid    = 1.3*g + 1.5*q/2   # 6.96
p_live_h = 1.5 * q/2          # 1.5

print(f"A板 λ: {lam_old:.2f}→{lam_new:.3f}")

# ============================================================
# 弯矩系数内插
# ============================================================
# 标准表(四边固定):
# λ:  1.40     1.50     1.60
std_table = {
    1.40: [0.0678, 0.0309, 0.0266, 0.0117, -0.0747, -0.0474],
    1.50: [0.0656, 0.0275, 0.0280, 0.0106, -0.0783, -0.0461],
    1.60: [0.0634, 0.0246, 0.0291, 0.0097, -0.0813, -0.0449],
}

def interp(lam):
    if lam <= 1.40: return std_table[1.40]
    if lam >= 1.60: return std_table[1.60]
    lo = 1.50 if lam > 1.50 else 1.40
    hi = 1.60 if lam > 1.50 else 1.50
    t = (lam - lo) / (hi - lo)
    return [std_table[lo][i] + t*(std_table[hi][i]-std_table[lo][i]) for i in range(6)]

# 缩放因子 (标准表新旧比值)
scales = [ns/os for ns, os in zip(interp(lam_new), interp(lam_old))]
# 原系数 (λ=1.39, 教科书版本)
coeff_old = [0.0659, 0.0304, 0.0347, 0.0131, -0.0765, -0.0572]
# 新系数
coeff_new = [round(coeff_old[i] * scales[i], 4) for i in range(6)]

print("系数: 简支短/长, 跨中短/长, 支座短/长")
print(f"  旧: {coeff_old}")
print(f"  新: {coeff_new}")

# ============================================================
# 弯矩计算 (lx² = 11.9025)
# ============================================================
lx2 = lx_A ** 2
M = {}
M['mid_s_0'] = coeff_new[2] * p_mid * lx2    # 跨中短(ν=0)
M['mid_l_0'] = coeff_new[3] * p_mid * lx2    # 跨中长(ν=0)
M['sup_s']   = coeff_new[4] * p_design * lx2  # 支座短
M['sup_l']   = coeff_new[5] * p_design * lx2  # 支座长

# 简支部分(活载半值)
M['simp_s'] = coeff_new[0] * p_live_h * lx2
M['simp_l'] = coeff_new[1] * p_live_h * lx2

# ν=0前的弯矩(用于T78显示)
M_mid_s_raw = M['mid_s_0'] + M['simp_s']
M_mid_l_raw = M['mid_l_0'] + M['simp_l']

# 泊松比调整: M_ν = M(ν=0) + ν × M_other(ν=0)
# 注意: 跨中短/长的交叉项需要用简支+跨中的总弯矩
M_mid_s_v0 = M_mid_s_raw
M_mid_l_v0 = M_mid_l_raw

M_mid_s = M_mid_s_v0 + nu * M_mid_l_v0  # 跨中短 ν=0.2
M_mid_l = M_mid_l_v0 + nu * M_mid_s_v0  # 跨中长 ν=0.2

# 支座泊松比调整
# 支座不进行ν调整 (直接使用ν=0值)
M_sup_s_raw = abs(M['sup_s'])
M_sup_l_raw = abs(M['sup_l'])

print(f"\n弯矩(ν=0): 跨中短={M_mid_s_v0:.2f}, 跨中长={M_mid_l_v0:.2f}")
print(f"弯矩(ν=0.2): 跨中短={M_mid_s:.2f}, 跨中长={M_mid_l:.2f}")
print(f"支座(ν=0): 短边={M_sup_s_raw:.2f}, 长边={M_sup_l_raw:.2f}")

# ============================================================
# 配筋计算
# ============================================================
fc, ft, fy = 14.3, 1.43, 360
h_slab, b = 120, 1000
rho_min = max(0.002, 0.45*ft/fy)
As_min = rho_min * b * h_slab  # 240mm²

def calc_As(M_val, d):
    alpha_s = M_val * 1e6 / (fc * b * d**2)
    xi = 1 - math.sqrt(1 - 2*alpha_s)
    if xi > 0.518: xi = 0.518
    gamma_s = 1 - xi/2
    As = M_val * 1e6 / (fy * gamma_s * d)
    return alpha_s, xi, max(As, As_min)

d_s, d_l = 100, 90
print(f"\n配筋 (As_min={As_min:.0f}mm²):")
As_results = {}
for name, M_val, d in [('跨中短', M_mid_s, d_s), ('跨中长', M_mid_l, d_l),
                        ('支座A/A', M_sup_s_raw, d_s), ('支座A/B', M_sup_l_raw, d_s)]:
    a, xi, As = calc_As(M_val, d)
    As_results[name] = (a, xi, As)
    print(f"  {name}: M={M_val:.2f}, αs={a:.4f}, As={As:.0f}mm²")

# ============================================================
# 写入DOCX
# ============================================================
print("\n写入DOCX...")

# T76: 高跨比
t = doc.tables[76]
t.rows[2].cells[1].text = f'{ly_A:.1f}/{lx_A:.2f}={lam_new:.2f}'
print("  T76 ✓")

# T77: 弯矩系数
t = doc.tables[77]
t.rows[3].cells[1].text = f'{coeff_new[0]:.4f}'
t.rows[3].cells[2].text = f'{coeff_new[1]:.4f}'
t.rows[3].cells[3].text = f'{coeff_new[2]:.4f}'
t.rows[3].cells[4].text = f'{coeff_new[3]:.4f}'
t.rows[3].cells[5].text = f'{coeff_new[4]:.4f}'
t.rows[3].cells[6].text = f'{coeff_new[5]:.4f}'
print("  T77 ✓")

# T78: 弯矩计算
t = doc.tables[78]
# R2: A支座短
t.rows[2].cells[3].text = f'{coeff_new[4]:.4f}×{p_design:.2f}×{lx_A:.2f}²={M["sup_s"]:.2f}'
t.rows[2].cells[4].text = f'{M["sup_s"]:.2f}'
# R3: A支座长
t.rows[3].cells[3].text = f'{coeff_new[5]:.4f}×{p_design:.2f}×{lx_A:.2f}²={M["sup_l"]:.2f}'
t.rows[3].cells[4].text = f'{M["sup_l"]:.2f}'
# R4: A跨中短
t.rows[4].cells[3].text = f'{coeff_new[2]:.4f}×{p_mid:.2f}×{lx_A:.2f}²+{coeff_new[0]:.4f}×{p_live_h:.2f}×{lx_A:.2f}²='
t.rows[4].cells[4].text = f'{M_mid_s_v0:.2f}+{M_mid_l_v0:.2f}×0.2={M_mid_s:.2f}'
# R5: A跨中长
t.rows[5].cells[3].text = f'{coeff_new[3]:.4f}×{p_mid:.2f}×{lx_A:.2f}²+{coeff_new[1]:.4f}×{p_live_h:.2f}×{lx_A:.2f}²='
t.rows[5].cells[4].text = f'{M_mid_l_v0:.2f}+{M_mid_s_v0:.2f}×0.2={M_mid_l:.2f}'
print("  T78 ✓")

# T79: 配筋
# R2=跨中A长边, R3=跨中A短边, R6=支座A/A, R7=支座A/B
t = doc.tables[79]

# R2: 跨中 A板 长边 d=90
a, xi, As = As_results['跨中长']
t.rows[2].cells[3].text = f'{d_l}'
t.rows[2].cells[4].text = f'{M_mid_l:.2f}'
t.rows[2].cells[5].text = f'{a:.4f}'
t.rows[2].cells[6].text = f'{xi:.4f}'
if len(t.rows[2].cells) > 7:
    t.rows[2].cells[7].text = f'{As:.0f}'

# R3: 跨中 A板 短边 d=100
a, xi, As = As_results['跨中短']
t.rows[3].cells[3].text = f'{d_s}'
t.rows[3].cells[4].text = f'{M_mid_s:.2f}'
t.rows[3].cells[5].text = f'{a:.4f}'
t.rows[3].cells[6].text = f'{xi:.4f}'
if len(t.rows[3].cells) > 7:
    t.rows[3].cells[7].text = f'{As:.0f}'

# R6: 支座 A/A (短边) — ν=0, 不调整
a, xi, As = As_results['支座A/A']
t.rows[6].cells[3].text = f'{d_s}'
t.rows[6].cells[4].text = f'{-M_sup_s_raw:.2f}'
t.rows[6].cells[5].text = f'{a:.4f}'
t.rows[6].cells[6].text = f'{xi:.4f}'
if len(t.rows[6].cells) > 7:
    t.rows[6].cells[7].text = f'{As:.0f}'

# R7: 支座 A/B (长边) — ν=0, 不调整
a, xi, As = As_results['支座A/B']
t.rows[7].cells[3].text = f'{d_s}'
t.rows[7].cells[4].text = f'{-M_sup_l_raw:.2f}'
t.rows[7].cells[5].text = f'{a:.4f}'
t.rows[7].cells[6].text = f'{xi:.4f}'
if len(t.rows[7].cells) > 7:
    t.rows[7].cells[7].text = f'{As:.0f}'

print("  T79 ✓")

# 保存
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n第9章更新完成! {DOC}")
