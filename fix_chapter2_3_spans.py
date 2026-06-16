#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""修复第2-3章残留旧跨度值: 4.8→5.4 + 级联合计值"""
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

DOC = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC)

# ============================================================
# T9 (表3-3 次梁自重): 4处 "1.54×4.8/2=3.7kN" → "1.54×5.4/2=4.16kN"
# ============================================================
t = doc.tables[9]
for r in [2, 8, 14, 18]:
    old = t.rows[r].cells[3].text
    t.rows[r].cells[3].text = '1.54×5.4/2=4.16kN'
    print(f"  T9 R{r} C3: '{old}' → '1.54×5.4/2=4.16kN'")

# ============================================================
# T12 (表3-6 屋面重力荷载代表值): R4内墙, R9合计
# ============================================================
t = doc.tables[12]

# R4 C2: 内墙自重公式 4.8→5.4, 结果 1063.8→1138.1
old = t.rows[4].cells[2].text
new_text = old.replace('(2.48×(3-0.4)×4.8×24×0.8)', '(2.48×(3-0.4)×5.4×24×0.8)')
new_text = new_text.replace('=1063.8kN', '=1138.1kN')
t.rows[4].cells[2].text = new_text
print(f"  T12 R4 C2: 4.8→5.4, 1063.8→1138.1")

# R9 C1: 合计公式, 1063.8/2→1138.1/2, 5181→5218
delta = (1138.1 - 1063.8) / 2  # 37.15
old = t.rows[9].cells[1].text
new_text = old.replace('1063.8/2', '1138.1/2').replace('=5181kN', '=5218kN')
t.rows[9].cells[1].text = new_text
print(f"  T12 R9 C1: 5181→5218")

# R9 C2: 同上
old = t.rows[9].cells[2].text
new_text = old.replace('1063.8/2', '1138.1/2').replace('=4924kN', '=4961kN')
t.rows[9].cells[2].text = new_text
print(f"  T12 R9 C2: 4924→4961")

# ============================================================
# T13 (表3-7 各层重力荷载代表值): R4,R15,R16内墙 + R8,R21合计
# ============================================================
t = doc.tables[13]

# R4 C2: 中间层内墙 (同T12)
old = t.rows[4].cells[2].text
new_text = old.replace('(2.48×(3-0.4)×4.8×24×0.8)', '(2.48×(3-0.4)×5.4×24×0.8)')
new_text = new_text.replace('=1063.8kN', '=1138.1kN')
t.rows[4].cells[2].text = new_text
print(f"  T13 R4 C2: 4.8→5.4, 1063.8→1138.1")

# R8 C1: 中间层合计, 1063.8→1138.1, 5401→5475
delta_mid = 1138.1 - 1063.8  # 74.3
old = t.rows[8].cells[1].text
new_text = old.replace('1063.8', '1138.1').replace('=5401kN', '=5475kN')
t.rows[8].cells[1].text = new_text
print(f"  T13 R8 C1: 5401→5475")

# R8 C2: 中间层合计(边跨)
old = t.rows[8].cells[2].text
new_text = old.replace('1063.8', '1138.1').replace('=5136kN', '=5210kN')
t.rows[8].cells[2].text = new_text
print(f"  T13 R8 C2: 5136→5210")

# R15 C2: 首层内墙 1484.01→1586.87
old = t.rows[15].cells[2].text
new_text = old.replace('(2.48×(4-0.4)×4.8×24×0.8)', '(2.48×(4-0.4)×5.4×24×0.8)')
new_text = new_text.replace('=1484.01kN', '=1586.87kN')
t.rows[15].cells[2].text = new_text
print(f"  T13 R15 C2: 4.8→5.4, 1484.01→1586.87")

# R16 C2: 二层内墙
old = t.rows[16].cells[2].text
new_text = old.replace('(2.48×(3-0.4)×4.8×24×0.8)', '(2.48×(3-0.4)×5.4×24×0.8)')
new_text = new_text.replace('=1063.8kN', '=1138.1kN')
t.rows[16].cells[2].text = new_text
print(f"  T13 R16 C2: 4.8→5.4, 1063.8→1138.1")

# R21 C1: 首层合计, (1484.01+1063.8)/2→(1586.87+1138.1)/2, 5813→5902
delta_1st = ((1586.87 + 1138.1) - (1484.01 + 1063.8)) / 2  # 88.58
old = t.rows[21].cells[1].text
new_text = old.replace('(1484.01+1063.8)/2', '(1586.87+1138.1)/2').replace('=5813kN', '=5902kN')
t.rows[21].cells[1].text = new_text
print(f"  T13 R21 C1: 5813→5902")

# R21 C2: 首层合计(边跨), 5548→5637
old = t.rows[21].cells[2].text
new_text = old.replace('(1484.01+1063.8)/2', '(1586.87+1138.1)/2').replace('=5548kN', '=5637kN')
t.rows[21].cells[2].text = new_text
print(f"  T13 R21 C2: 5548→5637")

# 保存
doc.save(DOC)
review = DOC.replace('修正版', '审阅版')
doc.save(review)
print(f"\n第2-3章修复完成! (含级联合计)")
print(f"修正版: {DOC}")
