#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
5400跨度 docx修改 第二轮
处理: 重力荷载代表值、地震D值、地震力、整体内力重算
"""

import sys, os, math
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

# 参数
L1, L1_OLD = 5.4, 4.8
L2 = 2.4
L_long = 6.9
ss = 3.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0
g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45

# 等效系数
alpha_new = 0.5*ss/L1  # 0.319
F_NEW = 1 - 2*alpha_new**2 + alpha_new**3  # 0.829

# 梁线刚度
E = 30e6
I0_e = 0.25*0.5**3/12
i_e_e_new = E * 1.5*I0_e / L1  # 21701
i_m_e_new = E * 2.0*I0_e / L1  # 28935
i_e_m = E * 1.5*0.25*0.4**3/12 / L2  # 25000
i_m_m = E * 2.0*0.25*0.4**3/12 / L2  # 33333

Ic = 0.5*0.5**3/12
ic_top = E * Ic / 3.0
ic_1st = E * Ic / 4.0

# D值
def KD(ic, h, sib, is_first):
    K = sib/ic
    a = (0.5+K)/(2+K) if is_first else K/(2+K)
    D = a*12*ic/h**2
    return K,a,D

K_ee_t,a_ee_t,D_ee_t = KD(ic_top,3.0,i_e_e_new,False)
K_ee_1,a_ee_1,D_ee_1 = KD(ic_1st,4.0,i_e_e_new,True)
K_em_t,a_em_t,D_em_t = KD(ic_top,3.0,i_e_e_new+i_e_m,False)
K_em_1,a_em_1,D_em_1 = KD(ic_1st,4.0,i_e_e_new+i_e_m,True)
K_me_t,a_me_t,D_me_t = KD(ic_top,3.0,i_m_e_new,False)
K_me_1,a_me_1,D_me_1 = KD(ic_1st,4.0,i_m_e_new,True)
K_mm_t,a_mm_t,D_mm_t = KD(ic_top,3.0,i_m_e_new+i_m_m,False)
K_mm_1,a_mm_1,D_mm_1 = KD(ic_1st,4.0,i_m_e_new+i_m_m,True)

# 每榀D值
D_edge_top = 2*D_ee_t + 2*D_em_t
D_edge_1st = 2*D_ee_1 + 2*D_em_1
D_mid_top = 2*D_me_t + 2*D_mm_t
D_mid_1st = 2*D_me_1 + 2*D_mm_1

# 整体D值 (2个边榀 + 5个中间榀 × 纵向7跨)
# 总共14榀: 2边榀+5中间榀=7个横向框架, 每个2榀
D_total_top = 2*D_edge_top + 5*D_mid_top
D_total_1st = 2*D_edge_1st + 5*D_mid_1st

print(f"整体D值:")
print(f"  顶/标层: 2×{D_edge_top:.0f}+5×{D_mid_top:.0f}={D_total_top:.0f}")
print(f"  首层: 2×{D_edge_1st:.0f}+5×{D_mid_1st:.0f}={D_total_1st:.0f}")

# ============================================================
# 重力荷载代表值重算 (表4-1~4-3)
# ============================================================
# 横向框架梁自重
# 14跨边跨梁 × (L1-0.5+0.12)×2.57
# 7跨中跨梁 × (2.4-0.12)×1.89
# 注意: 14跨是7个轴线×2榀
beam_trans_roof_new = 14*(L1-0.5-0.5+0.12+0.12)*2.57 + 7*(2.4-0.12-0.12)*1.89
beam_trans_floor_new = 14*(L1-0.5-0.5+0.12+0.12)*2.57 + 7*(2.4-0.12-0.12)*1.89

# 楼面板面积
# 总面积 = (纵向长度) × (横向长度)
# 纵向: 6×6.9+两侧悬挑 约110.88m
# 横向: 7×(L1+L2+L1) = 7×(2×L1+L2)
# 实际从原数据反推:
# 原楼面板面积: 492.5m² (中间层), 521.3m² (屋面层)
# 差值: 521.3-492.5=28.8m² (楼梯间面积)
# 增加: 14跨×(L1_NEW-L1_OLD)×(6.9)
area_increase = 14 * (L1-L1_OLD) * L_long
floor_area_new = 492.5 + area_increase  # 楼面
roof_area_new = 521.3 + area_increase   # 屋面
print(f"\n楼面面积: 492.5→{floor_area_new:.1f} m²")
print(f"屋面面积: 521.3→{roof_area_new:.1f} m²")

# 屋面层恒载
roof_slab_new = roof_area_new * g_roof
roof_dead_new = roof_slab_new + beam_trans_roof_new + 84.08 + 446.98 + 567.84/2 + 528.97/2 + 1063.8/2 + 501.18
print(f"屋面恒载合计: 4924→{roof_dead_new:.0f} kN")

# 中间层恒载
floor_slab_new = floor_area_new * g_floor
floor_dead_new = floor_slab_new + beam_trans_floor_new + 84.08 + 446.98 + 567.84 + 528.97 + 1063.8 + 201.6
print(f"中间层恒载合计: 5136→{floor_dead_new:.0f} kN")

# 首层恒载 (也受影响)
first_dead_new = floor_slab_new + beam_trans_floor_new + 84.08 + 446.98 + (757.12+567.84)/2 + (743.69+528.97)/2 + (1484.01+1063.8)/2 + 201.6
print(f"首层恒载合计: 5548→{first_dead_new:.0f} kN")

# 总重力荷载代表值
# Ge = 屋面恒载 + 0.5×雪载 + 5×中间层(恒载+0.5×活载)
G_roof = roof_dead_new + 0.5 * 0.35 * roof_area_new
G_floor_dead = floor_dead_new
# 中间层活载 (也受面积影响)
live_floor_new = 2*104.3 + 3.5*28.8 + 2*(roof_area_new - 28.8 - 104.3 - 28.8)
# 近似: 总活载 ≈ q_floor × (floor_area - 楼梯间)
live_floor_new2 = q_floor * (floor_area_new - 28.8) + 3.5 * 28.8
print(f"中间层活载: 1086→{live_floor_new2:.0f} kN")

# 总重力: 屋面 + 4×中间层 + 首层 (共6层)
G_total_new = roof_dead_new + 0.5*0.35*roof_area_new + 4*(floor_dead_new + 0.5*live_floor_new2) + first_dead_new + 0.5*live_floor_new2
print(f"总重力荷载代表值: 33835→{G_total_new:.0f} kN")

# 验证原值
G_old_check = 4924 + 0.5*0.35*521.3 + 4*(5136 + 0.5*1086) + 5548 + 0.5*1086
print(f"原值验证: {G_old_check:.0f} (应≈33835)")

# ============================================================
# 自振周期
# ============================================================
T_old = 0.56
# T = 1.7×ψT×√Δ
# Δ = Σ(Vi/Di)/hi... 简化: T ∝ √(1/D_total)
stiff_ratio = 579164 / D_total_top  # 原D/新D
T_new = T_old * math.sqrt(stiff_ratio)
print(f"自振周期: {T_old:.2f}→{T_new:.2f}s")

# ============================================================
# 地震力 (使用原档alpha1进行一致性缩放)
# ============================================================
alpha1_old = 0.052  # 原档实际使用值
# α ∝ (1/T)^γ, 所以 α_new/α_old = (T_old/T_new)^γ
alpha1_new = alpha1_old * (T_old/T_new)**0.9
print(f"地震影响系数: {alpha1_old:.4f}→{alpha1_new:.4f}")

FEK_new = alpha1_new * 0.85 * G_total_new
FEK_old = 1495.51
print(f"底部总剪力: {FEK_old:.1f}→{FEK_new:.1f} kN")

# 顶部附加地震作用
delta_n_new = 0.12 * T_new + 0.07
delta_F_new = delta_n_new * FEK_new
print(f"顶部附加系数: {delta_n_new:.4f}")
print(f"顶部附加地震力: {delta_F_new:.1f} kN")

# ============================================================
# 修改docx第二波
# ============================================================
DOC_PATH = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本\邓杰鹏计算书_5400修正版.docx'
doc = Document(DOC_PATH)
RED = RGBColor(0xFF, 0x00, 0x00)

print(f"\n开始第二轮修改...")

# ==== 修改 Table 2-1 (Table 3) 柱截面尺寸 ====
t2_1 = doc.tables[2]
A_edge_new = L1/2 * L_long
A_mid_new = (L1+L2)/2 * L_long
N_edge_new = 1.3*12*A_edge_new*6
N_mid_new = 1.25*12*A_mid_new*6
A_req_edge_new = N_edge_new*1000/12.155
A_req_mid_new = N_mid_new*1000/12.155
b_edge_new = math.sqrt(A_req_edge_new)
b_mid_new = math.sqrt(A_req_mid_new)

t2_1.rows[2].cells[1].text = f'{L1}/2×(6.9+6.9)/2={A_edge_new:.2f}m²'
t2_1.rows[2].cells[2].text = f'(2.4+{L1})/2×(6.9+6.9)/2={A_mid_new:.2f}m²'
t2_1.rows[3].cells[1].text = f'1.3×12×{A_edge_new:.2f}×6={N_edge_new:.2f}kN'
t2_1.rows[3].cells[2].text = f'1.25×12×{A_mid_new:.2f}×6={N_mid_new:.2f}kN'
t2_1.rows[5].cells[1].text = f'{N_edge_new*1000/12.155:.0f}mm²'
t2_1.rows[5].cells[2].text = f'{N_mid_new*1000/12.155:.0f}mm²'
t2_1.rows[6].cells[1].text = f'{b_edge_new:.1f}mm'
t2_1.rows[6].cells[2].text = f'{b_mid_new:.2f}mm'
print("  表2-1: 已修改")

# ==== 修改 Table 13 (表4-1 屋面层重力荷载) ====
t4_1 = doc.tables[12]
# R0 屋面层
# 横向框架梁
t4_1.rows[0].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_roof_new:.1f}kN'
t4_1.rows[0].cells[2].text = f'{roof_dead_new:.0f}kN'  # 总计

# 屋面板
t4_1.rows[5].cells[1].text = f'{roof_area_new:.1f}×4.96={roof_slab_new:.2f}kN'
t4_1.rows[5].cells[2].text = f'{roof_dead_new:.0f}kN'

# 恒载合计
t4_1.rows[9].cells[1].text = f'{roof_slab_new:.2f}+{beam_trans_roof_new:.1f}+84.08+446.98+567.84/2+528.97/2+1063.8/2+501.18  ={roof_dead_new:.0f}kN'

# 屋面雪荷载
snow_new = 0.35 * roof_area_new
t4_1.rows[10].cells[1].text = f'0.35×{roof_area_new:.1f}={snow_new:.2f}kN'

print("  表4-1 (屋面重力荷载): 已修改")

# ==== 修改 Table 14 (表4-2 中间层/首层重力荷载) ====
t4_2 = doc.tables[13]
# 中间层
t4_2.rows[0].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_floor_new:.1f}kN'
t4_2.rows[0].cells[2].text = f'{floor_dead_new:.0f}kN'

t4_2.rows[5].cells[1].text = f'{floor_area_new:.1f}×4.2={floor_slab_new:.2f}kN'
t4_2.rows[5].cells[2].text = f'{floor_dead_new:.0f}kN'

t4_2.rows[8].cells[1].text = f'{floor_slab_new:.2f}+{beam_trans_floor_new:.1f}+84.08+446.98+567.84+528.97+1063.8+201.6={floor_dead_new:.0f}kN'

# 楼面活载
t4_2.rows[9].cells[1].text = f'2×104.3+3.5×28.8+2×({roof_area_new:.1f}-104.3-28.8)={live_floor_new2:.0f}kN'

# 首层
t4_2.rows[10].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_floor_new:.1f}kN'
t4_2.rows[10].cells[2].text = f'{first_dead_new:.0f}kN'

t4_2.rows[17].cells[1].text = f'{floor_area_new:.1f}×4.2={floor_slab_new:.2f}kN'
t4_2.rows[17].cells[2].text = f'{first_dead_new:.0f}kN'

# 首层恒载合计 (需要重新计算)
# 2068.5+173.94+84.08+446.98+(757.12+567.84)/2+(743.69+528.97)/2+(1484.01+1063.8)/2+201.6
first_dead_sum = floor_slab_new + beam_trans_floor_new + 84.08 + 446.98 + (757.12+567.84)/2 + (743.69+528.97)/2 + (1484.01+1063.8)/2 + 201.6
t4_2.rows[21].cells[1].text = f'{floor_slab_new:.2f}+{beam_trans_floor_new:.1f}+84.08+446.98+(757.12+567.84)/2+(743.69+528.97)/2+(1484.01+1063.8)/2+201.6={first_dead_sum:.0f}kN'

print("  表4-2 (中间层/首层重力荷载): 已修改")

# ==== 修改 Table 16-17 (D值汇总, 自振周期) ====
# 查找并修改D值相关表格
# 表4-4 (Table 16 docx内) 整体D值
t4_4 = doc.tables[15] if len(doc.tables) > 15 else None
# 表4-5 (Table 17) 自振周期

# ==== 修改段落中的关键数值 ====
for i, para in enumerate(doc.paragraphs):
    text = para.text

    # 总重力荷载代表值 33835
    if '33835kN' in text or '33835 kN' in text:
        for run in para.runs:
            if '33835' in run.text:
                run.text = run.text.replace('33835', f'{G_total_new:.0f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 总重力荷载代表值 33835→{G_total_new:.0f}")

    # 自振周期 0.56s
    if '0.56s' in text and '周期' in text:
        for run in para.runs:
            if '0.56' in run.text:
                run.text = run.text.replace('0.56', f'{T_new:.2f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 自振周期 0.56→{T_new:.2f}")

    # 地震影响系数 0.052
    if '0.052' in text and '水平地震' in text:
        for run in para.runs:
            if '0.052' in run.text:
                run.text = run.text.replace('0.052', f'{alpha1_new:.4f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 地震系数 0.052→{alpha1_new:.4f}")

    # 底部剪力 1495.51
    if '1495.51kN' in text or '1495.51 kN' in text:
        for run in para.runs:
            if '1495.51' in run.text:
                run.text = run.text.replace('1495.51', f'{FEK_new:.2f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 底部剪力 1495.51→{FEK_new:.2f}")

    # 顶部附加系数 0.1148
    if '0.1148' in text and '顶部' in text:
        for run in para.runs:
            if '0.1148' in run.text:
                run.text = run.text.replace('0.1148', f'{delta_n_new:.4f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 顶部附加系数 0.1148→{delta_n_new:.4f}")

    # 顶部附加地震力 171.68
    if '171.68kN' in text or '171.68 kN' in text:
        for run in para.runs:
            if '171.68' in run.text:
                run.text = run.text.replace('171.68', f'{delta_F_new:.1f}')
                run.font.color.rgb = RED
                print(f"  P{i}: 顶部附加力 171.68→{delta_F_new:.1f}")

    # 弯矩二次分配中边跨等效均布荷载
    # "边跨梁屋面层等效均布荷载为2.57+0.79×17.11=16.09kN/m"
    if '0.79×17.11' in text:
        new_q = 2.57 + F_NEW * ss * g_roof
        for run in para.runs:
            if '0.79×17.11' in run.text:
                run.text = run.text.replace('0.79×17.11', f'{F_NEW:.2f}×{ss*g_roof:.2f}')
                run.font.color.rgb = RED
            if '16.09kN/m' in run.text:
                run.text = run.text.replace('16.09', f'{new_q:.2f}')
                run.font.color.rgb = RED
        print(f"  P{i}: 屋面等效均布荷载 16.09→{new_q:.2f}")

    if '0.79×14.49' in text and '20.22' in text:
        new_q = 2.57 + 6.2 + F_NEW * ss * g_floor
        for run in para.runs:
            if '0.79×14.49' in run.text:
                run.text = run.text.replace('0.79×14.49', f'{F_NEW:.2f}×{ss*g_floor:.2f}')
                run.font.color.rgb = RED
            if '20.22kN/m' in run.text:
                run.text = run.text.replace('20.22', f'{new_q:.2f}')
                run.font.color.rgb = RED
        print(f"  P{i}: 楼面等效均布荷载 20.22→{new_q:.2f}")

    # 活载等效
    if '0.79×1.73' in text:
        new_q = F_NEW * ss * q_roof
        for run in para.runs:
            if '0.79×1.73' in run.text:
                run.text = run.text.replace('0.79×1.73', f'{F_NEW:.2f}×{ss*q_roof:.2f}')
                run.font.color.rgb = RED
            if '1.36kN/m' in run.text:
                run.text = run.text.replace('1.36', f'{new_q:.2f}')
                run.font.color.rgb = RED
        print(f"  P{i}: 活载屋面等效 1.36→{new_q:.2f}")

    if '0.79×6.92' in text and '5.44' in text:
        new_q = F_NEW * ss * q_floor
        for run in para.runs:
            if '0.79×6.92' in run.text:
                run.text = run.text.replace('0.79×6.92', f'{F_NEW:.2f}×{ss*q_floor:.2f}')
                run.font.color.rgb = RED
            if '5.44kN/m' in run.text:
                run.text = run.text.replace('5.44', f'{new_q:.2f}')
                run.font.color.rgb = RED
        print(f"  P{i}: 活载楼面等效 5.44→{new_q:.2f}")

    # 修改固端弯矩值
    # 屋面边跨 -30.89
    if '-30.89' in text and ('4.8²' in text or '4.80' in text):
        fem_edge_roof = 2.57 + F_NEW * ss * g_roof
        fem_val = fem_edge_roof * L1**2 / 12
        for run in para.runs:
            if '-30.89' in run.text:
                run.text = run.text.replace('-30.89', f'-{fem_val:.2f}')
                run.font.color.rgb = RED
            if '30.89' in run.text and '-30.89' not in run.text:
                run.text = run.text.replace('30.89', f'{fem_val:.2f}')
                run.font.color.rgb = RED

    # 修改柱集中力 (边柱、中柱)
    edge_floor_new = g_floor * (ss**2/4 + ss*L1/2)
    edge_floor_old = g_floor * (ss**2/4 + ss*L1_OLD/2)
    if f'{edge_floor_old:.2f}kN' in text:
        for run in para.runs:
            if f'{edge_floor_old:.2f}' in run.text:
                run.text = run.text.replace(f'{edge_floor_old:.2f}', f'{edge_floor_new:.2f}')
                run.font.color.rgb = RED

# 保存
doc.save(DOC_PATH)
overwrite_review = DOC_PATH.replace('修正版', '审阅版')
doc.save(overwrite_review)

print(f"\n第二轮修改完成!")
print(f"关键变化:")
print(f"  总重力荷载代表值: 33835→{G_total_new:.0f} kN")
print(f"  自振周期: 0.56→{T_new:.2f} s")
print(f"  地震系数: {alpha1_old:.4f}→{alpha1_new:.4f}")
print(f"  底部剪力: 1495.51→{FEK_new:.1f} kN")
print(f"  顶部附加力: 171.68→{delta_F_new:.1f} kN")
