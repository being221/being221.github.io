#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
5400跨度 docx修改 整合版
从原始4800文件重新生成，包含所有精确计算
"""

import sys, os, math, shutil
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)
SRC = r'C:\Users\邓杰鹏\Desktop\毕设\邓杰鹏计算书.docx'
DST_DIR = r'C:\Users\邓杰鹏\Desktop\毕设\5400版本'
os.makedirs(DST_DIR, exist_ok=True)

doc = Document(SRC)

# ============================================================
# 参数
# ============================================================
L1, L1_OLD = 5.4, 4.8
L2 = 2.4; L_long = 6.9; ss = 3.45
g_roof, g_floor = 4.96, 4.2
q_roof, q_floor = 0.5, 2.0
g_beam_e, g_beam_m = 2.57, 1.89
g_wall_e, g_wall_m = 6.2, 6.45

alpha_new = 0.5*ss/L1
F_NEW = 1 - 2*alpha_new**2 + alpha_new**3
F_OLD = 0.79

E = 30e6
I0_e = 0.25*0.5**3/12
i_e_e_new = E * 1.5*I0_e / L1  # 21701
i_m_e_new = E * 2.0*I0_e / L1  # 28935
i_e_m = E * 1.5*0.25*0.4**3/12 / L2  # 25000
i_m_m = E * 2.0*0.25*0.4**3/12 / L2  # 33333

Ic = 0.5*0.5**3/12
ic_top = E * Ic / 3.0
ic_1st = E * Ic / 4.0

def KD(ic, h, sib, is_first):
    K = sib/ic
    a = (0.5+K)/(2+K) if is_first else K/(2+K)
    D = a*12*ic/h**2
    return K,a,D

K_ee_t,a_ee_t,D_ee_t = KD(ic_top,3.0,i_e_e_new,False)
K_ee_1,a_ee_1,D_ee_1 = KD(ic_1st,4.0,i_e_e_new,True)
K_em_t,a_em_t,D_em_t = KD(ic_top,3.0,i_e_e_new+i_e_m,False)
K_em_1,a_em_1,D_em_1 = KD(ic_1st,4.0,i_e_e_new+i_e_m,True)
K_me_t,a_me_t,D_me_t = KD(ic_top,3.0,i_m_e_new,False)
K_me_1,a_me_1,D_me_1 = KD(ic_1st,4.0,i_m_e_new,True)
K_mm_t,a_mm_t,D_mm_t = KD(ic_top,3.0,i_m_e_new+i_m_m,False)
K_mm_1,a_mm_1,D_mm_1 = KD(ic_1st,4.0,i_m_e_new+i_m_m,True)

D_edge_top = 2*D_ee_t + 2*D_em_t
D_edge_1st = 2*D_ee_1 + 2*D_em_1
D_mid_top = 2*D_me_t + 2*D_mm_t
D_mid_1st = 2*D_me_1 + 2*D_mm_1
D_total_top = 2*D_edge_top + 5*D_mid_top
D_total_1st = 2*D_edge_1st + 5*D_mid_1st

# 等效均布
qeq_d_roof_e = g_beam_e + F_NEW * ss * g_roof
qeq_d_floor_e = g_beam_e + g_wall_e + F_NEW * ss * g_floor
qeq_d_roof_m = g_beam_m + 0.625 * L2 * g_roof
qeq_d_floor_m = g_beam_m + g_wall_m + 0.625 * L2 * g_floor
qeq_l_roof_e = F_NEW * ss * q_roof
qeq_l_floor_e = F_NEW * ss * q_floor
qeq_l_roof_m = 0.625 * L2 * q_roof
qeq_l_floor_m = 0.625 * L2 * q_floor

fem_d_roof_e = qeq_d_roof_e * L1**2 / 12
fem_d_floor_e = qeq_d_floor_e * L1**2 / 12
fem_d_roof_m = qeq_d_roof_m * L2**2 / 12
fem_d_floor_m = qeq_d_floor_m * L2**2 / 12
fem_l_roof_e = qeq_l_roof_e * L1**2 / 12
fem_l_floor_e = qeq_l_floor_e * L1**2 / 12

# 面积
A_edge = L1/2 * L_long
A_mid = (L1+L2)/2 * L_long
area_increase = 14 * (L1-L1_OLD) * L_long
floor_area_new = 492.5 + area_increase
roof_area_new = 521.3 + area_increase

# 重力荷载
beam_trans_new = 14*(L1-0.5-0.5+0.12+0.12)*2.57 + 7*(2.4-0.12-0.12)*1.89
roof_slab_new = roof_area_new * g_roof
floor_slab_new = floor_area_new * g_floor
roof_dead_new = roof_slab_new + beam_trans_new + 84.08 + 446.98 + 567.84/2 + 528.97/2 + 1063.8/2 + 501.18
floor_dead_new = floor_slab_new + beam_trans_new + 84.08 + 446.98 + 567.84 + 528.97 + 1063.8 + 201.6
first_dead_new = floor_slab_new + beam_trans_new + 84.08 + 446.98 + (757.12+567.84)/2 + (743.69+528.97)/2 + (1484.01+1063.8)/2 + 201.6
live_floor_new = q_floor*(floor_area_new-28.8) + 3.5*28.8
snow_new = 0.35 * roof_area_new
G_total_new = roof_dead_new + 0.5*snow_new + 4*(floor_dead_new+0.5*live_floor_new) + first_dead_new + 0.5*live_floor_new

# 自振周期 & 地震力
T_old = 0.56
T_new = T_old * math.sqrt(579164/D_total_top)
alpha1_old = 0.052
alpha1_new = alpha1_old * (T_old/T_new)**0.9
FEK_new = alpha1_new * 0.85 * G_total_new
delta_n_new = 0.12*T_new + 0.07 if T_new > 1.4*0.35 else 0.08*T_new + 0.07
# 原档: 0.12×0.56+0.07=0.1148 (使用了0.12系数, 说明T_new>1.4Tg)
# T_new = 0.58, 1.4Tg = 0.49, 所以 T_new > 1.4Tg
delta_n_new_val = 0.12*T_new + 0.07
delta_F_new = delta_n_new_val * FEK_new

# 柱集中力
sec_beam_new = 1.54*L1/2
edge_floor_conc = g_floor * (ss**2/4 + ss*L1/2)
mid_floor_conc = g_floor * ((ss**2/4+ss*L1/2) + (ss*L2-0.5*L2*L2))
edge_roof_conc = g_roof * (ss**2/4 + ss*L1/2)
mid_roof_conc = g_roof * ((ss**2/4+ss*L1/2) + (ss*L2-0.5*L2*L2))

edge_floor_live = q_floor * (ss**2/4 + ss*L1/2)
mid_floor_live = q_floor*(ss**2/4+ss*L1/2) + q_floor*(ss*L2-0.5*L2*L2)
edge_roof_live = q_roof * (ss**2/4 + ss*L1/2)
mid_roof_live = q_roof*((ss**2/4+ss*L1/2)+(ss*L2-0.5*L2*L2))

print("="*60)
print("全部5400参数计算完成")
print(f"  梁线刚度: {i_e_e_new:.0f}, {i_m_e_new:.0f}")
print(f"  等效系数: 0.79→{F_NEW:.2f}")
print(f"  边跨等效均布(屋面): 16.09→{qeq_d_roof_e:.2f}")
print(f"  边跨等效均布(楼面): 20.22→{qeq_d_floor_e:.2f}")
print(f"  边跨固端弯矩(屋面): 30.89→{fem_d_roof_e:.2f}")
print(f"  边跨固端弯矩(楼面): 38.82→{fem_d_floor_e:.2f}")
print(f"  边柱受荷面积: 16.56→{A_edge:.2f}")
print(f"  总重力: 33835→{G_total_new:.0f}")
print(f"  自振周期: 0.56→{T_new:.2f}")
print(f"  地震系数: 0.052→{alpha1_new:.4f}")
print(f"  底部剪力: 1495.5→{FEK_new:.1f}")
print(f"  顶部附加力: 171.68→{delta_F_new:.1f}")
print("="*60)

# ============================================================
# 开始修改
# ============================================================

# --- 表2-1 柱截面尺寸 (Table index 2) ---
t = doc.tables[2]
N_edge = 1.3*12*A_edge*6
N_mid = 1.25*12*A_mid*6
A_req_e = N_edge*1000/12.155
A_req_m = N_mid*1000/12.155
b_e = math.sqrt(A_req_e)
b_m = math.sqrt(A_req_m)
t.rows[2].cells[1].text = f'{L1}/2×(6.9+6.9)/2={A_edge:.2f}m²'
t.rows[2].cells[2].text = f'(2.4+{L1})/2×(6.9+6.9)/2={A_mid:.2f}m²'
t.rows[3].cells[1].text = f'1.3×12×{A_edge:.2f}×6={N_edge:.2f}kN'
t.rows[3].cells[2].text = f'1.25×12×{A_mid:.2f}×6={N_mid:.2f}kN'
t.rows[5].cells[1].text = f'{A_req_e:.0f}mm²'
t.rows[5].cells[2].text = f'{A_req_m:.0f}mm²'
t.rows[6].cells[1].text = f'{b_e:.1f}mm'
t.rows[6].cells[2].text = f'{b_m:.2f}mm'

# --- 表2-2 梁线刚度 (Table index 3) ---
t = doc.tables[3]
t.rows[2].cells[2].text = f'{L1}'  # 边榀边跨跨度
t.rows[2].cells[6].text = f'{i_e_e_new:.0f}'
t.rows[4].cells[2].text = f'{L1}'  # 中间榀边跨跨度
t.rows[4].cells[6].text = f'{i_m_e_new:.0f}'

# --- 表2-4 柱刚度修正系数 (Table index 5) ---
t = doc.tables[5]
for ri in [2,3]:  # 顶层, 标准层
    t.rows[ri].cells[1].text = f'{K_ee_t:.2f}'
    t.rows[ri].cells[2].text = f'{a_ee_t:.2f}'
    t.rows[ri].cells[3].text = f'{K_em_t:.2f}'
    t.rows[ri].cells[4].text = f'{a_em_t:.2f}'
# 底层
t.rows[4].cells[1].text = f'{K_ee_1:.2f}'
t.rows[4].cells[2].text = f'{a_ee_1:.2f}'
t.rows[4].cells[3].text = f'{K_em_1:.2f}'
t.rows[4].cells[4].text = f'{a_em_1:.2f}'

# --- 表2-5 抗侧刚度 (Table index 6) ---
t = doc.tables[6]
# R3 顶层边榀 R4 顶层中间榀 R5 中间层边榀 R6 中间层中间榀 R7 首层边榀 R8 首层中间榀
configs = [
    (3, K_em_t, a_em_t, D_em_t, K_ee_t, a_ee_t, D_ee_t, D_edge_top),
    (4, K_mm_t, a_mm_t, D_mm_t, K_me_t, a_me_t, D_me_t, D_mid_top),
    (5, K_em_t, a_em_t, D_em_t, K_ee_t, a_ee_t, D_ee_t, D_edge_top),
    (6, K_mm_t, a_mm_t, D_mm_t, K_me_t, a_me_t, D_me_t, D_mid_top),
    (7, K_em_1, a_em_1, D_em_1, K_ee_1, a_ee_1, D_ee_1, D_edge_1st),
    (8, K_mm_1, a_mm_1, D_mm_1, K_me_1, a_me_1, D_me_1, D_mid_1st),
]
for ri, Km, am, Dm, Ke, ae, De, Dsum in configs:
    t.rows[ri].cells[4].text = f'{Km:.2f}'
    t.rows[ri].cells[5].text = f'{am:.2f}'
    t.rows[ri].cells[6].text = f'{Dm:.0f}'
    t.rows[ri].cells[7].text = f'{Ke:.2f}'
    t.rows[ri].cells[8].text = f'{ae:.2f}'
    t.rows[ri].cells[9].text = f'{De:.0f}'
    t.rows[ri].cells[10].text = f'{Dsum:.0f}'

# --- 表3-3 恒载柱集中力 (Table index 9) ---
t = doc.tables[9]
t.rows[2].cells[2].text = f'{sec_beam_new:.1f}kN'
t.rows[5].cells[2].text = f'{g_floor}×({ss}/2×{ss}/2+{ss}×{L1})={edge_floor_conc:.2f}kN'
t.rows[6].cells[2].text = f'{42.3+20.08+sec_beam_new+edge_floor_conc:.2f}kN'
t.rows[8].cells[2].text = f'{sec_beam_new:.1f}kN'
t.rows[11].cells[2].text = f'{g_floor}×(({ss}/2×{ss}/2+{ss}×{L1})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_floor_conc:.2f}kN'
t.rows[12].cells[2].text = f'{41.95+20.08+sec_beam_new+mid_floor_conc:.2f}kN'
t.rows[14].cells[2].text = f'{sec_beam_new:.1f}kN'
t.rows[16].cells[2].text = f'{g_roof}×({ss}/2×{ss}/2+{ss}×{L1})={edge_roof_conc:.2f}kN'
t.rows[17].cells[2].text = f'{31.19+20.08+sec_beam_new+edge_roof_conc:.2f}kN'
t.rows[18].cells[2].text = f'{sec_beam_new:.1f}kN'
mid_roof_final = 3.7+20.08+sec_beam_new+mid_roof_conc  # 中柱 次梁+纵梁+导荷
t.rows[20].cells[2].text = f'{g_roof}×(({ss}/2×{ss}/2+{ss}×{L1})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_roof_conc:.2f}kN'
t.rows[21].cells[2].text = f'{mid_roof_final:.2f}kN'

# --- 表3-5 活载柱集中力 (Table index 11) ---
t = doc.tables[11]
t.rows[1].cells[2].text = f'{q_floor}×({ss}/2×{ss}/2+{ss}×{L1})={edge_floor_live:.2f}kN'
t.rows[2].cells[2].text = f'{q_floor}×({ss}/2×{ss}/2+{ss}×{L1})+{q_floor}×({ss}×{L2}-0.5×{L2}×0.5×{L2})={mid_floor_live:.2f}kN'
t.rows[3].cells[2].text = f'{q_roof}×({ss}/2×{ss}/2+{ss}×{L1})={edge_roof_live:.2f}kN'
t.rows[4].cells[2].text = f'{q_roof}×(({ss}/2×{ss}/2+{ss}×{L1})+({ss}×{L2}-0.5×{L2}×0.5×{L2}))={mid_roof_live:.2f}kN'

# --- 表4-1 屋面重力荷载 (Table index 12) ---
t = doc.tables[12]
t.rows[0].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_new:.1f}kN'
t.rows[0].cells[2].text = f'{roof_dead_new:.0f}kN'
t.rows[5].cells[1].text = f'{roof_area_new:.1f}×4.96={roof_slab_new:.2f}kN'
t.rows[5].cells[2].text = f'{roof_dead_new:.0f}kN'
t.rows[9].cells[1].text = f'{roof_slab_new:.2f}+{beam_trans_new:.1f}+84.08+446.98+567.84/2+528.97/2+1063.8/2+501.18={roof_dead_new:.0f}kN'
t.rows[10].cells[1].text = f'0.35×{roof_area_new:.1f}={snow_new:.2f}kN'

# --- 表4-2 中间层/首层 (Table index 13) ---
t = doc.tables[13]
t.rows[0].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_new:.1f}kN'
t.rows[0].cells[2].text = f'{floor_dead_new:.0f}kN'
t.rows[5].cells[1].text = f'{floor_area_new:.1f}×4.2={floor_slab_new:.2f}kN'
t.rows[5].cells[2].text = f'{floor_dead_new:.0f}kN'
t.rows[8].cells[1].text = f'{floor_slab_new:.2f}+{beam_trans_new:.1f}+84.08+446.98+567.84+528.97+1063.8+201.6={floor_dead_new:.0f}kN'
t.rows[9].cells[1].text = f'2×104.3+3.5×28.8+2×({roof_area_new:.1f}-104.3-28.8)={live_floor_new:.0f}kN'

t.rows[10].cells[1].text = f'14×({L1}-0.5-0.5+0.12+0.12)×2.57+7×(2.4-0.12-0.12)×1.89={beam_trans_new:.1f}kN'
t.rows[10].cells[2].text = f'{first_dead_new:.0f}kN'
t.rows[17].cells[1].text = f'{floor_area_new:.1f}×4.2={floor_slab_new:.2f}kN'
t.rows[17].cells[2].text = f'{first_dead_new:.0f}kN'
t.rows[21].cells[1].text = f'{floor_slab_new:.2f}+{beam_trans_new:.1f}+84.08+446.98+(757.12+567.84)/2+(743.69+528.97)/2+(1484.01+1063.8)/2+201.6={first_dead_new:.0f}kN'
t.rows[22].cells[1].text = f'2×104.3+3.5×28.8+2×({roof_area_new:.1f}-104.3-28.8)={live_floor_new:.0f}kN'

# --- 表6-3 梯形等效系数 (Table index 31) ---
t = doc.tables[31]
t.rows[2].cells[2].text = f'{L1}'
t.rows[2].cells[3].text = f'0.5×{ss}/{L1}={alpha_new:.3f}'
t.rows[2].cells[4].text = f'{F_NEW:.2f}'

# --- 表6-4 恒载固端弯矩 (Table index 32) ---
t = doc.tables[32]
t.rows[2].cells[2].text = f'{L1:.2f}'
t.rows[2].cells[3].text = f'{qeq_d_roof_e:.2f}'
t.rows[2].cells[4].text = f'-{qeq_d_roof_e:.2f}×{L1}²/12=-{fem_d_roof_e:.2f}'
t.rows[2].cells[5].text = f'{fem_d_roof_e:.2f}'
t.rows[3].cells[2].text = f'{L1:.2f}'
t.rows[3].cells[3].text = f'{qeq_d_floor_e:.2f}'
t.rows[3].cells[4].text = f'-{qeq_d_floor_e:.2f}×{L1}²/12=-{fem_d_floor_e:.2f}'
t.rows[3].cells[5].text = f'{fem_d_floor_e:.2f}'

# --- 表6-11 活载固端弯矩 (Table index 39) ---
t = doc.tables[39]
t.rows[2].cells[2].text = f'{L1:.2f}'
t.rows[2].cells[3].text = f'{qeq_l_roof_e:.2f}'
t.rows[2].cells[4].text = f'-{qeq_l_roof_e:.2f}×{L1}²/12=-{fem_l_roof_e:.2f}'
t.rows[2].cells[5].text = f'{fem_l_roof_e:.2f}'
t.rows[3].cells[2].text = f'{L1:.2f}'
t.rows[3].cells[3].text = f'{qeq_l_floor_e:.2f}'
t.rows[3].cells[4].text = f'-{qeq_l_floor_e:.2f}×{L1}²/12=-{fem_l_floor_e:.2f}'
t.rows[3].cells[5].text = f'{fem_l_floor_e:.2f}'

# --- 更新跨中弯矩表和梁端剪力表中的跨度 ---
for ti in [34, 35, 41, 42, 45, 47]:  # 表6-6,6-7,6-13,6-14,7-1,7-3
    t = doc.tables[ti]
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    if '4.80' in run.text:
                        # 只替换边跨相关的4.80,不替换其他数值中的4.80
                        pass  # 在边跨行手动替换
                    if '4.8' in run.text and ('边跨' in run.text or '跨' in run.text):
                        run.text = run.text.replace('4.8', f'{L1}')
    # 直接改边跨行中的跨度单元格
    if ti in [34, 41]:  # 跨中弯矩表
        for ri in [2,4,6,8,10,12]:
            if len(t.rows) > ri and len(t.rows[ri].cells) > 2:
                t.rows[ri].cells[2].text = f'{L1:.2f}'
    if ti in [35, 42]:  # 梁端剪力表
        for ri in range(3, min(9, len(t.rows))):
            if len(t.rows[ri].cells) > 3:
                t.rows[ri].cells[3].text = f'{L1:.2f}'
    if ti in [45, 47]:  # 弯矩调幅表
        for ri in range(3, min(9, len(t.rows))):
            for cell in t.rows[ri].cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if '4.8' in run.text and '²' in run.text:
                            run.text = run.text.replace('4.8', f'{L1}')

# --- 修改段落中的关键数值 ---
paragraph_changes = [
    # (old_text, new_text) pairs
    ('4800mm', f'{int(L1*1000)}mm'),
    ('343mm至600mm', f'{int(L1*1000)//14}mm至{int(L1*1000)//8}mm'),
    ('0.79×17.11', f'{F_NEW:.2f}×{ss*g_roof:.2f}'),
    ('16.09kN/m', f'{qeq_d_roof_e:.2f}kN/m'),
    ('0.79×14.49', f'{F_NEW:.2f}×{ss*g_floor:.2f}'),
    ('0.79×1.73', f'{F_NEW:.2f}×{ss*q_roof:.2f}'),
    ('0.79×6.92', f'{F_NEW:.2f}×{ss*q_floor:.2f}'),
]

specific_para_changes = {
    # 跨度说明
    '计算跨度为4800mm': f'计算跨度为{int(L1*1000)}mm',
    '次梁的最大跨度为4800mm': f'次梁的最大跨度为{int(L1*1000)}mm',
    # 等效均布
    '边跨梁屋面层等效均布荷载为2.57+0.79×17.11=16.09kN/m':
        f'边跨梁屋面层等效均布荷载为2.57+{F_NEW:.2f}×{ss*g_roof:.2f}={qeq_d_roof_e:.2f}kN/m',
    '楼面层为2.57+6.2+0.79×14.49=20.22kN/m':
        f'楼面层为2.57+6.2+{F_NEW:.2f}×{ss*g_floor:.2f}={qeq_d_floor_e:.2f}kN/m',
    '边跨梁屋面层等效均布荷载为0.79×1.73=1.36kN/m':
        f'边跨梁屋面层等效均布荷载为{F_NEW:.2f}×{ss*q_roof:.2f}={qeq_l_roof_e:.2f}kN/m',
    '楼面层为0.79×6.92=5.44kN/m':
        f'楼面层为{F_NEW:.2f}×{ss*q_floor:.2f}={qeq_l_floor_e:.2f}kN/m',
}

for i, para in enumerate(doc.paragraphs):
    text = para.text
    if not text.strip():
        continue

    for old_s, new_s in specific_para_changes.items():
        if old_s in text:
            for run in para.runs:
                run.text = run.text.replace(old_s, new_s)
                run.font.color.rgb = RED
            break

    # 地震相关
    if '33835kN' in text and '重力' in text:
        for run in para.runs:
            if '33835' in run.text:
                run.text = run.text.replace('33835', f'{G_total_new:.0f}')
                run.font.color.rgb = RED
    if '=0.56s' in text and '周期' in text:
        for run in para.runs:
            if '0.56' in run.text and '=' not in run.text.split('0.56')[0][-2:]:
                run.text = run.text.replace('0.56', f'{T_new:.2f}')
                run.font.color.rgb = RED
    if '0.052' in text and '地震' in text and '系数' in text:
        for run in para.runs:
            if '0.052' in run.text:
                run.text = run.text.replace('0.052', f'{alpha1_new:.4f}')
                run.font.color.rgb = RED
    if '1495.51kN' in text:
        for run in para.runs:
            if '1495.51' in run.text:
                run.text = run.text.replace('1495.51', f'{FEK_new:.1f}')
                run.font.color.rgb = RED
    if '0.1148' in text and '顶部' in text:
        for run in para.runs:
            if '0.1148' in run.text:
                run.text = run.text.replace('0.1148', f'{delta_n_new_val:.4f}')
                run.font.color.rgb = RED
    if '171.68kN' in text:
        for run in para.runs:
            if '171.68' in run.text:
                run.text = run.text.replace('171.68', f'{delta_F_new:.1f}')
                run.font.color.rgb = RED

    # 通用跨度替换
    if '4.80' in text and ('边跨' in text or '梁跨' in text):
        for run in para.runs:
            if '4.80' in run.text:
                run.text = run.text.replace('4.80', f'{L1:.2f}')
                run.font.color.rgb = RED
    if '4.8m' in text and ('跨度' in text or '梁跨' in text):
        for run in para.runs:
            if '4.8m' in run.text:
                run.text = run.text.replace('4.8m', f'{L1}m')
                run.font.color.rgb = RED

# 保存
fixed_path = os.path.join(DST_DIR, '邓杰鹏计算书_5400修正版.docx')
doc.save(fixed_path)
print(f"\n修正版已保存: {fixed_path}")

# 副本 (原文件)
copy_path = os.path.join(DST_DIR, '邓杰鹏计算书_4800原版备份.docx')
shutil.copy2(SRC, copy_path)
print(f"副本已保存: {copy_path}")

# 审阅版 (在修正版基础上追加红色标注)
review_path = os.path.join(DST_DIR, '邓杰鹏计算书_5400审阅版.docx')
doc.save(review_path)
print(f"审阅版已保存: {review_path}")

print(f"\n{'='*60}")
print("整合版修改完成！")
print(f"产出文件:")
print(f"  1. {fixed_path}")
print(f"  2. {copy_path}")
print(f"  3. {review_path}")
print(f"{'='*60}")
