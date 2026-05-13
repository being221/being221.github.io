#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
完整处理张明宇计算书：AB/CD轴 4.8m→5.4m 所有变化 + 级联 + 黄色标注
修复版：正确单元格索引 + 防振荡
"""
import docx
import re
import sys
import io
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = r'C:\Users\邓杰鹏\Downloads\张明宇\张明宇\张明宇计算书_backup.docx'
OUT = r'C:\Users\邓杰鹏\Downloads\张明宇\张明宇\张明宇计算书_审阅版_v2.docx'

print('加载文档...')
doc = docx.Document(SRC)

# ============================================================
# 工具函数
# ============================================================
def replace_4800(text):
    return re.sub(r'(?<!\d)4800(?!\d)', '5400', text)

def replace_480(text):
    return re.sub(r'(?<!\d)4\.80(?!\d)', '5.40', text)

def replace_48(text):
    return re.sub(r'(?<!\d)4\.8(?![\d.])', '5.4', text)

def to_py(expr):
    s = expr.strip()
    s = s.replace('×', '*').replace('÷', '/').replace('²', '**2').replace('³', '**3')
    s = s.replace('（', '(').replace('）', ')')
    s = re.sub(r'(?<!\d)\.(\d)', r'0.\1', s)
    s = re.sub(r'\)(\d)', r')*\1', s)
    s = re.sub(r'(\d)\(', r'\1*(', s)
    return s

def eval_expr(expr):
    try:
        py = to_py(expr)
        if py.count('(') != py.count(')'):
            return None
        result = eval(py)
        if isinstance(result, (int, float)) and abs(result) < 1e100:
            r = round(result, 2)
            if r == int(r):
                return str(int(r))
            else:
                return f'{r:.2f}'
    except Exception:
        pass
    return None

FORMULA_RE = re.compile(
    r'([-(\d][\d\s\.\+\-\×\÷\*\/\²\³\(\)]*[\d\)²³])\s*=\s*(-?[\d.]+)([a-zA-Z/%・²³kNSm・]*)'
)

def _add_highlight(run, color):
    run.element.get_or_add_rPr()
    for existing in run.element.rPr.findall(qn('w:shd')):
        run.element.rPr.remove(existing)
    shading_el = run.element.rPr.makeelement(qn('w:shd'), {
        qn('w:fill'): color, qn('w:val'): 'clear'
    })
    run.element.rPr.append(shading_el)

# ============================================================
# 段落/单元格文本操作
# ============================================================
def get_para_text(p):
    return ''.join(run.text for run in p.runs)

def set_para_text(p, text):
    if p.runs:
        for i, run in enumerate(p.runs):
            run.text = text if i == 0 else ''
    else:
        p.add_run(text)

def get_cell_text(cell):
    return '\n'.join(get_para_text(p) for p in cell.paragraphs)

def set_cell_text(cell, text):
    lines = text.split('\n')
    for i, p in enumerate(cell.paragraphs):
        set_para_text(p, lines[i] if i < len(lines) else '')

# ============================================================
# 存储原始文本
# ============================================================
print('存储原始文本...')
orig_para_texts = [get_para_text(p) for p in doc.paragraphs]
orig_cell_texts = {}
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            orig_cell_texts[(ti, ri, ci)] = get_cell_text(cell)

# ============================================================
# 阶段1: 直接替换
# ============================================================
print('\n=== 阶段1: 直接替换 4.8→5.4, 4800→5400 ===')
span_count = 0

for p in doc.paragraphs:
    old = get_para_text(p)
    new = replace_48(replace_480(replace_4800(old)))
    if new != old:
        span_count += 1
        set_para_text(p, new)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                old = get_para_text(p)
                new = replace_48(replace_480(replace_4800(old)))
                if new != old:
                    span_count += 1
                    set_para_text(p, new)

print(f'  直接替换: {span_count} 处')

# ============================================================
# 阶段2: 上下文替换
# ============================================================
print('\n=== 阶段2: 上下文替换 ===')
ctx_count = 0

# --- T9 (表3-3) ---
t9 = doc.tables[9]
t9_targets = [
    (5, 3, '4.2×(3.45/2×3.45/2+3.45×2.4)=47.27kN',
     '4.2×(3.45/2×3.45/2+3.45×2.7)=47.27kN', 'T9 R5'),
    (11, 3, '4.2×((3.45/2×3.45/2+3.45×2.4)+(3.45×2.4-0.5×2.4×0.5×2.4))=76kN',
     '4.2×((3.45/2×3.45/2+3.45×2.7)+(3.45×2.4-0.5×2.4×0.5×2.4))=76kN', 'T9 R11'),
    (16, 3, '4.96×(3.45/2×3.45/2+3.45×2.4)=55.83kN',
     '4.96×(3.45/2×3.45/2+3.45×2.7)=55.83kN', 'T9 R16'),
    (20, 3, '4.96×((3.45/2×3.45/2+3.45×2.4)+(3.45×2.4-0.5×2.4×0.5×2.4))=89.75kN',
     '4.96×((3.45/2×3.45/2+3.45×2.7)+(3.45×2.4-0.5×2.4×0.5×2.4))=89.75kN', 'T9 R20'),
]

for ri, ci, old_pattern, new_pattern, label in t9_targets:
    txt = get_cell_text(t9.rows[ri].cells[ci])
    if txt == old_pattern:
        set_cell_text(t9.rows[ri].cells[ci], new_pattern)
        ctx_count += 1
        print(f'  {label}: 3.45×2.4→3.45×2.7')
    else:
        print(f'  {label}: 未匹配! 实际=[{txt}]')

# --- T11 (表3-5) ---
t11 = doc.tables[11]
t11_targets = [
    (1, 3, '2×(3.45/2×3.45/2+3.45×2.4)=22.51kN',
     '2×(3.45/2×3.45/2+3.45×2.7)=22.51kN', 'T11 R1'),
    (2, 3, '2×(3.45/2×3.45/2+3.45×2.4)+2×(3.45×2.4-0.5×2.4×0.5×2.4)=36.19kN',
     '2×(3.45/2×3.45/2+3.45×2.7)+2×(3.45×2.4-0.5×2.4×0.5×2.4)=36.19kN', 'T11 R2'),
    (3, 3, '0.5×(3.45/2×3.45/2+3.45×2.4)=5.63kN',
     '0.5×(3.45/2×3.45/2+3.45×2.7)=5.63kN', 'T11 R3'),
    (4, 3, '0.5×((3.45/2×3.45/2+3.45×2.4)+(3.45×2.4-0.5×2.4×0.5×2.4))=9.05kN',
     '0.5×((3.45/2×3.45/2+3.45×2.7)+(3.45×2.4-0.5×2.4×0.5×2.4))=9.05kN', 'T11 R4'),
]

for ri, ci, old_pattern, new_pattern, label in t11_targets:
    txt = get_cell_text(t11.rows[ri].cells[ci])
    if txt == old_pattern:
        set_cell_text(t11.rows[ri].cells[ci], new_pattern)
        ctx_count += 1
        print(f'  {label}: 3.45×2.4→3.45×2.7')
    else:
        print(f'  {label}: 未匹配! 实际=[{txt}]')

# --- T12/T13 4.04→4.64 ---
for ti in [12, 13]:
    t = doc.tables[ti]
    for ri in range(len(t.rows)):
        for ci in range(len(t.rows[ri].cells)):
            txt = get_cell_text(t.rows[ri].cells[ci])
            if '4.04' in txt.replace('14.04', ''):  # 避免14.04之类的
                new_txt = txt.replace('4.04', '4.64')
                set_cell_text(t.rows[ri].cells[ci], new_txt)
                ctx_count += 1
                print(f'  T{ti} R{ri} C{ci}: 4.04→4.64')

# --- P582, P636, T31 0.79→0.83 ---
for pi in [582, 636]:
    p = doc.paragraphs[pi]
    txt = get_para_text(p)
    if '0.79' in txt:
        set_para_text(p, txt.replace('0.79', '0.83'))
        ctx_count += 1
        print(f'  P{pi}: 0.79→0.83')

t31 = doc.tables[31]
txt = get_cell_text(t31.rows[2].cells[4])
if '0.79' in txt:
    set_cell_text(t31.rows[2].cells[4], txt.replace('0.79', '0.83'))
    ctx_count += 1
    print(f'  T31 R2 C4: 0.79→0.83')

# --- T47 活载等效荷载更新 (1.36→1.44, 5.44→5.74) ---
# 只替换系数，公式结果由阶段3的公式重算自动更新
t47 = doc.tables[47]
# R3: 屋面层 1.36→1.44
txt = get_cell_text(t47.rows[3].cells[1])
if '1.36×' in txt:
    txt = txt.replace('1.36×', '1.44×')
    set_cell_text(t47.rows[3].cells[1], txt)
    ctx_count += 1
    print(f'  T47 R3: 1.36→1.44')

# R4-R8: 楼面层 5.44→5.74
for ri in [4, 5, 6, 7, 8]:
    txt = get_cell_text(t47.rows[ri].cells[1])
    if '5.44×' in txt:
        txt = txt.replace('5.44×', '5.74×')
        set_cell_text(t47.rows[ri].cells[1], txt)
        ctx_count += 1
        print(f'  T47 R{ri}: 5.44→5.74')

# --- T9 小计更新 ---
# 新值: R6=118.12, R12=146.51, R17=116.37, R21=119.13
t9_subtotals = {
    6: ('113.35kN', '118.12kN'),
    12: ('141.73kN', '146.51kN'),
    17: ('110.8kN', '116.37kN'),
    21: ('113.53kN', '119.13kN'),
}
for ri, (old, new) in t9_subtotals.items():
    txt = get_cell_text(t9.rows[ri].cells[3])
    if old in txt:
        set_cell_text(t9.rows[ri].cells[3], txt.replace(old, new))
        ctx_count += 1
        print(f'  T9 R{ri} 小计: {old}→{new}')
    else:
        print(f'  T9 R{ri} 小计: 未匹配 [{txt}]')

print(f'  上下文替换: {ctx_count} 处')

# ============================================================
# 阶段3: 公式重算 Pass 1 + 构建级联映射
# ============================================================
print('\n=== 阶段3: 公式重算 ===')

# 全局级联映射，仅在Pass 1构建
cascade_map = {}

def process_all_formulas(build_cascade=True):
    """遍历文档重算公式，build_cascade控制是否更新cascade_map"""
    total = 0
    for p in doc.paragraphs:
        txt = get_para_text(p)
        if '=' in txt:
            new_txt, n = recalc_formulas(txt, build_cascade)
            if n > 0:
                set_para_text(p, new_txt)
                total += n
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = get_para_text(p)
                    if '=' in txt:
                        new_txt, n = recalc_formulas(txt, build_cascade)
                        if n > 0:
                            set_para_text(p, new_txt)
                            total += n
    return total

def recalc_formulas(text, build_cascade):
    """重算公式，build_cascade=False时不更新cascade_map"""
    changes = 0
    def replacer(m):
        nonlocal changes
        expr = m.group(1)
        old_result = m.group(2)
        unit = m.group(3) or ''
        new_result = eval_expr(expr)
        if new_result is not None and new_result != old_result:
            changes += 1
            # 只在初始Pass构建级联映射，且只收录正向荷载/内力值
            if build_cascade and len(old_result) >= 5 and not old_result.startswith('-'):
                # 检查表达式是否因跨度变化而变化
                if re.search(r'(?<!\d)5\.40(?!\d)|(?<!\d)5\.4(?![\d.])|0\.83', expr):
                    cascade_map[old_result] = new_result
            return f'{expr} = {new_result}{unit}'
        return m.group(0)
    new_text = FORMULA_RE.sub(replacer, text)
    return new_text, changes

p1_changes = process_all_formulas(build_cascade=True)
print(f'  Pass 1 (初始公式重算): {p1_changes} 处变更')

# 过滤级联映射：
# 1. >=5字符, 不包含0.xxx系数, 不包含负数
cascade_map = {k: v for k, v in cascade_map.items()
               if len(k) >= 5 and not k.startswith('0.') and not k.startswith('-')}

# 2. 关键过滤：如果old_value出现在任何与跨度无关的公式结果中，则它是"歧义值"，不应级联
#    因为替换它会错误地修改不相关的公式
ambiguous = set()
for ov in list(cascade_map.keys()):
    for p in doc.paragraphs:
        for m in FORMULA_RE.finditer(get_para_text(p)):
            result = m.group(2)
            expr = m.group(1)
            if result == ov and not re.search(r'(?<!\d)5\.40(?!\d)|(?<!\d)5\.4(?![\d.])|0\.83', expr):
                ambiguous.add(ov)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for m in FORMULA_RE.finditer(get_para_text(p)):
                        result = m.group(2)
                        expr = m.group(1)
                        if result == ov and not re.search(r'(?<!\d)5\.40(?!\d)|(?<!\d)5\.4(?![\d.])|0\.83', expr):
                            ambiguous.add(ov)

cascade_map = {k: v for k, v in cascade_map.items() if k not in ambiguous}
if ambiguous:
    print(f'  去除歧义值: {ambiguous}')

# 手动添加T9小计和T47的级联映射
manual_cascade = {
    '113.35': '118.12',  # T9 R6 边柱楼面小计
    '141.73': '146.51',  # T9 R12 中柱楼面小计
    '110.80': '116.37',  # T37中使用的屋面边柱荷载(带.80)
    '113.53': '119.13',  # T9 R21 屋面中柱小计
    '1.36': '1.44',      # T47 活载屋面等效荷载
    '5.44': '5.74',      # T47 活载楼面等效荷载
}
# 只添加>=5字符的键
for k, v in manual_cascade.items():
    if len(k) >= 5:
        cascade_map[k] = v
print(f'  手动添加 {len([k for k in manual_cascade if len(k)>=5])} 条级联')
print(f'  级联映射条目(过滤后): {len(cascade_map)}')
for k, v in sorted(cascade_map.items(), key=lambda x: -len(x[0])):
    print(f'    {k} → {v}')

# ============================================================
# 阶段4: 级联 (冻结cascade_map)
# ============================================================
print('\n=== 阶段4: 级联 ===')

def cascade_replace():
    """用级联映射替换文档中的旧值(带数字边界检查，避免子串误伤)"""
    total = 0
    sorted_items = sorted(cascade_map.items(), key=lambda x: -len(x[0]))

    def replace_in_text(txt):
        nonlocal total
        new_txt = txt
        for old_v, new_v in sorted_items:
            # 使用数字边界：确保old_v不是更大数字的子串
            # old_v可能包含小数点，用regex确保前后不是数字
            pattern = re.compile(r'(?<!\d)' + re.escape(old_v) + r'(?!\d)')
            matches = pattern.findall(new_txt)
            if matches:
                new_txt = pattern.sub(new_v, new_txt)
                total += len(matches)
        return new_txt

    for p in doc.paragraphs:
        txt = get_para_text(p)
        new_txt = replace_in_text(txt)
        if new_txt != txt:
            set_para_text(p, new_txt)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    txt = get_para_text(p)
                    new_txt = replace_in_text(txt)
                    if new_txt != txt:
                        set_para_text(p, new_txt)
    return total

for pass_num in range(2, 20):
    replaced = cascade_replace()
    # 级联时不更新cascade_map
    recalc = process_all_formulas(build_cascade=False)
    print(f'  Pass {pass_num}: {replaced} 处值替换, {recalc} 处公式重算')

    if replaced == 0 and recalc == 0:
        print(f'  → 收敛!')
        break
    if pass_num >= 10:
        print(f'  → 达到上限，强制停止')
        break

# ============================================================
# 阶段5: 黄色标注
# ============================================================
print('\n=== 阶段5: 黄色标注 ===')

highlight_para = 0
highlight_cell = 0

for i, p in enumerate(doc.paragraphs):
    new_text = get_para_text(p)
    orig_text = orig_para_texts[i] if i < len(orig_para_texts) else ''
    if new_text != orig_text:
        for run in p.runs:
            _add_highlight(run, 'FFFF00')
        highlight_para += 1

for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            new_text = get_cell_text(cell)
            orig_text = orig_cell_texts.get((ti, ri, ci), '')
            if new_text != orig_text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _add_highlight(run, 'FFFF00')
                highlight_cell += 1

print(f'  段落标注: {highlight_para}')
print(f'  单元格标注: {highlight_cell}')
print(f'  合计: {highlight_para + highlight_cell}')

# ============================================================
doc.save(OUT)
print(f'\n保存到: {OUT}')
print('完成!')
